"""
考研学习助手 - Streamlit版
运行: streamlit run app.py
"""
import streamlit as st
import os
import json
import sqlite3
import math
import time
import base64
import socket
socket.setdefaulttimeout(90)
from pathlib import Path
from datetime import datetime, date, timedelta
import urllib.request
import urllib.error
import re
import traceback
import secrets
import io
import kaoyan_predict
from recommend import generate_recommendation
import extra_streamlit_components as stx
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt

# Monkey-patch Streamlit 的 CachedWidgetWarning 检测（CookieManager 在 @st.cache_resource 中需要）
import streamlit.elements.lib.policies as _policies
import streamlit.components.v1.custom_component as _cc
_policies.check_cache_replay_rules = lambda: None
_cc.check_cache_replay_rules = lambda: None

# ==================== 配置 ====================
st.set_page_config(page_title="考研学习助手", page_icon="", layout="wide", initial_sidebar_state="expanded")

# API配置
API_KEY = "sk-c4f69ncnuomnc8pprclmhlasndea7tdjvxeo49jno3bzxpa6"
API_BASE = "https://api.xiaomimimo.com/v1"
MODEL_NAME = "mimo-v2.5"
UMI_OCR_URL = os.environ.get("UMI_OCR_URL", "http://localhost:1224")

# 考纲分类：数学一独有 / 数学三独有
MATH1_ONLY = {"020", "065", "067", "068", "069", "083", "084", "085", "086", "087", "101", "102", "103", "104", "105", "106", "109"}
MATH3_ONLY = {"107", "110"}

DATA_DIR = Path("data/corpus")
DEMO_DATA_DIR = Path("data/corpus_demo")
REFERENCE_DIR = Path("data/reference")
MEMORY_DB = "data/memory.db"
EXPERIENCE_FILE = "agent_experience.md"

# ==================== CSS样式 ====================
st.markdown("""
<style>
    /* ── Design Tokens ── */
    :root {
        --primary: #4f46e5;
        --primary-hover: #4338ca;
        --primary-soft: #eef2ff;
        --primary-light: rgba(79,70,229,0.08);
        --text-main: #1e293b;
        --text-sub: #64748b;
        --text-muted: #94a3b8;
        --border: #e2e8f0;
        --bg-base: #f1f5f9;
        --bg-surface: #ffffff;
        --bg-elevated: #f8fafc;
        --success: #10b981;
        --success-soft: #ecfdf5;
        --warning: #f59e0b;
        --danger: #ef4444;
        --radius-sm: 8px;
        --radius-md: 12px;
        --radius-lg: 18px;
        --shadow-sm: 0 1px 2px rgba(0,0,0,0.04);
        --shadow-md: 0 1px 3px rgba(0,0,0,0.04), 0 4px 12px rgba(0,0,0,0.05);
        --shadow-lg: 0 2px 6px rgba(0,0,0,0.06), 0 8px 24px rgba(79,70,229,0.08);
    }

    /* ── Body / App Shell ── */
    section.main { background-color: #f1f5f9 !important; }
    div[data-testid="stAppViewContainer"] { background: #f1f5f9; }
    .stApp { background: #f1f5f9; }
    header[data-testid="stHeader"] { background: transparent; }

    /* ── Sidebar (Glass + Gradient) ── */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(255,255,255,0.96) 0%, rgba(241,245,249,0.92) 50%, rgba(238,242,255,0.94) 100%) !important;
        backdrop-filter: blur(16px) !important;
        -webkit-backdrop-filter: blur(16px) !important;
        border-right: 1px solid rgba(226,232,240,0.6) !important;
    }
    section[data-testid="stSidebar"] div[data-testid="stSidebarNav"] { display: none; }
    /* Hide Streamlit built-in toolbar & keyboard shortcuts */
    [data-testid="stHeader"] { display: none !important; }
    [data-testid="stSidebarHeader"] { display: none !important; }
    button[data-testid="baseButton-headerNoPadding"],
    button[data-testid="baseButton-header"] { display: none !important; }
    /* Prevent sidebar horizontal scroll */
    section[data-testid="stSidebar"] { overflow-x: hidden !important; }

    /* ═══════════════════════════════════════════
       TYPOGRAPHY — Modern system font stack
       ═══════════════════════════════════════════ */
    * {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Hiragino Sans GB", "Helvetica Neue", Helvetica, Arial, sans-serif !important;
    }
    h1, h2, h3, h4, h5, h6 {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif !important;
    }
    h1 { font-weight: 700 !important; color: #1e293b !important; letter-spacing: -0.02em; }
    h2 { font-weight: 650 !important; color: #1e293b !important; letter-spacing: -0.01em; }
    h3 { font-weight: 600 !important; color: #1e293b !important; }
    p, li, label, span, div {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif !important;
    }
    /* Code stays monospace */
    code, pre, .stCodeBlock {
        font-family: "Cascadia Code", "Fira Code", "JetBrains Mono", Consolas, "Source Code Pro", monospace !important;
    }
    /* Sidebar text */
    section[data-testid="stSidebar"] * {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif !important;
    }
    /* Buttons */
    button, .stButton > button {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif !important;
        font-weight: 500 !important;
    }
    /* Captions & hints */
    .st-caption, small, .small-text {
        font-family: "Segoe UI", "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif !important;
        font-weight: 400 !important; color: #64748b !important;
    }

    /* ── Sidebar: hide collapse arrow ── */
    button[data-testid="stSidebarCollapseButton"] { display: none !important; }
    section[data-testid="stSidebar"] { min-width: 250px !important; }

    /* ── Sidebar Brand ── */
    .sidebar-brand {
        display: flex; align-items: center; gap: 10px;
        padding: 0.75rem 0.5rem 0.25rem 0.5rem; margin-bottom: 0.25rem;
    }
    .sidebar-brand-icon { animation: brandFloat 3s ease-in-out infinite; display: inline-flex; }
    .sidebar-brand-icon svg { width: 26px; height: 26px; display: block; color: #6366f1; }
    .sidebar-brand-text {
        font-size: 1.15rem; font-weight: 700;
        font-family: "STKaiti", "KaiTi", "STSong", "SimSun", serif !important;
        background: linear-gradient(135deg, #4f46e5 0%, #8b5cf6 45%, #a855f7 100%);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        background-clip: text; letter-spacing: 0.04em;
    }
    @keyframes brandFloat {
        0%, 100% { transform: translateY(0); }
        50% { transform: translateY(-2px); }
    }

    /* ── Sidebar Divider ── */
    .sidebar-divider {
        height: 1px; margin: 0.5rem 0.25rem;
        background: linear-gradient(90deg, transparent, #e2e8f0 20%, #e2e8f0 80%, transparent);
    }

    /* ── Sidebar Section Labels ── */
    .sidebar-section-label {
        font-size: 0.7rem; font-weight: 600; color: #94a3b8;
        letter-spacing: 0.08em;
        padding: 0.5rem 14px 0.2rem 14px; margin-top: 0.25rem;
        font-family: "YouYuan", "Yuanti SC", "Microsoft YaHei", sans-serif !important;
    }

    /* ── Navigation Items ── */
    .nav-item {
        display: flex; align-items: center; gap: 10px;
        padding: 10px 14px; margin: 2px 4px; border-radius: 10px;
        font-size: 0.88rem; font-weight: 500; color: #64748b;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
        cursor: pointer; position: relative; overflow: hidden;
        text-decoration: none !important;
        font-family: "STKaiti", "KaiTi", "Microsoft YaHei", sans-serif !important;
        letter-spacing: 0.02em;
    }
    .nav-item:hover {
        background: #f1f5f9; color: #4f46e5;
        transform: translateX(3px); text-decoration: none !important;
    }
    .nav-item-active {
        background: linear-gradient(135deg, #eef2ff 0%, #f0e6ff 100%) !important;
        color: #4f46e5 !important; font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(79,70,229,0.12);
        cursor: default; position: relative;
        font-family: "STKaiti", "KaiTi", "Microsoft YaHei", sans-serif !important;
        letter-spacing: 0.03em !important;
    }
    /* Glow border on active item */
    .nav-item-active::after {
        content: ''; position: absolute; inset: 0; border-radius: 10px;
        padding: 1.5px;
        background: linear-gradient(135deg, #6366f1, #a855f7, #ec4899, #6366f1);
        background-size: 300% 300%;
        animation: glowFlow 3s ease-in-out infinite;
        -webkit-mask: linear-gradient(#fff 0 0) content-box, linear-gradient(#fff 0 0);
        -webkit-mask-composite: xor; mask-composite: exclude;
        pointer-events: none;
    }
    @keyframes glowFlow {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    .nav-dot { width: 8px; height: 8px; border-radius: 50%; flex-shrink: 0; }
    @keyframes navDotPulse {
        0%, 100% { opacity: 1; transform: scale(1); }
        50% { opacity: 0.6; transform: scale(1.3); }
    }

    /* ── Button Press Animation ── */
    .st-key-nav_hub button:active, .st-key-nav_main button:active,
    .st-key-nav_english button:active, .st-key-nav_checkin button:active,
    .st-key-nav_popularity button:active, .st-key-nav_material button:active,
    .st-key-nav_suggest button:active {
        transform: scale(0.96) !important;
        transition: transform 0.1s ease !important;
    }

    /* Sidebar link — clean anchor style */
    a.nav-item { text-decoration: none !important; }
    a.nav-item:hover { text-decoration: none !important; }

    /* ── Colored dots for each nav button ── */
    .st-key-nav_hub button::before,
    .st-key-nav_main button::before,
    .st-key-nav_english button::before,
    .st-key-nav_checkin button::before,
    .st-key-nav_popularity button::before,
    .st-key-nav_material button::before,
    .st-key-nav_suggest button::before {
        content: '';
        display: inline-block; width: 8px; height: 8px; border-radius: 50%;
        margin-right: 10px; flex-shrink: 0;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .st-key-nav_hub button::before        { background: #4f46e5; box-shadow: 0 0 6px rgba(79,70,229,0.4); }
    .st-key-nav_main button::before       { background: #3b82f6; box-shadow: 0 0 6px rgba(59,130,246,0.4); }
    .st-key-nav_english button::before    { background: #059669; box-shadow: 0 0 6px rgba(5,150,105,0.4); }
    .st-key-nav_checkin button::before    { background: #16a34a; box-shadow: 0 0 6px rgba(22,163,74,0.4); }
    .st-key-nav_popularity button::before { background: #db2777; box-shadow: 0 0 6px rgba(219,39,119,0.4); }
    .st-key-nav_material button::before   { background: #ca8a04; box-shadow: 0 0 6px rgba(202,138,4,0.4); }
    .st-key-nav_suggest button::before    { background: #0284c7; box-shadow: 0 0 6px rgba(2,132,199,0.4); }

    /* Active nav-item dot — glowing */
    .nav-item-active .nav-dot {
        background: #4f46e5 !important;
        box-shadow: 0 0 8px rgba(79,70,229,0.5);
        animation: navDotPulse 1.8s ease-in-out infinite;
    }

    /* Sidebar buttons — match nav-item font + style */
    .st-key-nav_hub button, .st-key-nav_main button, .st-key-nav_english button,
    .st-key-nav_checkin button, .st-key-nav_popularity button,
    .st-key-nav_material button, .st-key-nav_suggest button {
        background: transparent !important; color: #64748b !important;
        border: none !important; border-radius: 10px !important;
        font-weight: 500 !important; text-align: left !important;
        justify-content: flex-start !important; padding: 10px 14px !important;
        margin: 2px 4px !important;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
        width: calc(100% - 8px) !important;
        font-family: "STKaiti", "KaiTi", "Microsoft YaHei", sans-serif !important;
        letter-spacing: 0.02em !important; font-size: 0.88rem !important;
    }
    .st-key-nav_hub button:hover, .st-key-nav_main button:hover,
    .st-key-nav_english button:hover, .st-key-nav_checkin button:hover,
    .st-key-nav_popularity button:hover, .st-key-nav_material button:hover,
    .st-key-nav_suggest button:hover {
        background: #f1f5f9 !important; color: #4f46e5 !important;
        transform: translateX(3px) !important;
    }

    /* ── Sidebar User Card ── */
    .sidebar-user {
        display: flex; align-items: center; gap: 12px;
        padding: 0.75rem 0.5rem; margin: 0.25rem 0;
        background: linear-gradient(135deg, rgba(255,255,255,0.7) 0%, rgba(238,242,255,0.5) 100%);
        border-radius: 14px; border: 1px solid rgba(226,232,240,0.5);
    }
    .sidebar-avatar {
        width: 38px; height: 38px; border-radius: 50%;
        background: linear-gradient(135deg, #6366f1, #a855f7, #ec4899);
        background-size: 200% 200%;
        animation: avatarShine 3s ease-in-out infinite;
        color: #ffffff; display: flex; align-items: center; justify-content: center;
        font-weight: 700; font-size: 0.95rem; flex-shrink: 0;
        box-shadow: 0 2px 12px rgba(99,102,241,0.3);
    }
    @keyframes avatarShine {
        0%, 100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    .sidebar-user-info { display: flex; flex-direction: column; }
    .sidebar-username {
        font-weight: 700; font-size: 0.85rem; color: #1e293b; line-height: 1.3;
    }
    .sidebar-subtitle {
        font-size: 0.72rem; color: #64748b;
    }
    /* Stats row */
    .sidebar-stats {
        display: flex; gap: 12px; padding: 0.25rem 0.5rem 0.5rem 0.5rem;
        font-size: 0.7rem; color: #64748b;
    }
    .sidebar-stat { display: flex; align-items: center; gap: 4px; }
    .sidebar-stat strong { color: #4f46e5; }

    /* ── Logout Button ── */
    .st-key-sidebar_logout button {
        background: transparent !important; color: #94a3b8 !important;
        border: 1px solid #e2e8f0 !important; border-radius: 10px !important;
        font-size: 0.82rem !important; padding: 8px !important;
        transition: all 0.2s !important;
    }
    .st-key-sidebar_logout button:hover {
        color: #ef4444 !important; border-color: #fecaca !important;
        background: #fef2f2 !important;
    }

    /* ═══════════════════════════════════════
       BREATHING PAGE TRANSITIONS
       ═══════════════════════════════════════ */

    /* Main content — breathe in on every page load */
    .stMainBlockContainer {
        animation: breatheIn 0.7s cubic-bezier(0.22, 0.61, 0.36, 1) both;
    }
    @keyframes breatheIn {
        0%   { opacity: 0; transform: scale(0.97) translateY(6px); filter: blur(4px); }
        30%  { opacity: 0.5; filter: blur(2px); }
        60%  { opacity: 0.9; transform: scale(1.005) translateY(0); filter: blur(0); }
        100% { opacity: 1; transform: scale(1) translateY(0); filter: blur(0); }
    }

    /* Sidebar — gentle slide-in */
    section[data-testid="stSidebar"] {
        animation: sidebarSlideIn 0.5s ease-out both;
    }
    @keyframes sidebarSlideIn {
        from { opacity: 0; transform: translateX(-12px); }
        to   { opacity: 1; transform: translateX(0); }
    }

    /* Feature cards — staggered shimmer entrance */
    .feature-card {
        animation: cardReveal 0.55s ease-out both, cardBreathe 4s ease-in-out infinite;
    }
    @keyframes cardReveal {
        0%   { opacity: 0; transform: translateY(16px) scale(0.96); }
        60%  { opacity: 0.7; }
        100% { opacity: 1; transform: translateY(0) scale(1); }
    }
    /* Stagger: each card appears slightly later than the previous */
    .feature-card:nth-child(1) { animation-delay: 0.05s, 0.05s; }
    .feature-card:nth-child(2) { animation-delay: 0.12s, 0.12s; }
    .feature-card:nth-child(3) { animation-delay: 0.19s, 0.19s; }
    .feature-card:nth-child(4) { animation-delay: 0.26s, 0.26s; }

    /* QA card — smooth rise */
    .qa-card {
        animation: qaRise 0.5s ease-out both;
    }
    @keyframes qaRise {
        from { opacity: 0; transform: translateY(10px); }
        to   { opacity: 1; transform: translateY(0); }
    }

    /* ── Loading Pulse (plays once on page load, feels intentional) ── */
    .main-title {
        animation: bannerReveal 0.6s ease-out both;
    }
    @keyframes bannerReveal {
        0%   { opacity: 0; transform: scale(0.95); }
        50%  { opacity: 0.8; }
        100% { opacity: 1; transform: scale(1); }
    }

    /* Navigation item — subtle delayed fade */
    .nav-item {
        animation: navItemFade 0.4s ease-out both;
    }
    @keyframes navItemFade {
        from { opacity: 0; transform: translateX(-8px); }
        to   { opacity: 1; transform: translateX(0); }
    }

    /* ── Main Title Banner ── */
    .main-title {
        background: linear-gradient(135deg, #1d4ed8 0%, #2563eb 35%, #3b82f6 100%) !important;
        background-size: 200% 200% !important;
        animation: bannerReveal 0.6s ease-out both, bannerShimmer 6s ease-in-out infinite !important;
        padding: 2rem 2rem !important; border-radius: 20px !important;
        color: #ffffff !important; text-align: center; margin-bottom: 1.2rem;
        border: 1px solid rgba(255,255,255,0.18) !important;
        box-shadow: 0 8px 32px rgba(29,78,216,0.2), inset 0 1px 0 rgba(255,255,255,0.15);
        position: relative; overflow: hidden;
    }
    .main-title::before {
        content: ''; position: absolute; top: 0; left: 0; right: 0; bottom: 0;
        background: linear-gradient(135deg, transparent 30%, rgba(255,255,255,0.08) 50%, transparent 70%);
        background-size: 200% 200%;
        animation: bannerShine 8s ease-in-out infinite;
    }
    @keyframes bannerShimmer {
        0%,100% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
    }
    @keyframes bannerShine {
        0%,100% { background-position: -100% -100%; }
        50% { background-position: 200% 200%; }
    }
    .main-title * { color: #ffffff !important; position: relative; z-index: 1; }
    .main-title h1 { font-size: 1.65rem !important; font-weight: 650 !important; margin: 0 !important; letter-spacing: -0.01em !important; }
    .main-title p { opacity: 0.82; margin-top: 0.35rem; font-size: 0.88rem; font-weight: 380; letter-spacing: 0.01em; }

    /* ── Feature Cards (glass morphism) ── */
    .feature-card {
        background: linear-gradient(135deg, rgba(255,255,255,0.9) 0%, rgba(248,250,252,0.85) 100%);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(226,232,240,0.8);
        border-radius: 20px; padding: 24px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.03), 0 8px 24px rgba(0,0,0,0.04);
        transition: all 0.4s cubic-bezier(0.25, 0.46, 0.45, 0.94);
        cursor: pointer; margin-bottom: 4px; position: relative; overflow: hidden;
        height: 100%; min-height: 190px;
        display: flex; flex-direction: column;
    }
    .feature-card::before {
        content: '';
        position: absolute; top: -50%; right: -50%; width: 100%; height: 100%;
        background: radial-gradient(circle, rgba(99,102,241,0.04) 0%, transparent 70%);
        transition: all 0.5s ease; pointer-events: none;
    }
    .feature-card:hover::before {
        top: -20%; right: -30%; width: 120%; height: 120%;
    }
    .feature-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.06), 0 12px 32px rgba(99,102,241,0.1);
        border-color: rgba(99,102,241,0.25);
    }
    .feature-card .card-icon {
        width: 48px; height: 48px; border-radius: 14px;
        display: inline-flex; align-items: center; justify-content: center;
        margin-bottom: 14px; flex-shrink: 0; position: relative; z-index: 1;
    }
    .feature-card .card-icon svg {
        width: 26px; height: 26px;
    }
    .feature-card .card-title {
        font-size: 1rem; font-weight: 650; color: #1e293b; margin-bottom: 4px;
        position: relative; z-index: 1;
        letter-spacing: -0.01em; line-height: 1.3;
    }
    .feature-card .card-desc {
        font-size: 0.8rem; color: #64748b; line-height: 1.5; margin-bottom: 10px;
        position: relative; z-index: 1; font-weight: 420;
    }
    .feature-card .card-tags {
        display: flex; gap: 5px; flex-wrap: wrap; margin-bottom: 8px;
        position: relative; z-index: 1;
    }
    .feature-card .card-tag {
        font-size: 0.68rem; padding: 3px 9px;
        background: rgba(241,245,249,0.8); color: #6366f1;
        border-radius: 20px; font-weight: 500; letter-spacing: 0.01em;
        backdrop-filter: blur(4px);
    }
    /* ── Card Icon Color Variants ── */
    .icon-math { background: linear-gradient(135deg, #eef2ff 0%, #e0e7ff 100%); color: #4f46e5; }
    .icon-fire { background: linear-gradient(135deg, #fdf2f8 0%, #fce7f3 100%); color: #db2777; }
    .icon-eng  { background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%); color: #059669; }
    .icon-fb   { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); color: #0284c7; }
    .icon-mat  { background: linear-gradient(135deg, #fefce8 0%, #fef9c3 100%); color: #ca8a04; }
    .icon-ck   { background: linear-gradient(135deg, #f0fdf4 0%, #dcfce7 100%); color: #16a34a; }

    /* ── Card Breathing Animation ── */
    @keyframes cardBreathe {
        0%, 100% { box-shadow: 0 1px 3px rgba(0,0,0,0.03), 0 8px 24px rgba(0,0,0,0.04); }
        50% { box-shadow: 0 2px 6px rgba(0,0,0,0.04), 0 10px 28px rgba(99,102,241,0.06); }
    }
    .feature-card:hover { animation: none; }

    /* ── Q&A Card ── */
    .qa-card {
        background: #ffffff; border-radius: 18px;
        padding: 24px; border: 1px solid #e2e8f0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.04), 0 4px 12px rgba(0,0,0,0.05);
        margin-bottom: 16px;
    }

    /* ── Memory / Knowledge Cards ── */
    .memory-card {
        padding: 12px 14px; margin: 6px 0;
        background: #ffffff; border: 1px solid #e2e8f0;
        border-radius: 12px; cursor: pointer;
        transition: all 0.25s ease;
    }
    .memory-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 2px 6px rgba(0,0,0,0.06), 0 8px 24px rgba(79,70,229,0.08);
        border-color: rgba(79,70,229,0.25);
    }
    .learning-card {
        padding: 8px 10px; margin: 3px 0;
        background: #fff7ed; border-radius: 6px;
        border-left: 3px solid #f97316;
        font-size: 12px; overflow: hidden; text-overflow: ellipsis;
    }
    .mastered-card {
        padding: 8px 10px; margin: 3px 0;
        background: #f0fdf4; border-radius: 6px;
        border-left: 3px solid #22c55e;
        font-size: 12px; overflow: hidden; text-overflow: ellipsis;
    }

    /* ── Reference Tags ── */
    .ref-tag {
        display: inline-block;
        background: #eef2ff; color: #4f46e5;
        padding: 3px 10px; border-radius: 20px;
        margin: 2px 4px; font-size: 12px;
        border: 1px solid rgba(79,70,229,0.15);
    }

    /* ── Calendar Grid ── */
    .cal-grid { display: grid; grid-template-columns: repeat(10, 1fr); gap: 2px; text-align: center; }
    .cal-grid .cal-cell { padding: 4px 0; }
    .cal-grid .cal-cell small { font-size: 11px; }

    /* ── Quiz Area ── */
    .quiz-area { max-height: 600px; overflow-y: auto; padding-right: 4px; }
    /* 信息卡片分类标签 */
    .info-badge {
        display: inline-block; background: linear-gradient(135deg, #4f46e5, #6366f1);
        color: #fff; padding: 2px 12px; border-radius: 20px;
        font-size: 0.75rem; font-weight: 600; margin-right: 8px;
        letter-spacing: 0.02em; vertical-align: middle;
    }

    /* ═══════════════════════════════════
       STREAMLIT COMPONENT OVERRIDES
       ═══════════════════════════════════ */

    /* Knowledge base buttons */
    .st-key-view_ button, .st-key-quiz_ button, .st-key-concept_ button {
        padding: 0.6rem 1rem !important;
        font-size: 0.88rem !important;
        font-weight: 600 !important;
        border-radius: 10px !important;
    }

    /* ── Buttons ── */
    button[kind="primary"], .st-key-hub_qa button, .st-key-hub_pop button,
    .st-key-hub_english button, .st-key-hub_suggest button,
    .st-key-hub_material button, .st-key-hub_checkin button,
    button[data-testid="baseButton-primary"],
    div[data-testid="stButton"] button[kind="primary"] {
        background: #2563eb !important; color: #ffffff !important;
        border: none !important; border-radius: 24px !important;
        font-weight: 700 !important; padding: 0.6rem 1.5rem !important;
        transition: all 0.2s !important;
        text-shadow: 0 1px 2px rgba(0,0,0,0.1) !important;
        letter-spacing: 0.02em !important;
    }
    button[kind="primary"]:hover, button[data-testid="baseButton-primary"]:hover {
        background: #4338ca !important; transform: scale(1.02);
    }
    button[kind="secondary"], button[kind="secondary"]:hover {
        background: #ffffff !important; color: #64748b !important;
        border: 1px solid #e2e8f0 !important; border-radius: 24px !important;
        font-weight: 500 !important;
    }

    /* ── Container cards (border=True) ── */
    div[data-testid="stVerticalBlockBorderWrapper"] {
        background: #ffffff !important;
        border: 1px solid #e2e8f0 !important;
        border-radius: 18px !important;
        box-shadow: 0 1px 3px rgba(0,0,0,0.04), 0 4px 12px rgba(0,0,0,0.05) !important;
        padding: 8px !important;
    }

    /* ── Tabs ── */
    div[data-testid="stTabs"] button {
        font-weight: 500 !important; color: #64748b !important;
        border-radius: 8px !important; padding: 0.45rem 1rem !important;
    }
    div[data-testid="stTabs"] button[aria-selected="true"] {
        color: #4f46e5 !important; font-weight: 600 !important;
        background: #eef2ff !important;
    }

    /* ── Expanders ── */
    div[data-testid="stExpander"] {
        border: 1px solid #e2e8f0 !important;
        border-radius: 12px !important;
        background: #ffffff !important;
    }
    div[data-testid="stExpander"] summary {
        font-weight: 600 !important; color: #64748b !important;
        padding: 10px 14px !important;
    }
    /* Expander: force-hide icon text overflow */
    div[data-testid="stExpander"] summary {
        overflow: hidden !important;
    }
    div[data-testid="stExpander"] summary p {
        display: inline !important;
    }
    div[data-testid="stExpander"] details summary > span:first-child,
    div[data-testid="stExpander"] summary svg,
    div[data-testid="stExpander"] summary [data-testid="stIconMaterial"],
    div[data-testid="stExpander"] summary [class*="material"] {
        font-size: 0 !important; width: 0 !important; height: 0 !important;
        overflow: hidden !important; display: inline-block !important;
        color: transparent !important; line-height: 0 !important;
    }

    /* File uploader: hide icon text */
    div[data-testid="stFileUploader"] [data-testid="stIconMaterial"],
    div[data-testid="stFileUploader"] [class*="material"],
    div[data-testid="stFileUploadDropzone"] span[class*="icon"],
    section[data-testid="stFileUploadDropzone"] [aria-hidden="true"] {
        font-size: 0 !important; width: 0 !important; overflow: hidden !important;
        color: transparent !important; display: inline-block !important;
    }

    /* ── Metrics ── */
    div[data-testid="stMetricValue"] { color: #4f46e5 !important; font-weight: 800 !important; }

    /* ── Select boxes / Inputs ── */
    div[data-testid="stSelectbox"], input, textarea, .stTextInput input {
        border-radius: 8px !important;
    }
    input:focus, textarea:focus, div[data-testid="stSelectbox"]:focus-within {
        border-color: #4f46e5 !important;
        box-shadow: 0 0 0 3px rgba(79,70,229,0.10) !important;
    }

    /* ── Radio buttons ── */
    div[data-testid="stRadio"] label {
        border: 1px solid #e2e8f0 !important; border-radius: 20px !important;
        padding: 5px 14px !important; margin-right: 4px !important;
    }

    /* ── Form submit ── */
    div[data-testid="stFormSubmitButton"] button {
        background: #4f46e5 !important; color: white !important;
        border: none !important; border-radius: 24px !important;
        font-weight: 600 !important; width: 100% !important;
    }
    div[data-testid="stFormSubmitButton"] button:hover { background: #4338ca !important; }

    /* ── Progress bars ── */
    div[data-testid="stProgress"] > div > div { background: #4f46e5 !important; }

    /* ── Success / Info / Warning / Error messages ── */
    div[data-testid="stSuccessMessage"], div[data-testid="stInfoMessage"],
    div[data-testid="stWarningMessage"], div[data-testid="stErrorMessage"] {
        border-radius: 10px !important; border: none !important;
    }

    /* ── Chat input ── */
    div[data-testid="stChatInput"] textarea {
        border-radius: 24px !important; border: 1px solid #e2e8f0 !important;
    }

    /* ═══════════════════════════════════
       RESPONSIVE
       ═══════════════════════════════════ */
    @media (max-width: 1024px) {
        .main-title h1 { font-size: 1.5rem !important; }
        .qa-card { padding: 18px !important; }
    }
    @media (max-width: 768px) {
        .main-title { padding: 1.2rem !important; }
        .main-title h1 { font-size: 1.25rem !important; }
        .main-title p { font-size: 0.8rem !important; }
        .qa-card { padding: 14px !important; font-size: 14px !important; }
        .learning-card, .mastered-card { white-space: normal !important; font-size: 11px !important; }
        .memory-card { padding: 8px !important; font-size: 13px !important; }
        .ref-tag { font-size: 11px !important; padding: 2px 6px !important; }
        .cal-grid { grid-template-columns: repeat(6, 1fr) !important; }
        div[data-testid="stMetricValue"] { font-size: 1.1rem !important; }
        div[data-testid="stMetricLabel"] { font-size: 0.75rem !important; }
        div[data-testid="stHorizontalBlock"] { gap: 0.5rem !important; }
    }
    @media (max-width: 480px) {
        .main-title { padding: 0.8rem !important; }
        .main-title h1 { font-size: 1rem !important; }
        .qa-card { padding: 10px !important; }
        .cal-grid { grid-template-columns: repeat(5, 1fr) !important; }
        div[data-testid="stMetricValue"] { font-size: 0.9rem !important; }
        div[data-testid="stMetricLabel"] { font-size: 0.7rem !important; }
    }
</style>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
<link href="https://fonts.googleapis.com/icon?family=Material+Icons" rel="stylesheet">
<link href="https://cdn.bootcdn.net/ajax/libs/material-design-icons/3.0.1/iconfont/material-icons.min.css" rel="stylesheet">
""", unsafe_allow_html=True)

st.html("""
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"></script>
<script>
(function(){
  function cleanIcons(){
    var els=document.querySelectorAll('[class*="material-icons"],[class*="material-symbols"],[class*="keyboard_"]');
    for(var i=0;i<els.length;i++){
      var el=els[i];
      el.innerHTML=''; el.textContent='';
      el.style.fontSize='0'; el.style.width='20px'; el.style.height='20px';
      el.style.display='inline-block'; el.style.overflow='hidden';
      el.style.color='transparent'; el.style.userSelect='none';
    }
    var walker=document.createTreeWalker(document.body,4,null,false), node;
    while(node=walker.nextNode()){
      var v=node.nodeValue;
      if(!v) continue;
      var t=v.trim();
      if(/^(expand_more|expand_less|chevron_right|chevron_left|keyboard_arrow|arrow_|keyboard_double|upload|file_upload|cloud_upload|attach_file|close|check|search|menu|more_vert|more_horiz|delete|edit|add|remove|visibility|visibility_off)/.test(t)){
        node.nodeValue='';
      } else if(/_arrow|upload/i.test(v)){
        node.nodeValue=v.replace(/_arrow[-_]?/g,'').replace(/\bupload\b/gi,'');
      }
    }
  }
  cleanIcons();
  new MutationObserver(cleanIcons).observe(document.body,{childList:true,subtree:true});
})();
</script>
""")

# ==================== 持久化登录（CookieManager 方案） ====================

@st.cache_resource
def get_cookie_manager():
    return stx.CookieManager()

cookie_manager = get_cookie_manager()

def generate_login_token():
    """生成 64 字符随机 token"""
    return secrets.token_hex(32)

def save_login_token(user_id, token):
    """将 token 存入数据库"""
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("UPDATE users SET login_token=? WHERE id=?", (token, user_id))
    conn.commit()
    conn.close()

def verify_login_token(token):
    """验证 token，返回 user_id 或 None"""
    if not token:
        return None
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT id, username FROM users WHERE login_token=?", (token,))
    row = c.fetchone()
    conn.close()
    if row:
        return {"user_id": row[0], "username": row[1]}
    return None

def clear_login_token(user_id):
    """清除数据库中的 token"""
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute("UPDATE users SET login_token=NULL WHERE id=?", (user_id,))
    conn.commit()
    conn.close()

# ==================== 核心功能 ====================

def _extract_content(msg):
    """从 API 响应中提取内容（兼容 MiMo 思维链模型）
    优先返回 content（最终回答），content 为空时才用 reasoning_content（思考过程）"""
    c = msg.get("content")
    if c is not None and c != "":
        return c
    return msg.get("reasoning_content") or ""

def _typing_display(placeholder, text, delay=0.02):
    """打字效果显示文本，LaTeX 公式整体插入不拆散"""
    import time as _time
    import re as _re

    # 将文本分割为 LaTeX 块和普通文本块
    # 匹配 $$...$$ 或 $...$（非贪婪）
    parts = _re.split(r'(\$\$[\s\S]*?\$\$|\$[^\$]+?\$)', text)

    displayed = ""
    for part in parts:
        if not part:
            continue
        if part.startswith("$"):
            # LaTeX 块：整体插入，不逐字
            displayed += part
            placeholder.markdown(displayed)
            _time.sleep(0.1)
        else:
            # 普通文本：逐字显示
            for char in part:
                displayed += char
                placeholder.markdown(displayed)
                _time.sleep(delay)

def read_file(p):
    try:
        return p.read_text(encoding="utf-8", errors="ignore")
    except:
        try:
            return p.read_text(encoding="gbk", errors="ignore")
        except:
            return ""

@st.cache_data
def load_corpus():
    docs = []
    if DATA_DIR.exists():
        for f in sorted(DATA_DIR.iterdir()):
            if f.is_file() and f.suffix.lower() in [".txt", ".md"]:
                t = read_file(f)
                if t and len(t) > 50:
                    docs.append({"id": f.name, "text": t})
    return docs

def _filter_corpus(corpus, math_type):
    """全部 / 数学一专属 / 数学三专属 过滤"""
    def _num(doc):
        return doc["id"].split("-", 1)[0]
    if math_type == "数学一专属":
        return [d for d in corpus if _num(d) in MATH1_ONLY]
    elif math_type == "数学三专属":
        return [d for d in corpus if _num(d) in MATH3_ONLY]
    return corpus

@st.cache_data
def load_demo_corpus():
    docs = []
    if DEMO_DATA_DIR.exists():
        for f in sorted(DEMO_DATA_DIR.iterdir()):
            if f.is_file() and f.suffix.lower() in [".txt", ".md"]:
                t = read_file(f)
                if t and len(t) > 50:
                    docs.append({"id": f.name, "text": t})
    return docs

def save_document(filename, content):
    file_path = DATA_DIR / filename
    try:
        file_path.write_text(content, encoding="utf-8")
        return True
    except:
        return False

def search_corpus(query, corpus, top_k=3):
    if not corpus or not query:
        return []
    query_lower = query.lower()
    results = []
    for doc in corpus:
        text = doc["text"].lower()
        score = sum(text.count(w) for w in query_lower.split() if w)
        if score > 0:
            results.append({"id": doc["id"], "score": score, "text": doc["text"][:500]})
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:top_k]

def get_knowledge_text(kid, corpus):
    for doc in corpus:
        if kid in doc["id"]:
            return doc["text"]
    return ""

import hashlib

# ==================== 用户管理 ====================

def hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def register_user(username, password):
    """注册新用户，返回 user_id 或 None"""
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT id FROM users WHERE username=?", (username,))
    if c.fetchone():
        conn.close()
        return None  # 用户名已存在
    pw_hash = hash_password(password)
    c.execute("INSERT INTO users (username, password_hash, display_name) VALUES (?, ?, ?)",
              (username, pw_hash, username))
    conn.commit()
    user_id = c.lastrowid
    conn.close()
    return user_id

def login_user(username, password):
    """登录，返回 user_id 或 None"""
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    pw_hash = hash_password(password)
    c.execute("SELECT id FROM users WHERE username=? AND password_hash=?", (username, pw_hash))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None

def get_experience_file():
    uid = st.session_state.get("user_id", 1)
    return Path(f"agent_experience_{uid}.md")

def load_agent_experience():
    exp_file = get_experience_file()
    if exp_file.exists():
        try:
            return exp_file.read_text(encoding="utf-8").strip()
        except:
            return ""
    return ""

def save_agent_experience(text):
    exp_file = get_experience_file()
    try:
        exp_file.write_text(text, encoding="utf-8")
        return True
    except:
        return False

def get_recent_experiences(count=5):
    exp = load_agent_experience()
    if not exp:
        return []
    parts = exp.split("---")
    return parts[-count:] if len(parts) >= count else parts

# ==================== 记忆系统 ====================

def init_memory_db():
    Path(MEMORY_DB).parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()

    # 确保 users 表
    c.execute("""CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY, username TEXT UNIQUE, display_name TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")
    try: c.execute("SELECT password_hash FROM users LIMIT 1")
    except: c.execute("ALTER TABLE users ADD COLUMN password_hash TEXT")
    try: c.execute("SELECT display_name FROM users LIMIT 1")
    except: c.execute("ALTER TABLE users ADD COLUMN display_name TEXT")
    try: c.execute("SELECT login_token FROM users LIMIT 1")
    except: c.execute("ALTER TABLE users ADD COLUMN login_token TEXT")

    c.execute("""CREATE TABLE IF NOT EXISTS knowledge_mastery (
        id INTEGER PRIMARY KEY, knowledge_id TEXT, user_id INTEGER DEFAULT 1,
        mastery_level REAL DEFAULT 0, status TEXT DEFAULT '陌生',
        times_correct INTEGER DEFAULT 0, times_wrong INTEGER DEFAULT 0,
        stability REAL DEFAULT 1.0, last_review TIMESTAMP,
        error_type TEXT DEFAULT '', wrong_reason TEXT DEFAULT '')""")

    try:
        c.execute("SELECT error_type FROM knowledge_mastery LIMIT 1")
    except:
        c.execute("ALTER TABLE knowledge_mastery ADD COLUMN error_type TEXT DEFAULT ''")

    try:
        c.execute("SELECT wrong_reason FROM knowledge_mastery LIMIT 1")
    except:
        c.execute("ALTER TABLE knowledge_mastery ADD COLUMN wrong_reason TEXT DEFAULT ''")

    try:
        c.execute("SELECT stability FROM knowledge_mastery LIMIT 1")
    except:
        c.execute("ALTER TABLE knowledge_mastery ADD COLUMN stability REAL DEFAULT 1.0")

    c.execute("""CREATE TABLE IF NOT EXISTS user_performance (
        id INTEGER PRIMARY KEY, user_id INTEGER DEFAULT 1,
        knowledge_id TEXT, is_correct INTEGER, error_type TEXT,
        mastery_score REAL, created_at TIMESTAMP)""")

    c.execute("""CREATE TABLE IF NOT EXISTS review_challenges (
        id INTEGER PRIMARY KEY, knowledge_id TEXT, user_id INTEGER DEFAULT 1,
        challenge_type TEXT, completed INTEGER DEFAULT 0,
        created_at TIMESTAMP)""")

    c.execute("""CREATE TABLE IF NOT EXISTS visit_log (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        username TEXT, action TEXT, detail TEXT)""")

    c.execute("""CREATE TABLE IF NOT EXISTS suggestions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT, content TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)""")

    # 打卡督学模块 - 每日打卡记录
    c.execute("""CREATE TABLE IF NOT EXISTS checkin_daily (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        checkin_date TEXT NOT NULL,
        subject TEXT,
        duration_minutes INTEGER,
        completion_rate REAL,
        mood TEXT,
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, checkin_date)
    )""")

    # 打卡督学模块 - 晚间复盘
    c.execute("""CREATE TABLE IF NOT EXISTS checkin_review (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        review_date TEXT NOT NULL,
        what_learned TEXT,
        what_difficult TEXT,
        what_improve TEXT,
        overall_rating INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, review_date)
    )""")

    # 添加日记内容列（如果不存在）
    try:
        c.execute("SELECT diary_content FROM checkin_review LIMIT 1")
    except:
        c.execute("ALTER TABLE checkin_review ADD COLUMN diary_content TEXT")

    # 打卡督学模块 - 学习计划
    c.execute("""CREATE TABLE IF NOT EXISTS checkin_plans (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        plan_name TEXT,
        target_date TEXT,
        tasks TEXT,
        progress REAL DEFAULT 0,
        status TEXT DEFAULT 'active',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # 打卡督学模块 - 番茄钟记录
    c.execute("""CREATE TABLE IF NOT EXISTS checkin_pomodoro (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        subject TEXT,
        duration_minutes INTEGER,
        actual_minutes INTEGER,
        completed INTEGER DEFAULT 0,
        started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # 费曼学习法 - 答题记录
    c.execute("""CREATE TABLE IF NOT EXISTS feynman_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        mode TEXT,
        question_text TEXT,
        user_answer TEXT,
        ai_evaluation TEXT,
        score_correct INTEGER,
        score_expression INTEGER,
        score_authentic INTEGER,
        total_score INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # ==================== 用户画像模块 - 3 张新表 ====================

    # 用户画像档案
    c.execute("""CREATE TABLE IF NOT EXISTS user_profiles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL UNIQUE,
        grade TEXT,
        major TEXT,
        undergraduate_school TEXT,
        cet4_score INTEGER,
        cet6_score INTEGER,
        target_year TEXT,
        target_major TEXT,
        target_region TEXT,
        target_schools TEXT,
        risk_preference TEXT,
        daily_hours REAL,
        schedule_preference TEXT,
        strong_subjects TEXT,
        weak_subjects TEXT,
        material_preference TEXT,
        math_exam_type TEXT,
        current_phase TEXT,
        completed_progress TEXT,
        mock_scores TEXT,
        anxiety_level INTEGER,
        procrastination_type TEXT,
        motivation_preference TEXT,
        common_errors TEXT,
        strong_areas TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # 补字段（服务器可能已有旧表）
    for col, typ in [
        ("grade", "TEXT"), ("major", "TEXT"), ("undergraduate_school", "TEXT"),
        ("undergraduate_major", "TEXT"), ("undergraduate_level", "TEXT"), ("is_cross_major", "TEXT"),
        ("cet4_score", "INTEGER"), ("cet6_score", "INTEGER"), ("target_year", "TEXT"),
        ("target_major", "TEXT"), ("target_region", "TEXT"), ("target_schools", "TEXT"),
        ("risk_preference", "TEXT"), ("daily_hours", "REAL"), ("schedule_preference", "TEXT"),
        ("strong_subjects", "TEXT"), ("weak_subjects", "TEXT"), ("material_preference", "TEXT"),
        ("math_exam_type", "TEXT"), ("current_phase", "TEXT"), ("completed_progress", "TEXT"),
        ("mock_scores", "TEXT"), ("anxiety_level", "INTEGER"), ("procrastination_type", "TEXT"),
        ("motivation_preference", "TEXT"), ("common_errors", "TEXT"), ("strong_areas", "TEXT"),
        ("updated_at", "TIMESTAMP"),
    ]:
        try:
            c.execute(f"SELECT {col} FROM user_profiles LIMIT 1")
        except:
            c.execute(f"ALTER TABLE user_profiles ADD COLUMN {col} {typ}")

    # 情节记忆
    c.execute("""CREATE TABLE IF NOT EXISTS episodic_memories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        memory_type TEXT,
        content TEXT,
        context TEXT,
        importance REAL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # 会话记忆
    c.execute("""CREATE TABLE IF NOT EXISTS session_memories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        session_id TEXT,
        role TEXT,
        content TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # ==================== 学习规划模块 - 3 张新表 ====================

    # 学习计划
    c.execute("""CREATE TABLE IF NOT EXISTS study_plans (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        plan_name TEXT,
        target_exam_date TEXT,
        math_type TEXT,
        daily_hours REAL,
        subjects_config TEXT,
        current_phase TEXT,
        status TEXT DEFAULT 'active',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # 计划任务
    c.execute("""CREATE TABLE IF NOT EXISTS plan_tasks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        plan_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        task_type TEXT,
        subject TEXT,
        task_name TEXT,
        description TEXT,
        target_date TEXT,
        estimated_hours REAL,
        actual_hours REAL DEFAULT 0,
        status TEXT DEFAULT 'pending',
        priority INTEGER DEFAULT 3,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        completed_at TIMESTAMP
    )""")

    # 进度记录
    c.execute("""CREATE TABLE IF NOT EXISTS plan_progress (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        plan_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        record_date TEXT,
        subject TEXT,
        planned_hours REAL,
        actual_hours REAL,
        tasks_completed INTEGER,
        tasks_total INTEGER,
        deviation_percent REAL,
        notes TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    # ==================== 专业知识库模块 ====================

    c.execute("""CREATE TABLE IF NOT EXISTS user_materials (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        subject TEXT,
        filename TEXT,
        chapter_name TEXT,
        file_path TEXT,
        file_type TEXT,
        processing_status TEXT DEFAULT 'pending',
        knowledge_count INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_knowledge (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        material_id INTEGER,
        subject TEXT,
        chapter_name TEXT,
        knowledge_name TEXT,
        content TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_wrong_questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        knowledge_id INTEGER,
        subject TEXT,
        chapter_name TEXT,
        question TEXT,
        user_answer TEXT,
        correct_answer TEXT,
        explanation TEXT,
        error_count INTEGER DEFAULT 1,
        status TEXT DEFAULT 'active',
        last_reviewed TIMESTAMP,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    c.execute("""CREATE TABLE IF NOT EXISTS user_review_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        knowledge_id INTEGER,
        review_date TEXT,
        mastered INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")

    conn.commit()
    conn.close()

import knowledge_base as kb

def log_visit(action, detail=""):
    try:
        conn = sqlite3.connect(MEMORY_DB)
        c = conn.cursor()
        username = st.session_state.get("username", "anon")
        c.execute("INSERT INTO visit_log (username, action, detail) VALUES (?, ?, ?)",
                  (username, action, detail[:200]))
        conn.commit()
        conn.close()
    except:
        pass

# ==================== 打卡督学模块 ====================

def checkin_today_str():
    return date.today().strftime("%Y-%m-%d")

def checkin_fetch_one(query, params=()):
    conn = sqlite3.connect(MEMORY_DB)
    conn.row_factory = sqlite3.Row
    row = conn.execute(query, params).fetchone()
    conn.close()
    return row

def checkin_fetch_all(query, params=()):
    conn = sqlite3.connect(MEMORY_DB)
    conn.row_factory = sqlite3.Row
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return rows

def save_checkin(user_id, checkin_date, subject, duration, completion, mood, notes):
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        """INSERT INTO checkin_daily
           (user_id, checkin_date, subject, duration_minutes, completion_rate, mood, notes)
           VALUES (?, ?, ?, ?, ?, ?, ?)
           ON CONFLICT(user_id, checkin_date) DO UPDATE SET
               subject=excluded.subject,
               duration_minutes=excluded.duration_minutes,
               completion_rate=excluded.completion_rate,
               mood=excluded.mood,
               notes=excluded.notes""",
        (user_id, checkin_date, subject, duration, completion, mood, notes))
    conn.commit()
    conn.close()

def get_today_checkin(user_id):
    return checkin_fetch_one(
        "SELECT * FROM checkin_daily WHERE user_id=? AND checkin_date=?",
        (user_id, checkin_today_str()))

def get_recent_checkins(user_id, days=30):
    start = (date.today() - timedelta(days=days - 1)).strftime("%Y-%m-%d")
    return checkin_fetch_all(
        """SELECT * FROM checkin_daily WHERE user_id=? AND checkin_date>=?
           ORDER BY checkin_date DESC""", (user_id, start))

def get_last_checkin_date(user_id):
    row = checkin_fetch_one(
        "SELECT checkin_date FROM checkin_daily WHERE user_id=? ORDER BY checkin_date DESC LIMIT 1",
        (user_id,))
    return row["checkin_date"] if row else None

def get_consecutive_days(user_id):
    rows = checkin_fetch_all(
        "SELECT checkin_date FROM checkin_daily WHERE user_id=? ORDER BY checkin_date DESC",
        (user_id,))
    checkin_dates = {row["checkin_date"] for row in rows}
    streak = 0
    cursor = date.today()
    while cursor.strftime("%Y-%m-%d") in checkin_dates:
        streak += 1
        cursor -= timedelta(days=1)
    return streak

def get_today_duration(user_id):
    row = checkin_fetch_one(
        "SELECT duration_minutes FROM checkin_daily WHERE user_id=? AND checkin_date=?",
        (user_id, checkin_today_str()))
    return int(row["duration_minutes"]) if row and row["duration_minutes"] else 0

def get_today_mood(user_id):
    row = checkin_fetch_one(
        "SELECT mood FROM checkin_daily WHERE user_id=? AND checkin_date=?",
        (user_id, checkin_today_str()))
    return row["mood"] if row else "未打卡"

def save_review(user_id, review_date, diary_content):
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        """INSERT INTO checkin_review
           (user_id, review_date, diary_content)
           VALUES (?, ?, ?)
           ON CONFLICT(user_id, review_date) DO UPDATE SET
               diary_content=excluded.diary_content""",
        (user_id, review_date, diary_content))
    conn.commit()
    conn.close()

def get_today_review(user_id):
    return checkin_fetch_one(
        "SELECT * FROM checkin_review WHERE user_id=? AND review_date=?",
        (user_id, checkin_today_str()))

def get_recent_reviews(user_id, limit=10):
    """获取最近的日记记录"""
    return checkin_fetch_all(
        """SELECT * FROM checkin_review WHERE user_id=? AND diary_content IS NOT NULL AND diary_content != ''
           ORDER BY review_date DESC LIMIT ?""",
        (user_id, limit))

def get_timeline(user_id, days=14):
    """获取日记+计划的穿插时间线"""
    start = (date.today() - timedelta(days=days - 1)).strftime("%Y-%m-%d")
    
    # 获取日记记录
    reviews = checkin_fetch_all(
        """SELECT review_date as date, 'diary' as type, diary_content as content, created_at
           FROM checkin_review 
           WHERE user_id=? AND review_date>=? AND diary_content IS NOT NULL AND diary_content != ''""",
        (user_id, start))
    
    # 获取计划活动（创建和完成）
    plan_activities = checkin_fetch_all(
        """SELECT DATE(created_at) as date, 'plan_create' as type, 
                  plan_name || '：创建了新计划' as content, created_at
           FROM checkin_plans 
           WHERE user_id=? AND DATE(created_at)>=?""",
        (user_id, start))
    
    # 合并并排序
    all_items = list(reviews) + list(plan_activities)
    all_items.sort(key=lambda x: str(x['created_at'] or ''), reverse=True)
    
    return all_items

def save_checkin_plan(user_id, plan_name, target_date, tasks):
    if isinstance(tasks, list):
        progress = calc_tasks_progress(tasks)
        tasks_json = json.dumps(tasks, ensure_ascii=False)
    else:
        progress = 0
        tasks_json = str(tasks)
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        """INSERT INTO checkin_plans (user_id, plan_name, target_date, tasks, progress, status)
           VALUES (?, ?, ?, ?, ?, 'active')""",
        (user_id, plan_name, target_date, tasks_json, progress))
    conn.commit()
    conn.close()

def _extract_text_from_pdf(file_path):
    """用 PyMuPDF 提取 PDF 文本（纯文本模式）"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text[:5000]
    except ImportError:
        return ""
    except Exception:
        return ""

def _check_umiocr_available():
    """检查 umi-ocr API 是否可用"""
    try:
        import requests
        resp = requests.get(f"{UMI_OCR_URL}/api/status", timeout=5)
        return resp.status_code == 200
    except:
        return False

def _extract_text_from_pdf_umiocr(file_path):
    """用 umi-ocr API 逐页识别 PDF（中文 OCR）"""
    import fitz
    doc = fitz.open(str(file_path))
    all_text = []
    total_pages = min(len(doc), 20)  # 最多处理 20 页
    
    for page_num in range(total_pages):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        img_b64 = base64.b64encode(img_bytes).decode()
        
        try:
            import requests
            resp = requests.post(
                f"{UMI_OCR_URL}/api/ocr",
                json={"base64": img_b64},
                timeout=30
            )
            result = resp.json()
            if result.get("text"):
                all_text.append(f"=== 第{page_num+1}页 ===\n{result['text']}")
        except Exception as e:
            st.warning(f"第{page_num+1}页 OCR 失败: {e}")
    
    doc.close()
    return "\n\n".join(all_text)[:8000]

def _extract_knowledge_from_pdf_images(file_path, subject, chapter_name):
    """将 PDF 每页转为图片，用多模态 AI 直接提取知识点"""
    import fitz  # PyMuPDF
    doc = fitz.open(str(file_path))
    all_knowledge = []
    
    for page_num in range(min(len(doc), 20)):  # 最多处理 20 页
        page = doc[page_num]
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("jpeg")
        img_b64 = base64.b64encode(img_bytes).decode()
        
        prompt = f"""请从这张图片中提取所有知识点。

学科：{subject}
章节：{chapter_name}
这是 PDF 第 {page_num+1} 页。

输出格式（严格遵守）：
知识点1: [知识点名称] - [1-2句话简要说明核心概念]
知识点2: [知识点名称] - [1-2句话简要说明核心概念]
...

要求：
- 提取所有可见的知识点
- 如果是公式或定理，写出名称和简要含义
- 如果没有知识点，输出「无」"""

        data = {
            "model": "mimo-v2.5",
            "messages": [{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"}}
            ]}],
            "max_tokens": 1500,
            "temperature": 0
        }
        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(data).encode("utf-8"),
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
            method="POST"
        )
        try:
            with urllib.request.urlopen(req, timeout=90) as resp:
                result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
            if "无" not in result[:10]:
                all_knowledge.append(result)
        except Exception:
            pass
    
    doc.close()
    return "\n".join(all_knowledge)

def _extract_text_from_image(file_bytes):
    """用 glm-4v-flash OCR 识别图片中的文字"""
    img_b64 = base64.b64encode(file_bytes).decode()
    data = {
        "model": "mimo-v2.5",
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": "请识别这张图片中的所有文字内容，只输出文字，不要添加任何说明。"},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
        ]}],
        "max_tokens": 2000,
        "temperature": 0
    }
    req = urllib.request.Request(
        API_BASE + "/chat/completions",
        data=json.dumps(data).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        return _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])

def _extract_knowledge_from_image(file_bytes, subject, chapter_name):
    """用多模态 AI 直接从图片提取知识点"""
    img_b64 = base64.b64encode(file_bytes).decode()
    prompt = f"""请仔细观察这张图片，从中提取所有知识点。

学科：{subject}
章节：{chapter_name}

输出格式（严格遵守）：
知识点1: [知识点名称] - [1-2句话简要说明核心概念]
知识点2: [知识点名称] - [1-2句话简要说明核心概念]
...

要求：
- 提取所有可见的知识点
- 知识点名称用中文
- 简要说明要准确、简洁
- 如果是公式或定理，写出名称和简要含义"""

    data = {
        "model": "mimo-v2.5",
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
        ]}],
        "max_tokens": 2000,
        "temperature": 0
    }
    req = urllib.request.Request(
        API_BASE + "/chat/completions",
        data=json.dumps(data).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=90) as resp:
        return _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])

def _extract_knowledge_from_text(content, subject, chapter_name):
    """用 LLM 从文本中提取知识点"""
    extract_prompt = f"""请从以下内容中提取知识点，输出格式为：
知识点1: [知识点名称]
知识点2: [知识点名称]
...
每个知识点简要说明其核心概念（1-2句话）。

学科：{subject}
章节：{chapter_name}

内容：
{content[:3000]}"""
    return call_llm_api(extract_prompt, model="mimo-v2.5", max_tokens=1500)

def _save_knowledge_points(user_id, material_id, subject, chapter_name, llm_result):
    """保存 LLM 提取的知识点到数据库"""
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    lines_kb = [l.strip() for l in llm_result.split("\n") if l.strip().startswith("知识点")]
    count = 0
    for line_kb in lines_kb:
        name_kb = line_kb.split(":", 1)[-1].strip() if ":" in line_kb else line_kb.strip()
        c.execute("""INSERT INTO user_knowledge
            (user_id, material_id, subject, chapter_name, knowledge_name, content)
            VALUES (?, ?, ?, ?, ?, ?)""",
            (user_id, material_id, subject, chapter_name, name_kb, llm_result))
        count += 1
    c.execute("UPDATE user_materials SET processing_status='done', knowledge_count=? WHERE id=?",
             (count, material_id))
    conn.commit()
    conn.close()
    return count

def get_checkin_plans(user_id):
    return checkin_fetch_all(
        "SELECT * FROM checkin_plans WHERE user_id=? AND status='active' ORDER BY target_date ASC, id DESC",
        (user_id,))

def calc_tasks_progress(tasks):
    if not tasks:
        return 0
    done_count = sum(1 for task in tasks if task.get("done"))
    return round(done_count / len(tasks) * 100, 1)

def update_plan_tasks(user_id, plan_id, tasks):
    progress = calc_tasks_progress(tasks)
    status = "completed" if tasks and progress >= 100 else "active"
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        "UPDATE checkin_plans SET tasks=?, progress=?, status=? WHERE id=? AND user_id=?",
        (json.dumps(tasks, ensure_ascii=False), progress, status, plan_id, user_id))
    conn.commit()
    conn.close()

def delete_plan(user_id, plan_id):
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute("UPDATE checkin_plans SET status='abandoned' WHERE id=? AND user_id=?", (plan_id, user_id))
    conn.commit()
    conn.close()

def get_checkin_plan_progress(user_id):
    rows = checkin_fetch_all(
        "SELECT progress FROM checkin_plans WHERE user_id=? AND status='active'",
        (user_id,))
    if not rows:
        return 0
    return round(sum(float(row["progress"] or 0) for row in rows) / len(rows))

def save_pomodoro(user_id, subject, duration, actual_minutes, completed):
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        """INSERT INTO checkin_pomodoro (user_id, subject, duration_minutes, actual_minutes, completed)
           VALUES (?, ?, ?, ?, ?)""",
        (user_id, subject, duration, actual_minutes, int(completed)))
    conn.commit()
    conn.close()

def get_today_pomodoros(user_id):
    today_prefix = checkin_today_str()
    row = checkin_fetch_one(
        """SELECT COUNT(*) AS count, COALESCE(SUM(actual_minutes), 0) AS minutes
           FROM checkin_pomodoro WHERE user_id=? AND DATE(started_at)=? AND completed=1""",
        (user_id, today_prefix))
    return int(row["count"] or 0), int(row["minutes"] or 0)

def check_checkin_reminders(user_id):
    reminders = []
    recent_3 = checkin_fetch_all(
        """SELECT checkin_date, completion_rate FROM checkin_daily
           WHERE user_id=? ORDER BY checkin_date DESC LIMIT 3""",
        (user_id,))
    expected_dates = [(date.today() - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(3)]
    recent_dates = [row["checkin_date"] for row in recent_3]
    if recent_dates == expected_dates and all(float(row["completion_rate"] or 0) < 60 for row in recent_3):
        reminders.append(("warning", "连续 3 天完成率低于 60%，建议降低任务颗粒度。"))
    last_date = get_last_checkin_date(user_id)
    if last_date:
        gap = (date.today() - datetime.strptime(last_date, "%Y-%m-%d").date()).days
        if gap >= 7:
            reminders.append(("error", f"已经 {gap} 天没有打卡了，建议今天先完成一个小任务。"))
    else:
        reminders.append(("info", "还没有打卡记录，先完成今天的第一次打卡。"))
    streak = get_consecutive_days(user_id)
    if streak in {7, 21, 50, 100}:
        reminders.append(("success", f"连续打卡 {streak} 天，已达成阶段里程碑。"))
    plan_progress = get_checkin_plan_progress(user_id)
    if 0 < plan_progress < 80:
        reminders.append(("warning", f"当前活跃计划平均完成率 {plan_progress}%，低于 80%，建议复盘。"))
    return reminders

def get_daily_goal_hours(user_id):
    """从活跃学习计划获取每日学习目标（小时），无计划时默认6小时"""
    row = checkin_fetch_one(
        "SELECT daily_hours FROM study_plans WHERE user_id=? AND status='active' ORDER BY created_at DESC LIMIT 1",
        (user_id,))
    return float(row["daily_hours"]) if row and row["daily_hours"] else 6.0

def get_time_period():
    """判断当前时段"""
    h = datetime.now().hour
    if 5 <= h < 9:
        return "清晨"
    elif 9 <= h < 12:
        return "上午"
    elif 12 <= h < 14:
        return "午间"
    elif 14 <= h < 18:
        return "下午"
    elif 18 <= h < 22:
        return "晚间"
    else:
        return "深夜"

def get_flow_focus_data(user_id):
    """汇聚今日心流数据：总时长、目标、进度、打卡状态、心情、连续天数"""
    today_minutes = get_today_duration(user_id)
    pomo_count, pomo_minutes = get_today_pomodoros(user_id)
    # 番茄钟时长与打卡时长取最大值（避免重复计算）
    total_minutes = max(today_minutes, pomo_minutes) if pomo_minutes else today_minutes
    goal_hours = get_daily_goal_hours(user_id)
    goal_minutes = goal_hours * 60
    progress_pct = min(round(total_minutes / goal_minutes * 100) if goal_minutes > 0 else 0, 100)
    current_ck = get_today_checkin(user_id)
    checked_in = current_ck is not None
    mood = current_ck["mood"] if current_ck else None
    streak = get_consecutive_days(user_id)
    return {
        "total_hours": round(total_minutes / 60, 1),
        "total_minutes": total_minutes,
        "goal_hours": goal_hours,
        "progress_pct": progress_pct,
        "checked_in": checked_in,
        "mood": mood,
        "streak": streak,
    }

# ── 心流寄语：规则兜底 + LLM 生成 ──

FLOW_FALLBACKS = {
    "清晨": [
        "清晨的每一分钟都在为未来铺路，今天加油。",
        "趁晨光正好，开启专注的一天。",
        "早起的你，已经领先了大多数人。",
    ],
    "上午": [
        "上午是大脑最清醒的时段，保持这份专注。",
        "稳步推进，今天的目标正在靠近。",
        "按自己的节奏来，每一步都算数。",
    ],
    "午间": [
        "适当休整后继续，下午还有目标等你。",
        "短暂的休息是为了更好的出发。",
        "保持节奏，不急不躁。",
    ],
    "下午": [
        "下午是攻坚的好时段，持续推进。",
        "还有半天的机会，把进度再推一步。",
        "专注当下，完成比完美更重要。",
    ],
    "晚间": [
        "今天的坚持值得肯定，回顾一下收获。",
        "夜深人静正是深度学习时，但别太晚。",
        "复盘今日所学，让努力有迹可循。",
    ],
    "深夜": [
        "夜深了，今天的努力已足够，早点休息。",
        "身体是革命的本钱，明天再战。",
        "今天的每一分钟都不会白费，晚安。",
    ],
}

def pick_flow_message(user_id):
    """根据时段 + 打卡数据选择心流寄语（纯规则，零延迟）"""
    data = get_flow_focus_data(user_id)
    period = get_time_period()
    candidates = FLOW_FALLBACKS.get(period, FLOW_FALLBACKS["上午"])
    pct = data["progress_pct"]

    # 根据进度微调选句
    if pct == 0 and data["streak"] >= 7:
        idx = 0  # 有连续打卡但今天还没开始：温和提醒
    elif pct >= 80:
        idx = 1  # 快完成了：肯定
    elif pct >= 50:
        idx = 0  # 过半：鼓励保持
    else:
        idx = min(pct // 30, len(candidates) - 1)  # 根据进度梯度选

    # 深夜 + 高完成度 → 劝休息
    if period == "深夜" and pct >= 80:
        return "今日目标已达成，这份坚持值得骄傲，去休息吧。"

    return candidates[idx]

def generate_flow_message_prompt(user_id):
    """
    构建心流寄语 LLM prompt 字符串。
    调用方用 call_llm_api(prompt_str, model="mimo-v2.5", max_tokens=100) 获取 AI 寄语。
    """
    data = get_flow_focus_data(user_id)
    period = get_time_period()

    prompt = f"""你是考研备考助手。根据以下用户数据，生成一句"今日心流寄语"。

## 用户数据
- 当前时段：{period}
- 今日已学习：{data['total_hours']} 小时（每日目标 {data['goal_hours']} 小时）
- 目标完成度：{data['progress_pct']}%
- 今日是否已打卡：{'是' if data['checked_in'] else '否'}
- 今日心情：{data['mood'] or '未记录'}
- 连续打卡：{data['streak']} 天

## 输出要求
1. 只输出一句话（15-30 字），不加前缀、引号或解释
2. 语气朴素理性，不使用"亲爱的""孩子""老师""同学"等称呼
3. 时段对应态度：
   - 清晨：鼓励开启新一天，简短有力
   - 上午：肯定早间努力，提醒保持节奏
   - 午间：提醒适当休息，储备下午精力
   - 下午：关注进度推进，给予方向感
   - 晚间：回顾今日收获，肯定坚持
   - 深夜：温和提醒休息，不鼓励透支
4. 完成度对应态度：
   - 0%（未开始）：温和提醒，不施加压力
   - 1-50%：鼓励推进，肯定已付出的努力
   - 50-80%：肯定进展，提醒保持节奏
   - 80%+：赞赏坚持，鼓励收尾
5. 连续打卡 ≥ 7 天时可含蓄肯定习惯的力量"""

    return prompt

# ==================== 用户画像模块 ====================

def _profile_columns():
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("PRAGMA table_info(user_profiles)")
    columns = [col[1] for col in c.fetchall()]
    conn.close()
    return columns

def get_user_profile(user_id):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT * FROM user_profiles WHERE user_id=?", (user_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return {}
    columns = _profile_columns()
    return dict(zip(columns, row))

def save_profile_field(user_id, field, value):
    allowed = set(_profile_columns()) - {"id", "user_id", "created_at", "updated_at"}
    if field not in allowed:
        raise ValueError(f"非法字段: {field}")
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT 1 FROM user_profiles WHERE user_id=?", (user_id,))
    exists = c.fetchone()
    if exists:
        c.execute(f"UPDATE user_profiles SET {field}=?, updated_at=CURRENT_TIMESTAMP WHERE user_id=?", (value, user_id))
    else:
        c.execute(f"INSERT INTO user_profiles (user_id, {field}) VALUES (?, ?)", (user_id, value))
    conn.commit()
    conn.close()

def profile_is_complete(user_id):
    profile = get_user_profile(user_id)
    required = ["grade", "major", "target_major", "daily_hours"]
    return all(profile.get(k) for k in required)

def _safe_json_loads(raw, default=None):
    if default is None:
        default = []
    if not raw:
        return default
    try:
        return json.loads(raw)
    except (json.JSONDecodeError, TypeError):
        return default

def _display_target_schools(profile):
    raw = profile.get("target_schools")
    if not raw:
        return "未设置"
    data = _safe_json_loads(raw, {})
    if isinstance(data, dict):
        parts = [f"{k}: {v}" for k, v in data.items() if v]
        return " · ".join(parts) if parts else "未设置"
    return str(raw)

def auto_generate_tags(user_id):
    profile = get_user_profile(user_id)
    tags = {
        "common_errors": [],
        "strong_areas": [],
        "current_phase": profile.get("current_phase") or "基础",
    }
    weak_subjects = _safe_json_loads(profile.get("weak_subjects"))
    strong_subjects = _safe_json_loads(profile.get("strong_subjects"))
    if "数学" in weak_subjects:
        tags["common_errors"].extend(["计算错误", "公式混淆"])
    if "英语" in weak_subjects:
        tags["common_errors"].extend(["语法错误", "词汇量不足"])
    if "政治" in weak_subjects:
        tags["common_errors"].append("知识点遗漏")
    if "专业课" in weak_subjects:
        tags["common_errors"].append("概念理解偏差")
    if "数学" in strong_subjects:
        tags["strong_areas"].extend(["逻辑推理", "公式应用"])
    if "英语" in strong_subjects:
        tags["strong_areas"].append("阅读理解")
    if "政治" in strong_subjects:
        tags["strong_areas"].append("时政敏感度")
    tags["common_errors"] = list(dict.fromkeys(tags["common_errors"]))
    tags["strong_areas"] = list(dict.fromkeys(tags["strong_areas"]))
    save_profile_field(user_id, "common_errors", json.dumps(tags["common_errors"], ensure_ascii=False))
    save_profile_field(user_id, "strong_areas", json.dumps(tags["strong_areas"], ensure_ascii=False))
    save_profile_field(user_id, "current_phase", tags["current_phase"])
    return tags

def update_profile_from_conversation(user_id, query, answer):
    extracted = {}
    if API_KEY:
        prompt = f"""从以下对话中提取用户信息，返回 JSON 格式：
用户：{query}
AI：{answer}
请提取以下信息（如果有的话）：target_major, target_school, math_type, weak_subject, anxiety_level(1-5整数), current_phase(基础/强化/冲刺/模考)
只返回 JSON，不要其他内容。如果没有信息，返回空 JSON {{}}"""
        try:
            result = call_llm_api(prompt, model="mimo-v2.5")
            text = result.strip()
            if text.startswith("```"):
                text = re.sub(r"^```(?:json)?\s*", "", text)
                text = re.sub(r"\s*```$", "", text)
            extracted = json.loads(text)
        except Exception:
            extracted = {}
    if not extracted:
        extracted = _rule_extract_profile(query)
    field_map = {"target_major": "target_major", "math_type": "math_exam_type", "anxiety_level": "anxiety_level", "current_phase": "current_phase"}
    for src, dst in field_map.items():
        if extracted.get(src):
            save_profile_field(user_id, dst, extracted[src])
    if extracted.get("target_school"):
        schools = _safe_json_loads(get_user_profile(user_id).get("target_schools"), {})
        schools["冲刺"] = extracted["target_school"]
        save_profile_field(user_id, "target_schools", json.dumps(schools, ensure_ascii=False))
    if extracted.get("weak_subject"):
        weak = _safe_json_loads(get_user_profile(user_id).get("weak_subjects"))
        if extracted["weak_subject"] not in weak:
            weak.append(extracted["weak_subject"])
        save_profile_field(user_id, "weak_subjects", json.dumps(weak, ensure_ascii=False))
    return extracted

def _rule_extract_profile(query):
    result = {}
    school_match = re.search(r"(清华|北大|复旦|上交|浙大|中科大|南大|武大|目标.*?([^\s，。]+))", query)
    if school_match:
        result["target_school"] = school_match.group(1).replace("目标", "").strip("是")
    if "焦虑" in query or "崩溃" in query:
        result["anxiety_level"] = 4
    for phase in ("基础", "强化", "冲刺", "模考"):
        if phase in query:
            result["current_phase"] = phase
            break
    for mt in ("数一", "数二", "数三", "199管综"):
        if mt in query:
            result["math_type"] = mt
            break
    return result

def check_content_safety(answer, context=None):
    context = context or {}
    violations = []
    if re.search(r"(报录比|招生人数|分数线).*\d+", answer):
        if "来源" not in answer and "数据来源" not in answer:
            violations.append("数据未标注来源")
    if re.search(r"(保证|承诺|一定).*(上岸|录取|通过)", answer):
        violations.append("禁止承诺录取")
    if re.search(r"(政策|规定|要求).*解读", answer):
        if "建议核实" not in answer and "官方" not in answer:
            violations.append("政策解读未标注核实建议")
    high_risk_keywords = ["分数线", "报录比", "招生人数", "政策变化"]
    query = context.get("query", "")
    if any(kw in query for kw in high_risk_keywords):
        if not context.get("rag_used"):
            violations.append("高风险问题未使用 RAG 检索")
    return violations

# ==================== 学习规划模块 ====================

def determine_phase():
    month = datetime.now().month
    if 3 <= month <= 6:
        return "基础阶段"
    elif 7 <= month <= 9:
        return "强化阶段"
    elif 10 <= month <= 11:
        return "提升阶段"
    elif month == 12:
        return "冲刺阶段"
    else:
        return "基础阶段"

PHASE_TEMPLATES = {
    "基础阶段": {"数学": ["教材通读", "基础概念理解", "基础题型练习", "公式推导"], "英语": ["词汇积累", "长难句解析", "阅读基础", "写作基础"], "政治": ["教材通读", "基本概念理解", "选择题练习"], "专业课": ["教材通读", "核心概念理解", "基础题型练习"]},
    "强化阶段": {"数学": ["专项突破", "大量刷题", "错题整理", "知识体系建立"], "英语": ["阅读强化", "写作强化", "翻译强化", "新题型练习"], "政治": ["重点章节强化", "选择题强化", "分析题练习"], "专业课": ["重点章节强化", "真题研究", "专题训练"]},
    "提升阶段": {"数学": ["真题实战", "查漏补缺", "模考检验", "高频考点强化"], "英语": ["真题实战", "写作模板", "阅读技巧", "完形填空"], "政治": ["真题实战", "时政热点", "分析题强化", "模拟考试"], "专业课": ["真题实战", "模拟考试", "重点难点突破"]},
    "冲刺阶段": {"数学": ["高频考点押题", "错题回顾", "公式速记", "模拟考试"], "英语": ["作文模板强化", "阅读技巧", "词汇巩固", "模拟考试"], "政治": ["时政热点", "分析题押题", "选择题速刷", "模拟考试"], "专业课": ["高频考点", "模拟考试", "重点难点回顾"]}
}

def get_subject_weights(math_type):
    weights = {
        "数一": {"数学": 0.35, "英语": 0.20, "政治": 0.15, "专业课": 0.30},
        "数二": {"数学": 0.35, "英语": 0.20, "政治": 0.15, "专业课": 0.30},
        "数三": {"数学": 0.35, "英语": 0.20, "政治": 0.15, "专业课": 0.30},
        "不考数学": {"英语": 0.30, "政治": 0.20, "专业课": 0.50},
        "199管综": {"管综": 0.40, "英语": 0.30, "政治": 0.15, "专业课": 0.15}
    }
    return weights.get(math_type, weights["数一"])

def calculate_daily_hours(daily_hours, math_type):
    weights = get_subject_weights(math_type)
    return {sub: round(daily_hours * w, 1) for sub, w in weights.items()}

def save_plan(user_id, plan_name, target_exam_date, math_type, daily_hours, weight_config, phase):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    now_str = datetime.now().strftime("%Y-%m-%d")
    config_json = json.dumps(weight_config, ensure_ascii=False)
    c.execute("INSERT INTO study_plans (user_id, plan_name, target_exam_date, math_type, daily_hours, subjects_config, current_phase, status, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)",
              (user_id, plan_name, target_exam_date, math_type, daily_hours, config_json, phase, now_str, now_str))
    plan_id = c.lastrowid
    conn.commit()
    conn.close()
    return plan_id

def save_task(user_id, plan_id, task_type, subject, task_name, description, target_date, est_hours, priority=3):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    now_str = datetime.now().strftime("%Y-%m-%d")
    c.execute("INSERT INTO plan_tasks (plan_id, user_id, task_type, subject, task_name, description, target_date, estimated_hours, priority, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
              (plan_id, user_id, task_type, subject, task_name, description, target_date, est_hours, priority, now_str))
    conn.commit()
    conn.close()

def get_user_tasks(user_id):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT id, subject, task_name, description, target_date, estimated_hours, status FROM plan_tasks WHERE user_id = ? ORDER BY target_date", (user_id,))
    res = []
    for row in c.fetchall():
        res.append({"id": row[0], "subject": row[1], "task_name": row[2], "description": row[3], "target_date": row[4], "estimated_hours": row[5], "status": row[6]})
    conn.close()
    return res

def update_task_status(task_id, new_status):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    completed_at = datetime.now().strftime("%Y-%m-%d") if new_status == "completed" else None
    c.execute("UPDATE plan_tasks SET status = ?, completed_at = ? WHERE id = ?", (new_status, completed_at, task_id))
    conn.commit()
    conn.close()

def calculate_progress(user_id):
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) total, COALESCE(SUM(CASE WHEN status='completed' THEN 1 ELSE 0 END), 0) completed FROM plan_tasks WHERE user_id = ?", (user_id,))
    row = c.fetchone()
    total_tasks = row[0] if row else 0
    completed_tasks = row[1] if row else 0
    c.execute("SELECT subject, COUNT(*) total, COALESCE(SUM(CASE WHEN status='completed' THEN 1 ELSE 0 END), 0) completed FROM plan_tasks WHERE user_id = ? GROUP BY subject", (user_id,))
    subjects = {}
    for row in c.fetchall():
        sub, t, cpl = row
        rate = round(cpl / max(t, 1) * 100, 1)
        subjects[sub] = {"total": t, "completed": cpl, "completion_rate": rate}
    c.execute("SELECT planned_hours, actual_hours FROM plan_progress WHERE user_id = ? ORDER BY record_date DESC LIMIT 7", (user_id,))
    recent = c.fetchall()
    deviation = 0.0
    if recent:
        total_plan = sum(r[0] for r in recent)
        total_act = sum(r[1] for r in recent)
        if total_plan > 0:
            deviation = round((total_plan - total_act) / total_plan * 100, 1)
    conn.close()
    completion_rate = round(completed_tasks / max(total_tasks, 1) * 100, 1)
    return {"total_tasks": total_tasks, "completed_tasks": completed_tasks, "completion_rate": completion_rate, "deviation": deviation, "subjects": subjects}

RECOVERY_STRATEGIES = {
    1: {"name": "删减低频考点", "description": "剔除近5年未考的知识点", "trigger": "deviation > 20%", "action": "删除低频考点任务，聚焦高频内容"},
    2: {"name": "聚焦高频核心", "description": "集中精力攻克高频+必考内容", "trigger": "deviation > 30%", "action": "将高频考点任务优先级提升为最高"},
    3: {"name": "压缩次要科目", "description": "提高优势科目用时，控制短板科目投入", "trigger": "deviation > 25%", "action": "调整科目权重，增加优势科目时间"},
    4: {"name": "切换速通模式", "description": "只看知识框架 + 重点题型，不全做", "trigger": "deviation > 40%", "action": "简化任务内容，只保留核心框架"},
    5: {"name": "错题优先", "description": "优先做错题，而非刷新题", "trigger": "连续3天完成率 < 60%", "action": "将错题复习任务优先级提升"},
    6: {"name": "调整作息", "description": "增加每日有效学习时长", "trigger": "连续5天完成率 < 70%", "action": "建议调整作息，增加学习时长"},
}

def select_recovery_strategy(deviation, recent_completion_rates):
    strategies = []
    if deviation > 40:
        strategies.append(4)
    if deviation > 30:
        strategies.append(2)
    if deviation > 25:
        strategies.append(3)
    if deviation > 20:
        strategies.append(1)
    if len(recent_completion_rates) >= 3:
        avg_rate = sum(recent_completion_rates[:3]) / 3
        if avg_rate < 60:
            strategies.append(5)
    if len(recent_completion_rates) >= 5:
        avg_rate5 = sum(recent_completion_rates[:5]) / 5
        if avg_rate5 < 70:
            strategies.append(6)
    return strategies

def generate_plan(user_id, target_date, math_type, daily_hours):
    # 获取用户画像
    profile = get_user_profile(user_id)
    weights = get_subject_weights(math_type)
    phase = determine_phase()
    daily_sub_hours = calculate_daily_hours(daily_hours, math_type)
    days_remaining = (target_date - datetime.now().date()).days

    # 从画像提取个性化信息（只取有效值）
    weak_subjects = _safe_json_loads(profile.get("weak_subjects"))
    strong_subjects = _safe_json_loads(profile.get("strong_subjects"))
    target_major = profile.get("target_major") or ""
    target_schools = _display_target_schools(profile)
    anxiety_level = profile.get("anxiety_level")
    undergraduate_major = profile.get("undergraduate_major") or ""
    undergraduate_level = profile.get("undergraduate_level") or ""
    is_cross_major = profile.get("is_cross_major") or ""

    # 弱科加权（优先补弱科，上限45%）
    if weak_subjects:
        for sub in weak_subjects:
            if sub in weights:
                weights[sub] = min(weights[sub] * 1.2, 0.45)
        total = sum(weights.values())
        weights = {k: round(v / total, 3) for k, v in weights.items()}
        daily_sub_hours = {k: round(daily_hours * v, 1) for k, v in weights.items()}

    tasks = []
    for subject, weight in weights.items():
        sub_h = daily_hours * weight
        task_list = PHASE_TEMPLATES[phase].get(subject, [])
        priority = 1 if subject in weak_subjects else 3
        for task_name in task_list:
            single_task_h = round(sub_h / max(len(task_list), 1), 1)
            tasks.append({"subject": subject, "task_name": task_name, "estimated_hours": single_task_h, "priority": priority})

    # 构建个性化 prompt（只包含有值的字段）
    profile_lines = []
    if target_schools and target_schools != "未设置":
        profile_lines.append(f"- 目标院校：{target_schools}")
    if target_major:
        profile_lines.append(f"- 目标专业：{target_major}")
    if undergraduate_major:
        profile_lines.append(f"- 本专业：{undergraduate_major}")
    if undergraduate_level:
        profile_lines.append(f"- 本科院校级别：{undergraduate_level}")
    if is_cross_major and is_cross_major == "是":
        profile_lines.append(f"- 是否跨考：是（跨考生需额外注意专业课基础）")
    if weak_subjects:
        profile_lines.append(f"- 弱科：{', '.join(weak_subjects)}")
    if strong_subjects:
        profile_lines.append(f"- 强科：{', '.join(strong_subjects)}")
    if anxiety_level:
        profile_lines.append(f"- 焦虑程度：{anxiety_level}/5")

    profile_text = "\n".join(profile_lines) if profile_lines else "（用户尚未填写画像信息）"

    prompt = f"""你是考研学习规划专家。请生成一份结构化的学习时间表。

## 用户画像
{profile_text}
- 每日学习时长：{daily_hours}小时
- 数学类型：{math_type}

## 考试规划
- 目标日期：{target_date.strftime("%Y-%m-%d")}
- 剩余天数：{days_remaining}天
- 当前阶段：{phase}
- 各科权重：{json.dumps(weights, ensure_ascii=False)}

## 输出格式要求

请以 **Markdown 表格 + 简要说明** 的格式输出，不要长篇抒情，语气简洁专业：

### 1. 每日时间表
用表格输出，例如：
```
| 时间段 | 科目 | 任务重点 | 建议时长 |
|--------|------|----------|----------|
| 08:00-12:00 | 数学 | 专题突破+真题训练 | 4h |
| 14:00-17:00 | 英语 | 阅读理解+单词 | 3h |
| 19:00-21:00 | 政治 | 章节梳理+选择题 | 2h |
| 21:00-22:00 | 总结整理 | 错题回顾+明日计划 | 1h |
```

### 2. 时间段分配原则
- 上午安排需要高度专注的科目（如数学、专业课）
- 下午安排语言类科目（如英语）
- 晚上安排记忆和政治类科目
- 根据弱科（{', '.join(weak_subjects) if weak_subjects else '无'}）优先分配黄金时间段
- 每科之间留10-15分钟休息

### 3. 每周计划概述
- 周一至周五：按时间表执行
- 周六：模拟测试+批改分析
- 周日：本周错题复习+下周计划调整

请直接输出，无需额外说明。"""
    description = call_llm_api(prompt, model="mimo-v2.5")
    return {"description": description, "tasks": tasks, "phase": phase, "weights": weights, "daily_sub_hours": daily_sub_hours}

def calc_recall(stability, days):
    if days <= 0:
        return 1.0
    return max(0, min(1, math.exp(-days / (stability + 0.1))))

def needs_review(recall_prob, threshold=0.3):
    return recall_prob < threshold

def save_feynman_record(user_id, mode, question_text, user_answer, ai_evaluation, score_correct, score_expression, score_authentic, total_score):
    """保存费曼学习法记录"""
    conn = sqlite3.connect(MEMORY_DB)
    conn.execute(
        """INSERT INTO feynman_records
           (user_id, mode, question_text, user_answer, ai_evaluation, score_correct, score_expression, score_authentic, total_score)
           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
        (user_id, mode, question_text, user_answer, ai_evaluation, score_correct, score_expression, score_authentic, total_score))
    conn.commit()
    conn.close()

def get_feynman_history(user_id, limit=10):
    """获取费曼学习法历史记录"""
    return checkin_fetch_all(
        """SELECT * FROM feynman_records WHERE user_id=?
           ORDER BY created_at DESC LIMIT ?""",
        (user_id, limit))

# ========== 调试日志系统 ==========
def _init_debug_log():
    """初始化调试日志缓冲区"""
    if "debug_logs" not in st.session_state:
        st.session_state.debug_logs = []
    if "debug_mode" not in st.session_state:
        st.session_state.debug_mode = False


def _append_debug_log(entry):
    """追加一条调试日志（保留最近 50 条）"""
    _init_debug_log()
    st.session_state.debug_logs.append(entry)
    if len(st.session_state.debug_logs) > 50:
        st.session_state.debug_logs = st.session_state.debug_logs[-50:]


def call_llm_api(prompt, model="mimo-v2.5", max_tokens=2000, temperature=0.3):
    """调用 LLM API（非流式）+ 调试日志 + 自动重试"""
    _init_debug_log()
    t0 = datetime.now()
    log_entry = {
        "time": t0.strftime("%H:%M:%S"),
        "model": model,
        "max_tokens": max_tokens,
        "temperature": temperature,
        "prompt_len": len(prompt),
        "prompt_preview": prompt[:250],
    }
    data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": temperature
    }
    last_error = None
    max_retries = 5
    for attempt in range(max_retries):
        if attempt > 0:
            wait_s = 2 ** attempt  # 2s, 4s, 8s, 16s 指数退避
            time.sleep(wait_s)
        try:
            req = urllib.request.Request(
                API_BASE + "/chat/completions",
                data=json.dumps(data).encode("utf-8"),
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                method="POST"
            )
            with urllib.request.urlopen(req, timeout=90) as resp:
                raw_body = resp.read().decode("utf-8")
                msg = json.loads(raw_body)["choices"][0]["message"]
                c = msg.get("content")
                # 修复: content 可能是空列表 []（MiMo 思维链特征）
                raw_full = c if isinstance(c, str) and c != "" else ""
                used_reasoning = False
                if not raw_full:
                    raw_full = msg.get("reasoning_content") or ""
                    used_reasoning = bool(raw_full)
                # 记录原始响应信息
                log_entry["raw_content_len"] = len(c) if isinstance(c, str) else (len(c) if isinstance(c, list) else 0)
                log_entry["raw_reasoning_len"] = len(msg.get("reasoning_content") or "")
                log_entry["raw_full_len"] = len(raw_full)
                log_entry["raw_preview"] = raw_full[:300]
                log_entry["used_reasoning"] = used_reasoning
                if attempt > 0:
                    log_entry["retry_attempt"] = attempt
                # 清洗
                cleaned = _clean_mimo_output(raw_full, prompt, used_reasoning=used_reasoning)
                log_entry["cleaned_len"] = len(cleaned)
                log_entry["cleaned_preview"] = cleaned[:300]
                log_entry["elapsed_ms"] = int((datetime.now() - t0).total_seconds() * 1000)
                log_entry["status"] = "ok"
                _append_debug_log(log_entry)
                return cleaned
        except (urllib.error.URLError, ConnectionResetError, TimeoutError, OSError) as e:
            last_error = e
            if attempt < max_retries - 1:
                continue
        except Exception as e:
            # 非网络错误（如 JSON 解析失败），不重试
            log_entry["status"] = "error"
            log_entry["error"] = str(e)
            log_entry["traceback"] = traceback.format_exc()[-500:]
            log_entry["elapsed_ms"] = int((datetime.now() - t0).total_seconds() * 1000)
            _append_debug_log(log_entry)
            raise
    # 所有重试均失败
    log_entry["status"] = "error"
    log_entry["error"] = str(last_error)
    log_entry["retries"] = 3
    log_entry["elapsed_ms"] = int((datetime.now() - t0).total_seconds() * 1000)
    _append_debug_log(log_entry)
    raise last_error


def _clean_mimo_output(raw_text, prompt="", used_reasoning=False):
    """清洗 MiMo 思维链输出：优先找答案起始标记 → 行过滤 → 编号行兜底
    used_reasoning=True 表示 content 为空、全量回退到 reasoning_content"""
    if not raw_text or len(raw_text) < 5:
        return raw_text
    text = raw_text.strip()
    # 1. 去掉 prompt 回显
    if prompt and text.startswith(prompt.strip()[:40]):
        text = text[len(prompt.strip()):].strip()

    # === 核心策略：优先找「答案起始」标记，截取之后的内容 ===
    _answer_markers = (
        # 中文答案标记
        '正式回答', '给出答案', '输出如下', '答案如下', '解答如下',
        '最终答案', '回答如下', '现在回答', '开始答题',
        # 结构化标记（来自 prompt 的 format 要求）
        '[题目]', '[解答]', '[答案]',
        # 英文
        'Answer:', 'Solution:', 'Here is',
    )
    best_marker_pos = len(text) + 1  # 取最早出现的标记（最小位置）
    for marker in _answer_markers:
        pos = text.find(marker)
        if pos != -1 and pos < best_marker_pos:
            best_marker_pos = pos
    if best_marker_pos < len(text) and best_marker_pos > 0:
        # 从标记所在行开始截取
        line_start = text.rfind('\n', 0, best_marker_pos)
        text = text[line_start + 1:] if line_start != -1 else text[best_marker_pos:]

    # === 行级过滤 ===
    _think_starts = (
        '好的', '让我', '首先', '根据', '综上', '因此', '注意', '这个', '该知',
        '我们', '需要', '可以', '这里', '现在', '接下来', '最后', '总的', '所以',
        'Okay', 'Let', 'First', 'I need', 'The user',
        '用户', '问题是', '题目是', '要解', '已知', '分析', '考虑', '理解',
        '回顾', '判断', '比较', '计算', '推导', '综合', '综上所',
        '让我来', '我来', '我明白', '收到', '嗯，', '嗯,', '好，', '好,',
        '首先，', '接下来，', '最后，', '然后，', '接着，',
        '根据题目', '从题目', '由题意', '观察', '对比', '得到',
        # 更激进的过滤
        '这道题', '此题', '该题', '本题', '考察', '涉及', '属于',
        '目的是', '目标是', '任务是', '要求是',
    )
    _think_keywords = (
        '知识点是', '用户要求', '定义如下', '给出答案', '任务是', '需要生成',
        '我的思考', '思路如下', '步骤是', '解题思路', '推理过程',
        '思考过程', '分析如下', '详细说明', '解释如下',
        '我来分析', '逐步思考', '逐步分析',
    )
    lines = text.split('\n')
    total_lines = len(lines)
    clean_lines = []
    filtered_reasons = []
    think_line_count = 0
    for line in lines:
        s = line.strip()
        if not s:
            clean_lines.append(line)  # 保留空行，维持排版
            continue
        if s.startswith(_think_starts):
            think_line_count += 1
            if len(filtered_reasons) < 10:
                filtered_reasons.append(f"think_start: {s[:30]}")
            continue
        if any(kw in s for kw in _think_keywords):
            think_line_count += 1
            if len(filtered_reasons) < 10:
                filtered_reasons.append(f"keyword: {s[:30]}")
            continue
        clean_lines.append(line)
    all_filtered = (len([l for l in clean_lines if l.strip()]) == 0)
    if not clean_lines or all_filtered:
        _record_clean_stats(total_lines, total_lines, 0, len(lines),
                           False, True, filtered_reasons)
        return text  # 全被过滤了，返回原文

    # === 找第一个编号行 ===
    start_idx = 0
    for i, line in enumerate(clean_lines):
        if re.match(r'^\d+[\.\、\)\)]\s', line.strip()):
            start_idx = i
            break
    result_lines = clean_lines[start_idx:]

    # === reasoning 回退时的强力兜底 ===
    think_ratio = think_line_count / max(total_lines, 1)
    if used_reasoning and think_ratio > 0.4:
        # 优先提取编号行；没有则取非空行
        numbered = [l for l in clean_lines if re.match(r'^\d+[\.\、\)\)]', l.strip())]
        if numbered:
            result_lines = numbered
            filtered_reasons.append("reasoning_fb: 编号行兜底")
        else:
            # 取后 50% 的行（思维链通常在前半部分）
            non_empty = [l for l in clean_lines if l.strip()]
            half = max(1, len(non_empty) // 2)
            result_lines = non_empty[-half:]
            filtered_reasons.append("reasoning_fb: 后半段兜底")

    # === 截断保护 ===
    was_truncated = len(result_lines) > 20
    if was_truncated:
        result_lines = result_lines[:15]

    _record_clean_stats(total_lines, total_lines - len(clean_lines), start_idx,
                       len(result_lines), was_truncated, all_filtered, filtered_reasons,
                       think_ratio=think_ratio, used_reasoning=used_reasoning)
    return '\n'.join(result_lines).strip()


def _record_clean_stats(total_lines, filtered_count, start_idx, result_count,
                        was_truncated, all_filtered, filtered_reasons,
                        think_ratio=0, used_reasoning=False):
    """记录清洗统计到最近一条调试日志"""
    logs = st.session_state.get("debug_logs", [])
    if logs and logs[-1].get("status") in ("ok", None):
        logs[-1]["clean_stats"] = {
            "total_lines": total_lines,
            "filtered_count": filtered_count,
            "think_ratio": f"{think_ratio:.0%}",
            "used_reasoning": used_reasoning,
            "start_idx": start_idx,
            "result_count": result_count,
            "was_truncated": was_truncated,
            "all_filtered": all_filtered,
            "filtered_reasons": filtered_reasons[:5],
        }


def call_llm_stream(prompt, model="mimo-v2.5", max_tokens=800, system_prompt=""):
    """流式调用 LLM，返回原始完整文本"""
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})
    data = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.3,
        "stream": True,
    }
    req = urllib.request.Request(
        API_BASE + "/chat/completions",
        data=json.dumps(data).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
        method="POST"
    )
    raw_full = ""
    with urllib.request.urlopen(req, timeout=120) as resp:
        buffer = ""
        while True:
            chunk = resp.read(1024)
            if not chunk:
                break
            buffer += chunk.decode("utf-8", errors="ignore")
            while "\n" in buffer:
                line, buffer = buffer.split("\n", 1)
                line = line.strip()
                if not line.startswith("data: "):
                    continue
                payload = line[6:]
                if payload == "[DONE]":
                    break
                try:
                    obj = json.loads(payload)
                    delta_obj = obj.get("choices", [{}])[0].get("delta", {})
                    c = delta_obj.get("content")
                    delta = c if isinstance(c, str) else ""
                    if not delta:
                        delta = delta_obj.get("reasoning_content") or ""
                    if delta:
                        raw_full += delta
                except Exception:
                    pass
    return raw_full.strip()

# 费曼学习法评价提示词
CONCEPT_EVAL_PROMPT = """你是考研数学辅导专家，同时也是教育心理学专家。你的任务是评价学生对数学概念的理解和表达能力。

## 评价维度

### 1. 概念理解（6分）
- 6分：理解准确，能用自己的话清晰表达
- 5分：理解准确，但表达不够清晰
- 4分：理解基本正确，但有遗漏
- 3分：理解有偏差，但核心思想正确
- 2分：理解有重大偏差
- 1分：理解完全错误
- 0分：未作答

### 2. 表达能力（2分）
- 2分：表达清晰，逻辑连贯，有个人风格
- 1分：表达基本清晰，但缺乏个人理解
- 0分：表达混乱或照搬教材

### 3. 书写真实性（2分）
- 2分：明显是自主思考，有个人表达习惯
- 1分：可能是自主思考，但有部分可疑特征
- 0分：明显是复制粘贴

## 输出格式

[总分] X/10分

[概念理解] X/6分
- 理解准确性：[准确/基本准确/有偏差/错误]
- 具体评价：...

[表达能力] X/2分
- 表达清晰度：[清晰/基本清晰/混乱]
- 具体评价：...

[书写真实性] X/2分
- 判断结果：[自主作答/可能复制/明显复制]
- 判断依据：...

[详细评价]
（具体分析）

[改进建议]
（给出建议）

## 题目
{question}

## 学生答案
{answer}"""

PROBLEM_EVAL_PROMPT = """你是考研数学辅导专家，同时也是教育心理学专家。你的任务是评价学生的解题能力和思维方式。

## 评价维度

### 1. 解题正确性（5分）
- 5分：答案完全正确，步骤完整
- 4分：答案正确，但有小瑕疵
- 3分：答案基本正确，但有重要错误
- 2分：答案错误，但思路基本正确
- 1分：答案错误，思路也有问题
- 0分：未作答或完全错误

### 2. 解题过程（3分）
- 3分：步骤清晰，逻辑连贯，有个人理解
- 2分：步骤基本清晰，但缺乏个人理解
- 1分：步骤混乱，但有尝试
- 0分：未展示解题过程

### 3. 书写真实性（2分）
- 2分：明显是自主思考，有自然思考痕迹
- 1分：可能是自主思考，但有部分可疑特征
- 0分：明显是复制粘贴

## 输出格式

[总分] X/10分

[解题正确性] X/5分
- 答案正确性：[正确/部分正确/错误]
- 具体评价：...

[解题过程] X/3分
- 步骤完整性：[完整/基本完整/不完整]
- 逻辑清晰度：[清晰/基本清晰/混乱]
- 具体评价：...

[书写真实性] X/2分
- 判断结果：[自主作答/可能复制/明显复制]
- 判断依据：...

[详细评价]
（具体分析）

[改进建议]
（给出建议）

## 题目
{question}

## 学生答案
{answer}"""

def update_memory(kid, is_mastered, error_type="", mastery_score=0):
    init_memory_db()
    uid = st.session_state.get("user_id", 1)
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT times_correct, times_wrong, stability FROM knowledge_mastery WHERE knowledge_id=? AND user_id=?", (kid, uid))
    row = c.fetchone()

    if row:
        old_correct, old_wrong, old_stability = row
        times_correct = old_correct + (1 if is_mastered else 0)
        times_wrong = old_wrong + (0 if is_mastered else 1)
        stability = old_stability * 1.1 if is_mastered else max(0.5, old_stability * 0.9)
    else:
        times_correct = 1 if is_mastered else 0
        times_wrong = 0 if is_mastered else 1
        stability = 1.0

    status = "掌握" if is_mastered else "学习中"

    # 先删除旧记录（避免重复堆积）
    c.execute("DELETE FROM knowledge_mastery WHERE knowledge_id=? AND user_id=?", (kid, uid))
    c.execute("""INSERT INTO knowledge_mastery
        (knowledge_id, user_id, status, times_correct, times_wrong, stability, last_review, error_type)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
        (kid, uid, status, times_correct, times_wrong, stability, datetime.now(), error_type))

    c.execute("""INSERT INTO user_performance
        (user_id, knowledge_id, is_correct, error_type, mastery_score, created_at)
        VALUES (?, ?, ?, ?, ?, ?)""",
        (uid, kid, 1 if is_mastered else 0, error_type, mastery_score, datetime.now()))

    conn.commit()
    conn.close()

def get_memory_stats():
    init_memory_db()
    uid = st.session_state.get("user_id", 1)
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM knowledge_mastery WHERE status='掌握' AND user_id=?", (uid,))
    mastered = c.fetchone()[0] or 0
    c.execute("SELECT COUNT(*) FROM knowledge_mastery WHERE status='学习中' AND user_id=?", (uid,))
    learning = c.fetchone()[0] or 0
    c.execute("SELECT COUNT(*) FROM knowledge_mastery WHERE user_id=?", (uid,))
    total = c.fetchone()[0] or 0
    conn.close()
    return {"mastered": mastered, "learning": learning, "total": total}

def get_weak_points():
    init_memory_db()
    uid = st.session_state.get("user_id", 1)
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("""SELECT knowledge_id, times_wrong, times_correct, status, stability, error_type
        FROM knowledge_mastery WHERE times_wrong > 0 AND user_id=? ORDER BY times_wrong DESC LIMIT 10""", (uid,))
    results = c.fetchall()
    conn.close()
    weak_points = []
    for r in results:
        recall = calc_recall(r[4] or 1.0, 3)
        weak_points.append({"knowledge_id": r[0], "times_wrong": r[1], "times_correct": r[2], "status": r[3] or "学习中", "recall": recall, "error_type": r[5] or "理解不清"})
    return weak_points

def get_review_candidates():
    init_memory_db()
    uid = st.session_state.get("user_id", 1)
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("""SELECT knowledge_id, mastery_level, status, stability, last_review
        FROM knowledge_mastery WHERE user_id=? ORDER BY last_review DESC""", (uid,))
    results = c.fetchall()
    conn.close()

    candidates = []
    for r in results:
        kid, mastery, status, stability, last_review = r
        if last_review:
            days = (datetime.now() - datetime.fromisoformat(str(last_review))).days
        else:
            days = 30

        recall = calc_recall(stability or 1.0, days)

        if needs_review(recall) or status != "掌握":
            candidates.append({
                "knowledge_id": kid,
                "mastery_level": mastery or 0,
                "status": status or "陌生",
                "recall": recall,
                "urgency": 1 - recall
            })

    candidates.sort(key=lambda x: x["urgency"], reverse=True)
    return candidates[:10]

def create_review_challenge(kid):
    init_memory_db()
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("""INSERT INTO review_challenges (knowledge_id, challenge_type, created_at)
        VALUES (?, '自动复习', ?)""", (kid, datetime.now()))
    conn.commit()
    conn.close()

# ==================== 复习题目生成 ====================

def _fix_latex(text):
    r"""Convert \(\) to $, \[\] to $$"""
    text = text.replace("\\( ", "$ ").replace(" \\)", " $").replace("\\(", "$").replace("\\)", "$").replace("\\[", "$$").replace("\\]", "$$")
    return text

def _collapse_math(text):
    """Merge line breaks inside $$...$$ and $...$ blocks to prevent st.markdown <br> splitting"""
    if "$$" not in text and "$" not in text:
        return text
    text = re.sub(r'(\$\$)(.+?)(\$\$)', lambda m: m.group(1) + m.group(2).replace('\n', ' ') + m.group(3), text, flags=re.DOTALL)
    text = re.sub(r'(?<!\$)(\$)([^$\n]+)(\$)(?!\$)', lambda m: m.group(1) + m.group(2).replace('\n', ' ') + m.group(3), text)
    return text

def _escape_md(text):
    """Prevent st.markdown from eating LaTeX \\ (matrix row sep, newline commands)"""
    if "\\" not in text:
        return text
    # Protect $$...$$ blocks: double backslashes inside them
    text = re.sub(r'(\$\$.+?\$\$)', lambda m: m.group(1).replace("\\\\", "\\\\\\\\"), text, flags=re.DOTALL)
    # Protect $...$ blocks
    text = re.sub(r'(?<!\$)(\$[^$\n]+\$)(?!\$)', lambda m: m.group(1).replace("\\\\", "\\\\\\\\"), text)
    return text

def _katex_refresh():
    st.html("<script>if(typeof renderMathInElement!=='undefined'){renderMathInElement(document.body,{delimiters:[{left:'$$',right:'$$',display:true},{left:'$',right:'$',display:false},{left:'\\\\(',right:'\\\\)',display:false}],throwOnError:!1,strict:!1})}</script>")

def render_qa_cards(raw_text, columns=2, typing=False):
    """渲染练习题：全宽卡片，选项直接显示，答案/解析折叠。typing=True 时逐字打字效果"""
    if not raw_text:
        return
    import time as _time
    blocks = raw_text.split("---")
    qi = 0
    for block in blocks:
        block = block.strip()
        if not block or "Q:" not in block:
            continue
        lines = block.split("\n")
        question = ""
        options = []
        answer = ""
        explain = ""
        collecting_question = False
        collecting_explain = False
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.startswith("Q:") or line.startswith("Q："):
                collecting_question = True
                collecting_explain = False
                q_text = line.split(":", 1)[-1].split("：", 1)[-1].strip()
                if q_text:
                    question = q_text
            elif line.startswith(("A)", "A.", "A、", "B)", "B.", "C)", "C.", "D)", "D.")):
                collecting_question = False
                collecting_explain = False
                # Wrap bare math in $ for KaTeX rendering
                m = re.match(r'^([A-D][).、])\s*(.+)', line)
                if m:
                    pref, cont = m.group(1), m.group(2)
                    # 检测 $$ 和中文混合：将 $$ 替换为 $（内联模式）
                    if "$$" in cont and re.search(r'[\u4e00-\u9fff]', cont):
                        cont = cont.replace("$$", "$")
                    elif "\\" in cont and "$" not in cont:
                        cont = f"${cont}$"
                    options.append(_fix_latex(f"{pref} {cont}"))
                else:
                    options.append(_fix_latex(line))
            elif line.startswith("ANSWER:") or line.startswith("答案:"):
                collecting_question = False
                collecting_explain = False
                answer = line.split(":", 1)[-1].split("：", 1)[-1].strip()
            elif line.startswith("EXPLAIN:") or line.startswith("解析:"):
                collecting_question = False
                collecting_explain = True
                explain = line.split(":", 1)[-1].split("：", 1)[-1].strip()
            elif line.startswith(("[ANSWER", "[QUIZ]", "[END]", "[KNOWLEDGE")):
                continue
            else:
                if collecting_explain:
                    explain = (explain + " " + line).strip()
                elif collecting_question:
                    question = (question + " " + line).strip()
                elif not question:
                    question = line
        # 题干中裸数学记号自动包 $（仅当无任何已有分隔符 \$ \[ \( 时触发）
        if re.search(r'\\[a-zA-Z]', question) and not re.search(r'[\$\\[]', question):
            question = re.sub(r'(\\[a-zA-Z]+(?:\{[^}]*\})*(?:_\{[^}]*\})*(?:\^\{[^}]*\})*|\w+\^\{?\d+\}?|\\,?[a-z]+|\w+\'\(\d+\))', r'$\1$', question)

        st.markdown(f"<div style='background:#fff;border-radius:16px;padding:clamp(14px,3vw,40px);box-shadow:0 1px 3px rgba(0,0,0,0.04);margin-bottom:24px;font-size:16px;overflow-x:auto;'>", unsafe_allow_html=True)
        st.caption(f"第{qi+1}题")

        if typing:
            # 逐字打字效果：题干（LaTeX 公式整体插入）
            question_placeholder = st.empty()
            _typing_display(question_placeholder, _escape_md(_collapse_math(_fix_latex(question))), delay=0.03)
            # 逐字打字效果：选项
            if options:
                for opt in options[:4]:
                    opt_placeholder = st.empty()
                    _typing_display(opt_placeholder, _escape_md(_collapse_math(opt)), delay=0.02)
        else:
            st.markdown(_escape_md(_collapse_math(_fix_latex(question))))
            if options:
                for opt in options[:4]:
                    st.markdown(_escape_md(_collapse_math(opt)))

        if answer or explain:
            with st.expander("答案与解析", expanded=False):
                if answer:
                    st.markdown(f"**正确答案**: {_escape_md(_collapse_math(_fix_latex(answer)))}")
                if explain:
                    st.markdown(_escape_md(_collapse_math(_fix_latex(explain))))
                _katex_refresh()
        st.markdown("</div>", unsafe_allow_html=True)
        qi += 1
        if qi >= 1:
            break
    _katex_refresh()

def _extract_summary(text, max_lines=3):
    """从 corpus 正文提取 1-2 句摘要（开篇核心定义 + 第一行要点）"""
    lines = [l.strip() for l in text.split("\n") if l.strip() and not l.startswith("#")]
    summary = []
    for l in lines:
        if l.startswith("- **") or l.startswith("* **"):
            summary.append(l.strip("-* "))
        elif len(l) > 20 and l[0].isalpha() and len(summary) < max_lines:
            summary.append(l[:200])
        if len(summary) >= max_lines:
            break
    return "；".join(summary) if summary else ""

def _clean_knowledge_name(kid):
    """004-导数的定义与几何意义.md → 导数的定义与几何意义"""
    if not kid:
        return kid
    name = kid.replace(".md", "")
    m = re.match(r'^\d+\-(.+)$', name)
    return m.group(1) if m else name

def generate_review_questions(knowledge_points):
    if not knowledge_points:
        return {"error": "无复习知识点", "questions": ""}

    try:
        # 构建知识点列表 + 全文摘要
        corpus = load_corpus()
        kb_lines = []
        contexts = []
        for i, kp in enumerate(knowledge_points[:3]):
            kid = kp.get("knowledge_id", "")
            clean_name = _clean_knowledge_name(kid)
            kb_lines.append(f"{i+1}. {clean_name}")
            # 从 corpus 找到正文并提取摘要
            doc_text = ""
            for doc in corpus:
                if doc["id"] == kid:
                    doc_text = doc["text"]
                    break
            summary = _extract_summary(doc_text) if doc_text else ""
            if summary:
                contexts.append(f"知识点「{clean_name}」核心内容：{summary[:300]}")
        kb_list = "\n".join(kb_lines)
        context_text = "\n\n".join(contexts) if contexts else ""

        system_prompt = r"""你是考研数学辅导专家。请直接输出1道练习题，不要输出任何思考过程或内心独白。

⚠️ 题目必须紧扣知识点核心概念，不得偏题。
⚠️ 直接输出题目内容，不要输出"首先"、"我需要"等思考过程。

⚠️ 数学公式强制规则（必须遵守，否则无法显示）：
- 所有公式必须用 $...$ 包裹，例如 $f(x)$、$\int_{a}^{b}$、$\frac{a}{b}$
- 独立公式用 $$...$$，例如 $$\lim_{x \to 0} \frac{\sin x}{x} = 1$$
- 禁止使用 \\(\\) 或 \\[\\]
- 禁止在 $ 外面写 \\frac、\\int、\\lim、\\pi 等 LaTeX 命令

严格按以下格式输出（不要输出格式说明之外的任何内容）：
Q: 题目（用文字描述）
A) 选项A
B) 选项B
C) 选项C
D) 选项D
ANSWER: 正确选项字母
EXPLAIN: 解析过程
---"""

        # 注入激活的 Skill prompt
        skill_prompt = build_system_prompt_with_skills(st.session_state.get("active_skills", []))
        full_system = system_prompt + ("\n\n---\n\n" + skill_prompt if skill_prompt else "")

        user_prompt = f"为以下知识点出1道选择题：\n\n{kb_list}\n\n{context_text}"

        request_data = {
            "model": "mimo-v2.5",
            "messages": [
                {"role": "system", "content": full_system},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": 5000,
            "temperature": 0.3
        }

        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(request_data).encode('utf-8'),
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {API_KEY}'
            },
            method='POST'
        )

        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            msg = result['choices'][0]['message']
            # MiMo 思维链模型：content 和 reasoning_content 都可能包含答案
            c = msg.get('content')
            content = c if isinstance(c, str) else ''
            reasoning = msg.get('reasoning_content') or ''
            raw = content if content and 'Q:' in content else ''
            # 从 reasoning_content 中提取 Q:...--- 格式
            if not raw:
                q_match = re.search(r'(Q:.*?---)', reasoning, re.DOTALL)
                if q_match:
                    raw = q_match.group(1)
            # 兜底：从合并文本中提取
            if not raw:
                combined = content + '\n' + reasoning
                q_match = re.search(r'(Q:.*?---)', combined, re.DOTALL)
                if q_match:
                    raw = q_match.group(1)
            if not raw:
                raw = content or reasoning
            return {
                "success": True,
                "questions": raw,
                "knowledge_points": [kp['knowledge_id'] for kp in knowledge_points[:3]]
            }

    except Exception as e:
        print(f"生成题目失败: {e}")
        return generate_local_questions(knowledge_points)

def generate_local_questions(knowledge_points):
    if not knowledge_points:
        return {"error": "无复习知识点", "questions": ""}

    kid = knowledge_points[0].get("knowledge_id", "知识点")
    questions = f"""Q: 请回忆 {kid} 的定义和基本概念
A) 查看知识点
B) 看文档
C) 翻资料
D) 点展开
ANSWER: A
EXPLAIN: 在知识库中查看完整内容
---"""
    return {
        "success": True,
        "questions": questions,
        "knowledge_points": [kp['knowledge_id'] for kp in knowledge_points[:3]]
    }

# ==================== 学习资料 - 辅助函数 ====================

def _read_reference_docx_structure():
    """读取 data/reference/ 下的参考 docx 文档，提取段落层级结构作为格式参考"""
    structure_desc = []
    for ref_file in sorted(REFERENCE_DIR.glob("*.docx")):
        try:
            doc = Document(ref_file)
            lines = []
            for p in doc.paragraphs[:60]:  # 只取前60段看清结构
                style = p.style.name if p.style else "Normal"
                text = (p.text or "").strip()
                if text:
                    lines.append(f"[{style}] {text[:120]}")
            if lines:
                structure_desc.append(f"### 《{ref_file.stem}》结构示例：\n" + "\n".join(lines))
        except Exception:
            pass
    return "\n\n".join(structure_desc) if structure_desc else "（暂无参考文档）"


def _build_material_prompt(selected_topics, user_requirement):
    """根据用户选择的知识点和需求，构建发给 AI 的 prompt"""
    # 读取 corpus 内容
    corpus_parts = []
    corpus_files = sorted(DATA_DIR.glob("*.md"))
    for fp in corpus_files:
        if selected_topics and fp.stem not in selected_topics:
            continue
        try:
            content = fp.read_text(encoding="utf-8")[:1500]  # 每个知识点最多取1500字
            corpus_parts.append(f"### {fp.stem}\n{content}")
        except Exception:
            pass

    corpus_text = "\n\n".join(corpus_parts) if corpus_parts else "（使用全部知识点）"

    # 读取参考 docx 格式
    ref_structure = _read_reference_docx_structure()

    # 读取 LaTeX 格式规范 skill
    latex_skill_path = Path("skills/latex-formatter/SKILL.md")
    latex_rules = ""
    if latex_skill_path.exists():
        latex_rules = latex_skill_path.read_text(encoding="utf-8")
    else:
        # 兜底：内嵌最基础的 LaTeX 规则
        latex_rules = """## LaTeX 格式强制规则
- 行内公式只用 $...$，禁止 \\(...\\)
- 独立公式只用 $$...$$，禁止 \\[...\\]
- \\\\ 只能出现在 $$...$$ 内
- 禁止在 $...$ 外用 \\frac、\\lim 等 LaTeX 命令"""

    prompt = f"""你是考研数学辅导专家。请根据提供的知识点内容，仿照参考文档的格式，生成一份考研数学学习/习题资料。

## 格式要求（参考 data/reference/ 下的文档结构）

参考文档采用以下层级：
{ref_structure}

请你用 Markdown 格式输出，层级规则：
- # 一级标题：章标题（如 # 第一章 凑元换元法）
- ## 二级标题：节标题（如 ## 含参变量积分）
- ### 三级标题：题号/子标题（如 ### 第 1 题）
- 每个题目或知识点包含：题目/概念 → 分析/提示 → 解答/推导 → 方法总结

{latex_rules}

## 知识点参考内容

{corpus_text}

## 用户需求

{user_requirement}

请直接输出生成的内容，无需额外说明。"""

    return prompt


def _build_english_material_prompt(category, user_requirement, vocab_category=""):
    """根据用户选择的分类和需求，构建发给 AI 的英语资料 prompt"""
    if category == "单词归类总结":
        prompt = f"""你是考研英语词汇专家。请根据用户选择的分类方式，生成一份考研英语词汇归类总结。

## 输出要求

1. **结构清晰**：按分类维度组织词汇，每个类别用标题分隔
2. **信息完整**：每个单词包含词性、中文释义、英文例句、记忆技巧
3. **实用性强**：提供真题出处和常见搭配
4. **格式规范**：
   - 一级标题：分类维度（如"经济类词汇"）
   - 二级标题：子分类（如"宏观经济学"）
   - 每个单词：词性 + 释义 + 例句 + 记忆技巧 + 真题出处

## 分类维度

{vocab_category}

## 用户需求

{user_requirement}

请直接输出生成的内容，无需额外说明。"""
    else:
        category_prompts = {
            "语法专题": "你是考研英语语法专家。请生成一份语法专题学习资料，包含：语法规则、例句、真题、易错点、记忆技巧。",
            "阅读技巧": "你是考研英语阅读专家。请生成一份阅读技巧学习资料，包含：题型特征、解题步骤、真题演示、注意事项。",
            "写作模板": "你是考研英语写作专家。请生成一份写作模板学习资料，包含：模板结构、高分句型、真题范文、评分标准。",
            "翻译技巧": "你是考研英语翻译专家。请生成一份翻译技巧学习资料，包含：翻译原则、技巧、真题演示、常见错误。",
            "完形填空": "你是考研英语完形填空专家。请生成一份完形填空学习资料，包含：逻辑关系、固定搭配、真题演示、解题策略。",
            "新题型": "你是考研英语新题型专家。请生成一份新题型学习资料，包含：题型特征、解题步骤、真题演示、注意事项。",
        }
        system_prompt = category_prompts.get(category, "你是考研英语辅导专家。")
        prompt = f"""{system_prompt}

## 输出要求

1. **结构清晰**：使用 Markdown 标题层级组织内容
2. **知识点完整**：每个知识点包含定义、规则、例句、真题
3. **例句真实**：使用考研真题例句或符合考研难度的例句
4. **实用性强**：提供解题技巧和常见错误提示
5. **格式规范**：
   - 一级标题：章节名称
   - 二级标题：知识点名称
   - 三级标题：子知识点
   - 每个知识点包含：定义 → 规则 → 例句 → 真题 → 技巧

## 用户需求

{user_requirement}

请直接输出生成的内容，无需额外说明。"""
    return prompt


def _generate_material(prompt):
    """调用 AI 生成资料内容，返回 (思考过程, 最终结果)"""
    data = {
        "model": "mimo-v2.5",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 4096,
        "temperature": 0.3,
    }
    try:
        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(data).encode("utf-8"),
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=300) as resp:
            msg = json.loads(resp.read().decode("utf-8"))["choices"][0]["message"]
        reasoning = msg.get("reasoning_content") or ""
        content = _extract_content(msg)
        # 如果模型不区分思考/结果，则全部作为结果
        if not content and reasoning:
            content = reasoning
            reasoning = ""
        return reasoning, content
    except Exception as e:
        raise RuntimeError(f"AI 调用失败: {e}")


def _ai_output_to_docx_via_pandoc(markdown_text):
    """用 Pandoc 将 Markdown+LaTeX 转为 DOCX（LaTeX 完美渲染）"""
    import tempfile
    import subprocess

    def _try_subprocess_pandoc(md_path):
        """尝试用 shell pandoc 转换"""
        template = str(Path(__file__).parent / "data" / "reference" / "template.docx")
        docx_path = md_path.replace(".md", ".docx")
        cmd = ["pandoc", md_path, "-o", docx_path, "--mathml", "--from", "markdown", "--to", "docx"]
        if os.path.exists(template):
            cmd += ["--reference-doc", template]
        subprocess.run(cmd, check=True, capture_output=True)
        with open(docx_path, "rb") as f:
            result = f.read()
        # cleanup
        try: os.unlink(docx_path)
        except Exception: pass
        return result

    def _try_pypandoc(md_path):
        """用 pypandoc_binary 内置的 pandoc 转换"""
        import pypandoc
        docx_path = md_path.replace(".md", ".docx")
        template = str(Path(__file__).parent / "data" / "reference" / "template.docx")
        extra_args = ["--mathml", "--from", "markdown", "--to", "docx"]
        if os.path.exists(template):
            extra_args += ["--reference-doc", template]
        pypandoc.convert_file(
            source_file=md_path, to="docx", format="markdown",
            outputfile=docx_path, extra_args=extra_args
        )
        with open(docx_path, "rb") as f:
            result = f.read()
        try: os.unlink(docx_path)
        except Exception: pass
        return result

    with tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w", encoding="utf-8") as md:
        clean = markdown_text
        # 去掉所有 Markdown 加粗/粗斜体标记，保留标题层级：
        # 1. ***粗斜体*** → 内容
        clean = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', clean, flags=re.DOTALL)
        # 2. **加粗** → 内容（跨行匹配）
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean, flags=re.DOTALL)
        # 3. __加粗__（下划线语法也是粗体）→ 内容
        clean = re.sub(r'__(.+?)__', r'\1', clean, flags=re.DOTALL)
        md.write(clean)
        md_path = md.name

    try:
        # 方案1: 优先用 pypandoc_binary（自带 pandoc，不依赖 PATH）
        try:
            return _try_pypandoc(md_path)
        except Exception:
            pass
        # 方案2: 尝试 shell pandoc
        try:
            return _try_subprocess_pandoc(md_path)
        except (FileNotFoundError, subprocess.SubprocessError):
            pass
        # 方案3: python-docx 兜底（去掉 LaTeX 标记让公式变纯文本）
        from docx import Document as DocxDoc
        import re as _re
        doc = DocxDoc()
        # 预处理：把 LaTeX 转成可读的纯文本
        clean_text = markdown_text
        clean_text = _re.sub(r'\$\$([^$]+)\$\$', r'\1', clean_text)  # $$...$$ → 内容
        clean_text = _re.sub(r'\$([^$]+)\$', r'\1', clean_text)      # $...$ → 内容
        clean_text = _re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', clean_text)  # \frac{a}{b} → (a)/(b)
        clean_text = _re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', clean_text)  # \sqrt{x} → √(x)
        clean_text = _re.sub(r'\\int_\{([^}]+)\}\^\{([^}]+)\}', r'∫[\1,\2]', clean_text)  # 积分
        clean_text = _re.sub(r'\\sum_\{([^}]+)\}\^\{([^}]+)\}', r'∑[\1,\2]', clean_text)  # 求和
        clean_text = _re.sub(r'\\lim_\{([^}]+)\}', r'lim[\1]', clean_text)
        clean_text = _re.sub(r'\\alpha', 'α', clean_text)
        clean_text = _re.sub(r'\\beta', 'β', clean_text)
        clean_text = _re.sub(r'\\theta', 'θ', clean_text)
        clean_text = _re.sub(r'\\pi', 'π', clean_text)
        clean_text = _re.sub(r'\\infty', '∞', clean_text)
        clean_text = _re.sub(r'\\to', '→', clean_text)
        clean_text = _re.sub(r'\\rightarrow', '→', clean_text)
        clean_text = _re.sub(r'\\times', '×', clean_text)
        clean_text = _re.sub(r'\\cdot', '·', clean_text)
        clean_text = _re.sub(r'\\pm', '±', clean_text)
        clean_text = _re.sub(r'\\leq', '≤', clean_text)
        clean_text = _re.sub(r'\\geq', '≥', clean_text)
        clean_text = _re.sub(r'\\neq', '≠', clean_text)
        clean_text = _re.sub(r'\\approx', '≈', clean_text)
        clean_text = _re.sub(r'\\Delta', 'Δ', clean_text)
        clean_text = _re.sub(r'\\delta', 'δ', clean_text)
        clean_text = _re.sub(r'\\lambda', 'λ', clean_text)
        clean_text = _re.sub(r'\\mu', 'μ', clean_text)
        clean_text = _re.sub(r'\\sigma', 'σ', clean_text)
        clean_text = _re.sub(r'\\varphi', 'φ', clean_text)
        clean_text = _re.sub(r'\\partial', '∂', clean_text)
        clean_text = _re.sub(r'\\nabla', '∇', clean_text)
        clean_text = _re.sub(r'\\varepsilon', 'ε', clean_text)
        clean_text = _re.sub(r'\\omega', 'Ω', clean_text)
        clean_text = _re.sub(r'\\begin\{[^}]*\}', '', clean_text)
        clean_text = _re.sub(r'\\end\{[^}]*\}', '', clean_text)
        clean_text = _re.sub(r'\\text\{([^}]*)\}', r'\1', clean_text)
        clean_text = _re.sub(r'\\mathbf\{([^}]*)\}', r'\1', clean_text)  # 粗体变普通
        clean_text = _re.sub(r'\\hat\{([^}]*)\}', r'\1̂', clean_text)
        clean_text = _re.sub(r'\\bar\{([^}]*)\}', r'\1̄', clean_text)
        clean_text = _re.sub(r'\\vec\{([^}]*)\}', r'\1⃗', clean_text)

        for line in clean_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("# "):
                p = doc.add_paragraph(stripped[2:].replace("***", "").replace("**", "").replace("__", ""))
                p.style = doc.styles["Heading 1"]
            elif stripped.startswith("## "):
                p = doc.add_paragraph(stripped[3:].replace("***", "").replace("**", "").replace("__", ""))
                p.style = doc.styles["Heading 2"]
            elif stripped.startswith("### "):
                p = doc.add_paragraph(stripped[4:].replace("***", "").replace("**", "").replace("__", ""))
                p.style = doc.styles["Heading 3"]
            elif stripped.startswith("#### "):
                p = doc.add_paragraph(stripped[5:].replace("***", "").replace("**", "").replace("__", ""))
                p.style = doc.styles["Heading 4"]
            else:
                doc.add_paragraph(stripped.replace("***", "").replace("**", "").replace("__", ""))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    finally:
        if os.path.exists(md_path):
            os.unlink(md_path)




# ==================== 多Agent管线 ====================

def extract_json(text):
    text = text.strip()
    if "```json" in text:
        text = text.split("```json",1)[1].split("```",1)[0].strip()
    elif text.startswith("```"):
        text = text.split("```",2)[1].strip()
    return text

ROUTER_PROMPT = """判断以下考研问题的学科类型，只输出JSON：
- english: 英语作文、翻译、阅读、完形、词汇、语法
- politics: 政治理论、马原、毛中特、近代史、思修、时政
- math: 数学计算、求导、积分、证明、公式、矩阵、概率

输出 {"type":"english"|"politics"|"math"}"""

ENGLISH_PROMPT = """你是考研英语辅导专家。专精：作文模板、长难句分析、翻译技巧、阅读策略。
回答简洁实用，给出可操作的建议。不编造具体分数线或统计数据。"""

POLITICS_PROMPT = """你是考研政治辅导专家。专精：马原原理、毛中特体系、近代史脉络、思修要点、时政热点。
回答结构清晰，先给出核心结论再展开。不编造具体分值或命题预测。"""

def classify_query(query):
    """Router: 判断问题属于 english/politics/math"""
    data = {
        "model": "mimo-v2.5",
        "messages": [
            {"role": "system", "content": ROUTER_PROMPT},
            {"role": "user", "content": query}
        ],
        "max_tokens": 30, "temperature": 0.3
    }
    req = urllib.request.Request(API_BASE + "/chat/completions",
        data=json.dumps(data).encode('utf-8'),
        headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}'},
        method='POST')
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            raw = json.loads(resp.read().decode('utf-8'))['choices'][0]['message']['content']
            return json.loads(extract_json(raw)).get("type", "math")
    except:
        return "math"

def parse_multi_output(raw_text):
    """解析 LLM 一次输出的 [ANSWER]/[KNOWLEDGE]/[QUIZ]"""
    if "[ANSWER]" not in raw_text:
        cleaned = raw_text.replace("\\(", "$").replace("\\)", "$").replace("\\[", "$$").replace("\\]", "$$")
        return {"answer": cleaned[:2000], "knowledge": [], "quiz": ""}
    def extract(begin, end):
        if begin in raw_text and end in raw_text:
            return raw_text.split(begin, 1)[1].split(end, 1)[0].strip()
        return ""
    knowledge_part = raw_text.split("[KNOWLEDGE]", 1)[-1] if "[KNOWLEDGE]" in raw_text else ""
    knowledge_raw = knowledge_part.split("[", 1)[0].strip() if "[" in knowledge_part else knowledge_part.strip()
    return {
        "answer": _fix_latex(extract("[ANSWER]", "[KNOWLEDGE]") or raw_text[:1500]),
        "knowledge": [k.strip() for k in knowledge_raw.split(",") if k.strip()],
    }

def run_pipeline(query, results, model_name, img_data=None):
    """统一管线: 流式调用 LLM，逐 token 返回"""
    pipeline_log = []
    
    skill_prompt = build_system_prompt_with_skills(st.session_state.get("active_skills", []))
    context = "\n\n".join([f"【{d['id']}】\n{d['text'][:800]}" for d in results[:3]]) if results else ""

    math_rules = r"""- 所有公式必须用 $...$ 包裹，例如 $f(x)$、$\int_{a}^{b}$、$\frac{a}{b}$
- 独立公式用 $$...$$，例如 $$\lim_{x \to 0} \frac{\sin x}{x} = 1$$
- 禁止使用 \( \) 或 \[ \]
- 禁止在 $ 外面写 \frac、\int、\lim、\pi 等 LaTeX 命令"""

    system_prompt = f"""你是考研数学辅导专家。请完成以下任务并用标签输出：

任务1：根据参考资料回答用户问题。{"严格遵循 Skill 的格式要求。" if skill_prompt else ""}

任务2：判断问题涉及的知识点，输出概念名称（如：导数, 定积分, 矩阵）。

⚠️ 数学公式强制规则（必须遵守，否则无法显示）：
{math_rules}

输出格式：
[ANSWER]
（回答）

[KNOWLEDGE]
（概念名，逗号分隔）

{skill_prompt if skill_prompt else ""}

参考资料：
{context}"""

    if img_data:
        user_content = [
            {"type": "text", "text": f"问题：{query}"},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_data}"}}
        ]
    else:
        user_content = f"问题：{query}"
    model = model_name
    max_tok = 800 if img_data else 1500
    temp = 0.3
    data = {
        "model": model,
        "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_content}],
        "max_tokens": max_tok,
        "temperature": temp,
        "stream": True,
    }
    # 先尝试流式
    try:
        req = urllib.request.Request(API_BASE + "/chat/completions", data=json.dumps(data).encode('utf-8'), headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}'}, method='POST')
        raw_full = ""
        with urllib.request.urlopen(req, timeout=180) as resp:
            buffer = ""
            while True:
                chunk = resp.read(1024)
                if not chunk:
                    break
                buffer += chunk.decode("utf-8", errors="ignore")
                while "\n" in buffer:
                    line, buffer = buffer.split("\n", 1)
                    line = line.strip()
                    if not line.startswith("data: "):
                        continue
                    payload = line[6:]
                    if payload == "[DONE]":
                        break
                    try:
                        obj = json.loads(payload)
                        delta_obj = obj.get("choices", [{}])[0].get("delta", {})
                        # MiMo 是思维链模型，内容在 reasoning_content 中
                        c = delta_obj.get("content")
                        delta = c if isinstance(c, str) else ""
                        if not delta:
                            delta = delta_obj.get("reasoning_content") or ""
                        if delta:
                            raw_full += delta
                            yield {"type": "token", "content": delta}
                    except json.JSONDecodeError:
                        pass
        result = parse_multi_output(raw_full)
        result["_raw_debug"] = raw_full[:500]
        result["qtype"] = "math"
        result["pipeline_log"] = pipeline_log
        yield {"type": "done", "result": result}
    except Exception:
        # 流式失败，降级为非流式
        try:
            data["stream"] = False
            req = urllib.request.Request(API_BASE + "/chat/completions", data=json.dumps(data).encode('utf-8'), headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}'}, method='POST')
            with urllib.request.urlopen(req, timeout=180) as resp:
                msg = json.loads(resp.read().decode('utf-8'))['choices'][0]['message']
                c = msg.get('content')
                raw_full = c if isinstance(c, str) else ''
                if not raw_full:
                    raw_full = msg.get('reasoning_content') or ''
                yield {"type": "token", "content": raw_full}
            result = parse_multi_output(raw_full)
            result["_raw_debug"] = raw_full[:500]
            result["qtype"] = "math"
            result["pipeline_log"] = pipeline_log
            yield {"type": "done", "result": result}
        except Exception as e:
            yield {"type": "done", "result": {"answer": f"[系统提示] API调用失败: {str(e)[:100]}", "knowledge": [], "quiz": "", "qtype": "math", "pipeline_log": pipeline_log}}

# ==================== LLM调用 ====================

def call_llm(query, context_docs, model_name=None):
    """调用LLM API - 支持RAG和纯LLM两种模式"""
    if model_name is None:
        model_name = MODEL_NAME

    try:
        experience = load_agent_experience()

        # 模式判断：有检索结果用RAG，无检索结果用纯LLM
        has_context = context_docs and len(context_docs) > 0

        # 加载动态经验库
        experience = load_agent_experience()

        if has_context:
            # RAG模式：结合知识库
            context = "\n\n".join([f"【{d['id']}】\n{d['text'][:800]}" for d in context_docs[:3]])

            # 不可变约束 + 动态经验库
            static_rules = """## 铁律：不可变约束 (绝对不可修改)
1. **信息溯源**：回答必须严格基于提供的参考资料。资料中信息不足时，请如实说明。
2. **禁止编造数据**：不编造具体数字、百分比、机构名、人名，除非资料中明确出现。
3. **禁止无关延伸**：不补充资料未提及的内容。

## 动态经验与偏好库 (自学习记录)
"""
            system_prompt = static_rules + (experience if experience else "暂无追加规则")
            system_prompt += "\n\n请直接回答，不要多余的开场或结尾闲聊。"

            user_prompt = f"""【用户问题】
{query}

【参考资料】
{context}

请根据以上参考资料回答："""
        else:
            # 纯LLM模式
            static_rules = """## 铁律：不可变约束
1. 如果无法确定答案，请诚实说明。
2. 不编造具体数字、研究来源、统计报告。
3. 回答简洁、有据可查。

## 动态经验与偏好库
"""
            system_prompt = static_rules + (experience if experience else "暂无追加规则")
            system_prompt += "\n\n请直接回答，不要多余闲聊。"

            user_prompt = f"""【用户问题】
{query}

请回答："""

        # 注入激活的 Skill
        skill_prompt = build_system_prompt_with_skills(st.session_state.get("active_skills", []))
        if skill_prompt:
            system_prompt = skill_prompt + "\n\n---\n\n" + system_prompt

        # 不同的max_tokens
        max_tokens = 800 if has_context else 1200

        request_data = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": max_tokens,
            "temperature": 0.3
        }

        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(request_data).encode('utf-8'),
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {API_KEY}'
            },
            method='POST'
        )

        with urllib.request.urlopen(req, timeout=30) as response:
            result = json.loads(response.read().decode('utf-8'))
            return result['choices'][0]['message']['content']

    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else ""
        st.error(f"API错误 {e.code}: {error_body}")
        return fallback_answer(query, context_docs)
    except Exception as e:
        st.error(f"API调用失败: {e}")
        return fallback_answer(query, context_docs)

def fallback_answer(query, docs):
    if not docs:
        return "未找到相关资料"
    best = docs[0]
    text = best["text"]
    return f"""根据检索到的资料回答【{query}】：

{text[:600]}...

---
参考来源：{best['id']} (相关性: {best['score']})"""

# ==================== 幻觉检测 ====================

MATH_EVAL_PROMPT = """你是考研数学事实核查员。评估回答是否在上下文中存在有害幻觉。

## 三类声明
1. **严格支持**: 回答直接来源于Context
2. **专业常识拓展**: Context未提及，但属于大学数学公认定理/定义（如子数列收敛性、零点定理、极限四则运算）- 宽容通过
3. **有害幻觉**: 捏造考情/分值/频率/历史/应用领域

## 输出JSON
{"is_hallucinating": true/false, "hallucinated_claims": [...], "common_sense_claims": [...]}"""

def evaluate_hallucination(user_query: str, context: str, agent_response: str, model_name=None):
    """调用LLM评估回答是否存在有害幻觉"""
    if model_name is None:
        model_name = MODEL_NAME
    try:
        prompt = f"""[User Query]: {user_query}

[Context]:
{context}

[Agent Response]:
{agent_response}"""

        data = {
            "model": model_name,
            "messages": [
                {"role": "system", "content": MATH_EVAL_PROMPT},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 600,
            "temperature": 0.1
        }
        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(data).encode('utf-8'),
            headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}'},
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            content = json.loads(resp.read().decode('utf-8'))['choices'][0]['message']['content']
            return json.loads(content)
    except Exception as e:
        return {"is_hallucinating": False, "error": str(e), "hallucinated_claims": [], "common_sense_claims": []}

# ==================== Agent自我反思 ====================

def trigger_self_learning(rule_text: str) -> str:
    """将新规则追加到动态经验库，并返回确认信息"""
    existing = load_agent_experience()
    # 找到最后一条编号
    lines = existing.split("\n")
    last_num = 1
    for line in lines:
        if line.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            try:
                num = int(line.split(".")[0])
                last_num = max(last_num, num)
            except:
                pass

    next_num = last_num + 1
    today = datetime.now().strftime("%Y-%m-%d")
    new_entry = f"{next_num}. [{today}] {rule_text}"

    # 追加到"结束记录"之前
    if "--- 结束记录 ---" in existing:
        updated = existing.replace("--- 结束记录 ---", f"{new_entry}\n--- 结束记录 ---")
    else:
        updated = f"{existing}\n{new_entry}\n--- 结束记录 ---"

    save_agent_experience(updated)
    return f"\n**自学习已触发** — **已将以下规则追加至经验库**：{rule_text}\n**当前状态**：底层逻辑未受影响，增量规则已生效。"

def agent_reflect(question, answer, feedback):
    try:
        prompt = f"""你是一个Agent，正在进行自我反思。用户对你的回答提供了反馈：

问题: {question}
你的回答: {answer}
用户反馈: {feedback}

请从反馈中提炼出1条可以在后续任务中复用的具体规则（一句话即可，不要编号）。

直接输出规则文字，不要多余内容。"""

        request_data = {
            "model": MODEL_NAME,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 200,
            "temperature": 0.3
        }

        req = urllib.request.Request(
            API_BASE + "/chat/completions",
            data=json.dumps(request_data).encode('utf-8'),
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {API_KEY}'
            },
            method='POST'
        )

        with urllib.request.urlopen(req, timeout=20) as response:
            result = json.loads(response.read().decode('utf-8'))
            rule = result['choices'][0]['message']['content'].strip()

            # 追加到动态经验库
            confirm = trigger_self_learning(rule)
            return {"success": True, "reflection": rule, "confirm": confirm}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ==================== 会话状态 ====================
if "thinking_log" not in st.session_state:
    st.session_state.thinking_log = []
if "current_knowledge_ids" not in st.session_state:
    st.session_state.current_knowledge_ids = []
if "selected_model" not in st.session_state:
    st.session_state.selected_model = "gpt-4o"

def add_thinking(msg):
    """添加思考日志"""
    timestamp = datetime.now().strftime("%H:%M:%S")
    st.session_state.thinking_log.append(f"[{timestamp}] {msg}")

# ==================== Skill 技能系统 ====================

SKILLS_DIR = Path("skills")

def load_all_skills():
    """自动扫描 skills/ 目录，加载所有 SKILL.md"""
    skills = {}
    if not SKILLS_DIR.exists():
        return skills
    for skill_dir in sorted(SKILLS_DIR.iterdir()):
        if not skill_dir.is_dir() or skill_dir.name.startswith("_"):
            continue
        skill_file = skill_dir / "SKILL.md"
        if not skill_file.exists():
            continue
        try:
            content = skill_file.read_text(encoding="utf-8")
            meta, body = parse_skill_frontmatter(content)
            if meta.get("hidden"):
                continue
            meta["_dir"] = str(skill_dir)
            meta["_body"] = body.strip()
            skills[meta.get("name", skill_dir.name)] = meta
        except:
            pass
    return skills

def parse_skill_frontmatter(content):
    """解析 YAML frontmatter，返回 (meta_dict, body)"""
    lines = content.strip().split("\n")
    meta = {}
    body_start = 0
    if lines and lines[0].strip() == "---":
        i = 1
        while i < len(lines):
            line = lines[i]
            if line.strip() == "---":
                body_start = i + 1
                break
            if ":" in line:
                key, _, val = line.partition(":")
                key = key.strip()
                val = val.strip().strip('"').strip("'")
                if val.startswith("[") and val.endswith("]"):
                    val = [x.strip().strip('"') for x in val[1:-1].split(",")]
                meta[key] = val
            i += 1
    body = "\n".join(lines[body_start:]).strip() if body_start > 0 else content.strip()
    return meta, body

def build_system_prompt_with_skills(active_skills):
    """将激活的 Skill prompts 注入 system_prompt"""
    skill_prompts = []
    for name in active_skills:
        skills = load_all_skills()
        if name in skills:
            body = skills[name].get("_body", "")
            if body:
                skill_prompts.append(f"## Skill: {skills[name].get('description', name)}\n\n{body}")
    return "\n\n---\n\n".join(skill_prompts) if skill_prompts else ""

# ==================== 智能知识点匹配 ====================

def _to_bigrams(text):
    """中文 2-gram 切分"""
    text = text.strip()
    if len(text) < 2:
        return {text}
    return {text[i:i+2] for i in range(len(text) - 1)}

def _jaccard(set_a, set_b):
    if not set_a or not set_b:
        return 0.0
    return len(set_a & set_b) / len(set_a | set_b)

def _build_knowledge_index(corpus):
    """构建知识点倒排索引（启动时执行一次）"""
    idx = {"doc_names": [], "title_map": {}, "title_kw": {}, "content_terms": {}}
    for doc in corpus:
        fname = doc["id"]
        text = doc["text"]
        idx["doc_names"].append(fname)
        # 提取标题
        title_line = ""
        for line in text.split("\n"):
            s = line.strip()
            if s.startswith("# "):
                title_line = s.lstrip("# ").strip()
                break
        idx["title_map"][fname] = title_line or fname
        # 标题关键词：字符级 bigram
        idx["title_kw"][fname] = _to_bigrams(title_line or fname)
        # 全文 bigram 词频（用于 TF 重叠）
        idx["content_terms"][fname] = set(_to_bigrams(text[:3000]))
    return idx

def match_knowledge_v2(concepts, index):
    """用索引匹配 LLM 提取的概念 → 文件名列表"""
    if not concepts or not index:
        return []
    _NOISE = {"函数", "公式", "定理", "法则", "方法", "计算", "概念", "性质", "定义", "应用", "意义"}
    results = []
    for concept_raw in concepts:
        concept_raw = concept_raw.strip()
        if not concept_raw:
            continue
        # 拆分停用词：尝试多粒度匹配
        variants = [concept_raw]
        for noise in _NOISE:
            if noise in concept_raw and len(concept_raw.replace(noise, "")) >= 2:
                variants.append(concept_raw.replace(noise, ""))
        scores = {}
        for fname in index["doc_names"]:
            title = index["title_map"].get(fname, "")
            best_var_score = 0.0
            for variant in variants[:2]:
                vs = 0.0
                v_bigrams = _to_bigrams(variant)
                pos = title.find(variant)
                if pos >= 0:
                    # 位置加权：越靠近标题开头分越高
                    vs += 0.5 * max(0.1, 1 - pos / max(len(title), 1))
                else:
                    vs += _jaccard(v_bigrams, index["title_kw"].get(fname, set())) * 0.3
                vs += _jaccard(v_bigrams, index["content_terms"].get(fname, set())) * 0.2
                if vs > best_var_score:
                    best_var_score = vs
            if best_var_score > 0:
                scores[fname] = best_var_score
        best = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:2]
        for fname, _ in best:
            results.append(fname)
    return list(dict.fromkeys(results))

def smart_match_knowledge(query):
    """LLM 提取概念 → 向量/关键词双重匹配"""
    # ① LLM 提取概念
    try:
        data = {
            "model": "mimo-v2.5",
            "messages": [
                {"role": "system", "content": "从以下考研数学问题中提取1-3个核心知识点名称（每行一个，不要编号）。"},
                {"role": "user", "content": query}
            ],
            "max_tokens": 500, "temperature": 0.3
        }
        req = urllib.request.Request(API_BASE + "/chat/completions",
            data=json.dumps(data).encode('utf-8'),
            headers={'Content-Type': 'application/json', 'Authorization': f'Bearer {API_KEY}'},
            method='POST')
        with urllib.request.urlopen(req, timeout=20) as resp:
            concepts = json.loads(resp.read().decode('utf-8'))['choices'][0]['message']['content']
            concepts = [c.strip().strip("-•*") for c in concepts.split("\n") if c.strip()]
    except:
        return []
    # ② 用预建索引匹配
    idx = st.session_state.get("_knowledge_index")
    if not idx:
        idx = _build_knowledge_index(load_corpus())
        st.session_state["_knowledge_index"] = idx
    return match_knowledge_v2(concepts, idx)

# ==================== UI界面 ====================

# 登录状态
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "page" not in st.session_state:
    st.session_state.page = "hub"

# 确保数据库表存在（登录前就必须建好）
init_memory_db()

# 自动登录（CookieManager 方案）
if not st.session_state.logged_in:
    token = cookie_manager.get("auth_token")
    if token:
        user_info = verify_login_token(token)
        if user_info:
            st.session_state.logged_in = True
            st.session_state.user_id = user_info["user_id"]
            st.session_state.username = user_info["username"]
            st.rerun()

if not st.session_state.logged_in:
    # ─── 登录/注册页 ───
    if not API_KEY:
        st.warning("⚠️ 未设置 API Key。请设置环境变量 `AI_API_KEY` 后重启。")
        st.code("export AI_API_KEY='sk-xxx'  # Linux/Mac\nset AI_API_KEY=sk-xxx  # Windows", language="bash")
        st.stop()
    st.markdown("""
    <div class="main-title">
        <h1>考研学习助手</h1>
        <p>多用户知识问答系统</p>
    </div>
    """, unsafe_allow_html=True)

    tab_login, tab_register = st.tabs(["登录", "注册"])
    
    with tab_login:
        with st.form("login_form"):
            username = st.text_input("用户名")
            password = st.text_input("密码", type="password")
            submitted = st.form_submit_button("登录", use_container_width=True, type="primary")
            if submitted and username and password:
                uid = login_user(username, password)
                if uid:
                    token = generate_login_token()
                    save_login_token(uid, token)
                    cookie_manager.set("auth_token", token, expires_at=datetime.now() + timedelta(days=30))
                    st.session_state.logged_in = True
                    st.session_state.user_id = uid
                    st.session_state.username = username
                    st.success("登录成功！")
                    st.rerun()
                else:
                    st.error("用户名或密码错误")

    with tab_register:
        with st.form("register_form"):
            new_user = st.text_input("新用户名")
            new_pass = st.text_input("新密码", type="password")
            new_pass2 = st.text_input("确认密码", type="password")
            reg_submitted = st.form_submit_button("注册", use_container_width=True)
            if reg_submitted and new_user and new_pass:
                if new_pass != new_pass2:
                    st.error("两次密码不一致")
                elif len(new_pass) < 3:
                    st.error("密码至少3位")
                else:
                    uid = register_user(new_user, new_pass)
                    if uid:
                        token = generate_login_token()
                        save_login_token(uid, token)
                        cookie_manager.set("auth_token", token, expires_at=datetime.now() + timedelta(days=30))
                        st.session_state.logged_in = True
                        st.session_state.user_id = uid
                        st.session_state.username = new_user
                        st.success(f"注册成功！欢迎 {new_user}")
                        st.rerun()
                    else:
                        st.error("用户名已存在")

    st.stop()

# ==================== 全局侧边栏导航 ====================
_username = st.session_state.get('username', '?')
with st.sidebar:
    # 品牌区 — 渐变标题 (SVG icon)
    st.markdown("""
    <div class="sidebar-brand">
        <span class="sidebar-brand-icon"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/></svg></span>
        <span class="sidebar-brand-text">考研学习助手</span>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

    # 导航分组 1: 核心功能
    st.markdown('<div class="sidebar-section-label">核心功能</div>', unsafe_allow_html=True)
    _group1 = [
        ("hub",    "备考看板"),
        ("main",   "数学问答"),
        ("english","英语专家"),
        ("checkin","打卡督学"),
    ]
    current_page = st.session_state.get("page", "hub")
    for p, label in _group1:
        if current_page == p:
            st.markdown(f'<div class="nav-item nav-item-active" data-nav="{p}"><span class="nav-dot"></span>{label}</div>', unsafe_allow_html=True)
        else:
            if st.button(label, key=f"nav_{p}", use_container_width=True):
                st.session_state.page = p
                st.rerun()

    # 导航分组 2: 辅助工具
    st.markdown('<div class="sidebar-section-label">辅助工具</div>', unsafe_allow_html=True)
    _group2 = [
        ("popularity","高校热度"),
        ("material", "学习资料"),
        ("suggest",  "提建议"),
    ]
    for p, label in _group2:
        if current_page == p:
            st.markdown(f'<div class="nav-item nav-item-active" data-nav="{p}"><span class="nav-dot"></span>{label}</div>', unsafe_allow_html=True)
        else:
            if st.button(label, key=f"nav_{p}", use_container_width=True):
                st.session_state.page = p
                st.rerun()

    st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

    # 用户信息卡片 (升级版)
    _uid = st.session_state.get("user_id")
    _study_days = 5  # 默认值
    if _uid:
        try:
            conn = sqlite3.connect(MEMORY_DB)
            row = conn.execute("SELECT COUNT(DISTINCT date(checkin_time)) FROM checkin_records WHERE user_id=?", (_uid,)).fetchone()
            if row and row[0]: _study_days = row[0]
            conn.close()
        except Exception:
            pass
    st.markdown(f"""
    <div class="sidebar-user">
        <div class="sidebar-avatar">{_username[0]}</div>
        <div class="sidebar-user-info">
            <span class="sidebar-username">{_username}</span>
            <span class="sidebar-subtitle">2026 届硕士</span>
        </div>
    </div>
    <div class="sidebar-stats">
        <span class="sidebar-stat">· <strong>{_study_days}</strong> 天学习</span>
        <span class="sidebar-stat">· <strong>{_study_days}</strong> 连击</span>
    </div>
    """, unsafe_allow_html=True)

    if st.button("退出登录", key="sidebar_logout", use_container_width=True):
        clear_login_token(st.session_state.get("user_id", 0))
        cookie_manager.delete("auth_token")
        st.session_state.logged_in = False
        st.session_state.user_id = None
        st.session_state.page = "hub"
        st.rerun()

    # ===== 调试模式 =====
    st.markdown("---")
    debug_on = st.checkbox("🔍 调试模式", value=st.session_state.get("debug_mode", False), key="debug_toggle",
                           help="开启后记录每次 API 请求的完整链路")
    st.session_state.debug_mode = debug_on
    if debug_on:
        logs = st.session_state.get("debug_logs", [])
        st.caption(f"📋 已记录 {len(logs)} 条日志")
        if logs:
            with st.expander("查看最近日志", expanded=False):
                for log in reversed(logs[-10:]):
                    idx = logs.index(log) + 1
                    status_icon = "✅" if log.get("status") == "ok" else "❌"
                    elapsed = log.get("elapsed_ms", "?")
                    st.markdown(
                        f"**{status_icon} #{idx}** `{log.get('time','?')}` | "
                        f"{elapsed}ms | `{log.get('model','?')}`"
                    )
                    with st.expander(f"#{idx} 详情"):
                        st.caption(f"**Prompt 预览** ({log.get('prompt_len',0)} 字符)")
                        st.code(log.get("prompt_preview", "")[:300], language=None)
                        if log.get("status") == "ok":
                            st.caption(f"**原始输出** ({log.get('raw_full_len',0)} 字符)")
                            st.code(log.get("raw_preview", "")[:400], language=None)
                            cs = log.get("clean_stats", {})
                            if cs:
                                st.caption(
                                    f"**清洗统计**: 总{cs.get('total_lines','?')}行 → "
                                    f"过滤{cs.get('filtered_count',0)}行 → "
                                    f"结果{cs.get('result_count',0)}行"
                                    f"{' ⚠️截断' if cs.get('was_truncated') else ''}"
                                    f"{' ⚠️全过滤' if cs.get('all_filtered') else ''}"
                                )
                                if cs.get("filtered_reasons"):
                                    st.caption(f"过滤原因: {'; '.join(cs['filtered_reasons'][:3])}")
                            st.caption(f"**清洗后输出** ({log.get('cleaned_len',0)} 字符)")
                            st.code(log.get("cleaned_preview", "")[:400], language=None)
                        else:
                            st.error(f"**异常**: {log.get('error','')[:300]}")
                            st.code(log.get("traceback", "")[:400], language=None)
        if st.button("🗑️ 清空日志", key="clear_debug_logs"):
            st.session_state.debug_logs = []
            st.rerun()

# ==================== 数学问答 ====================
if st.session_state.page == "main":
    if st.button("← 返回首页", key="back_hub_math"):
        st.session_state.page = "hub"
        st.rerun()

    st.markdown("""
    <div class="main-title" style="text-align:left;padding:1.2rem 1.8rem;">
        <div style="display:flex;align-items:center;gap:14px;">
            <div style="width:42px;height:42px;border-radius:12px;background:linear-gradient(135deg,#4f46e5,#8b5cf6);display:flex;align-items:center;justify-content:center;color:#fff;flex-shrink:0;box-shadow:0 4px 12px rgba(79,70,229,0.3);"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" width="24" height="24"><rect x="3" y="3" width="7" height="7" rx="1"/><rect x="14" y="3" width="7" height="7" rx="1"/><rect x="3" y="14" width="7" height="7" rx="1"/><rect x="14" y="14" width="7" height="7" rx="1"/></svg></div>
            <div>
                <h1 style="margin:0;font-size:1.4rem!important;">考研数学助手</h1>
                <p style="margin:2px 0 0;font-size:0.82rem;opacity:0.75;">110 个知识点 · 智能检索 · AI 精准回答</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Layout: Left Panel + Main ──
    left, main = st.columns([0.22, 0.78], gap="medium")

    with left:
        # User card
        st.markdown(f"""
        <div class="qa-card" style="padding:14px;">
            <div style="display:flex;align-items:center;gap:10px;">
                <div style="width:36px;height:36px;border-radius:50%;background:linear-gradient(135deg,#4f46e5,#a855f7);color:#fff;display:flex;align-items:center;justify-content:center;font-weight:700;font-size:0.9rem;">{st.session_state.get('username','?')[0]}</div>
                <div>
                    <div style="font-weight:700;font-size:0.85rem;color:#1e293b;">{st.session_state.get('username','?')}</div>
                    <div style="font-size:0.7rem;color:#64748b;">2026 届硕士</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # QA Style Selector
        st.markdown('<p style="font-size:0.78rem;font-weight:700;color:#1e293b;margin:10px 0 6px 0;">回答模式选择</p>', unsafe_allow_html=True)
        _styles = ["默认", "分步解题", "概念讲解", "错题分析", "纯要点", "一问一答", "纯公式"]
        _style = st.session_state.get("qa_style", "默认")
        for s in _styles:
            if st.button(s, key=f"style_{s}", use_container_width=True,
                         type="secondary" if _style != s else "primary"):
                st.session_state.qa_style = s
                st.rerun()

    with main:
        # Guide — 蓝色渐变卡片
        st.markdown("""
        <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border:1px solid #bae6fd;border-radius:12px;padding:18px 20px;margin-bottom:16px;">
            <div style="font-size:0.92rem;font-weight:700;color:#0369a1;margin-bottom:10px;">从这里开始</div>
            <div style="font-size:0.84rem;color:#1e40af;line-height:1.8;">
                <p style="margin-bottom:8px;">这里收录了 <strong>110 个考研数学核心知识点</strong>，从高数到线代再到概率论，按 2025 年最新发布的考研数学大纲整理好了。问什么，它就去知识库里翻最相关的内容，再让 AI 组织成你能看懂的答案。</p>
                <p style="margin-bottom:8px;">左边可以选回答风格——想要一步一步推导就选「分步解题」，想搞清楚概念就选「概念讲解」，当然也可以什么都不选，它会自己判断。</p>
                <p style="margin-bottom:8px;">如果想要自己测试一下对知识的理解效果，可以试试网页下方的费曼学习法——挑一个你觉得已经掌握的知识点，试着用最简单的话写出来，提交后 AI 会帮你打分，看看哪里讲得清楚、哪里还需要补一补。分数和记录都会留下来，回头能翻。</p>
                <p>问完问题之后，AI 也会顺便帮你记着哪些知识点还不熟，自动存到「复习挑战」里，下次打开就能直接练，不用自己惦记。</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # QA Input
        st.markdown("### 智能回答")
        _upload = st.file_uploader("上传题目截图（可选）", type=["png","jpg","jpeg"], key="qa_upload")
        _q = st.text_input("输入你的考研问题", placeholder="例如：什么是导数？如何求泰勒展开？",
                          key="qa_input", label_visibility="collapsed")
        _ask = st.button("提问", use_container_width=True, type="primary")

        if _ask and (_q.strip() or _upload):
            with st.spinner("AI 正在思考..."):
                user_input = _q.strip()
                # 处理图片上传：用 MiMo 多模态识别题目文字
                if _upload and not user_input:
                    import base64 as _b64
                    img_bytes = _upload.read()
                    img_b64 = _b64.b64encode(img_bytes).decode()
                    ocr_prompt = "请识别这张考研数学题目图片中的文字和公式，直接输出题目原文，不要添加任何解释。如有公式请用 LaTeX 格式。"
                    ocr_data = {
                        "model": "mimo-v2.5",
                        "messages": [{"role": "user", "content": [
                            {"type": "text", "text": ocr_prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
                        ]}],
                        "max_tokens": 500, "temperature": 0.1
                    }
                    try:
                        ocr_req = urllib.request.Request(
                            API_BASE + "/chat/completions",
                            data=json.dumps(ocr_data).encode("utf-8"),
                            headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                            method="POST")
                        with urllib.request.urlopen(ocr_req, timeout=60) as ocr_resp:
                            ocr_msg = json.loads(ocr_resp.read().decode("utf-8"))["choices"][0]["message"]
                        user_input = _extract_content(ocr_msg) or "（图片识别失败，请手动输入问题）"
                        with st.expander("📷 识别结果"):
                            st.caption(user_input)
                    except Exception as e:
                        st.error(f"图片识别失败: {e}")
                        user_input = ""
                elif _upload and user_input:
                    # 有文本也有图片：附加图片内容到 prompt
                    import base64 as _b64
                    img_bytes = _upload.read()
                    img_b64 = _b64.b64encode(img_bytes).decode()
                    user_input += "\n[附题目截图]"
                    # 多模态调用：同时传文本和图片
                    corpus = load_corpus()
                    docs = search_corpus(user_input, corpus, top_k=3)
                    style = st.session_state.get("qa_style", "默认")
                    style_hint = "" if style == "默认" else f"请用{style}的方式回答。"
                    prompt = f"""{style_hint}你是考研数学辅导专家。用户上传了一道题目截图，请根据以下参考资料解答。

用户问题: {user_input}

参考资料:
"""
                    for i, doc in enumerate(docs):
                        prompt += f"\n[{i+1}] {doc['id']}: {doc['text'][:800]}\n"
                    prompt += "\n请给出准确、有深度的回答。如有公式请用 LaTeX 格式。"
                    mm_data = {
                        "model": "mimo-v2.5",
                        "messages": [{"role": "user", "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
                        ]}],
                        "max_tokens": 2000, "temperature": 0.3
                    }
                    mm_req = urllib.request.Request(
                        API_BASE + "/chat/completions",
                        data=json.dumps(mm_data).encode("utf-8"),
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                        method="POST")
                    with urllib.request.urlopen(mm_req, timeout=120) as mm_resp:
                        mm_msg = json.loads(mm_resp.read().decode("utf-8"))["choices"][0]["message"]
                    answer = _extract_content(mm_msg) or "（AI 回复为空，请重试）"
                    st.markdown("### 回答")
                    st.markdown(f'<div class="qa-card">{answer}</div>', unsafe_allow_html=True)
                    if docs:
                        ref_html = "".join(f'<span class="ref-tag">{d["id"]}</span>' for d in docs)
                        st.markdown(f'<div style="margin-top:8px;">{ref_html}</div>', unsafe_allow_html=True)
                    st.stop()

                if not user_input:
                    st.warning("请输入问题或上传题目截图")
                    st.stop()

                corpus = load_corpus()
                docs = search_corpus(user_input, corpus, top_k=3)
                style = st.session_state.get("qa_style", "默认")
                style_hint = "" if style == "默认" else f"请用{style}的方式回答。"
                prompt = f"""{style_hint}你是考研数学辅导专家。请根据以下参考资料回答用户问题。

用户问题: {user_input}

参考资料:
"""
                for i, doc in enumerate(docs):
                    prompt += f"\n[{i+1}] {doc['id']}: {doc['text'][:800]}\n"
                prompt += "\n请给出准确、有深度的回答。如有公式请用 LaTeX 格式。"
                answer = call_llm_api(prompt, model="mimo-v2.5", max_tokens=2000)
            st.markdown("### 回答")
            st.markdown(f'<div class="qa-card">{answer}</div>', unsafe_allow_html=True)
            if docs:
                ref_html = "".join(f'<span class="ref-tag">{d["id"]}</span>' for d in docs)
                st.markdown(f'<div style="margin-top:8px;">{ref_html}</div>', unsafe_allow_html=True)

        # ── Bottom Tabs ──
        st.markdown("---")
        _tabs = st.tabs(["知识库", "复习挑战", "费曼学习法", "记忆系统"])

        with _tabs[0]:
            corpus = load_corpus()
            _q_search = st.text_input("搜索知识库", placeholder="输入关键词筛选...", key="kb_search")
            _filter = st.radio("筛选", ["全部", "数学一专属", "数学三专属"], horizontal=True, key="kb_filter")
            _filtered = _filter_corpus(corpus, {"全部": "all", "数学一专属": "math1", "数学三专属": "math3"}.get(_filter, "all"))
            if _q_search:
                _filtered = [d for d in _filtered if _q_search.lower() in d.get("id","").lower() or _q_search.lower() in d.get("text","").lower()]

            # 简洁标题
            st.markdown(f"""
            <div style="display:flex;align-items:center;gap:8px;margin-bottom:14px;padding:8px 0;border-bottom:1px solid #e2e8f0;">
                <span style="font-size:0.88rem;font-weight:600;color:#1e293b;">共 <span style="color:#4f46e5;">{len(_filtered)}</span> 个知识点</span>
            </div>
            """, unsafe_allow_html=True)

            # 分页显示（每页20条）
            _total = len(_filtered)
            _per_page = 20
            _total_pages = max(1, (_total + _per_page - 1) // _per_page)
            _page = st.number_input("页码", min_value=1, max_value=_total_pages, value=1, key="kb_page")
            _start = (_page - 1) * _per_page
            _end = min(_start + _per_page, _total)

            # 知识点列表
            for doc in _filtered[_start:_end]:
                doc_id = doc.get("id", "未知")
                doc_text = doc.get("text", "无内容")

                # 简化文件名：去掉 .md 后缀
                clean_name = doc_id.replace('.md', '').replace('_', ' ')

                # 卡片：整个横条淡蓝色渐变
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border-radius:12px;padding:16px 18px;margin-bottom:10px;cursor:pointer;border:1px solid #bae6fd;">
                    <div style="font-size:0.92rem;font-weight:600;color:#0369a1;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">
                        {clean_name}
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # 按钮行（大尺寸）
                b1, b2, b3 = st.columns([1.5, 1, 1])
                with b1:
                    if st.button("展开查看", key=f"view_{doc_id}", use_container_width=True):
                        st.session_state[f"show_{doc_id}"] = not st.session_state.get(f"show_{doc_id}", False)
                with b2:
                    if st.button("出题", key=f"quiz_{doc_id}", use_container_width=True):
                        prompt = f"""<role>考研数学命题专家</role>
<task>根据知识点出一道考研数学题，并给出详细解答</task>
<knowledge>
{doc_text[:500]}
</knowledge>
<rules>
- 必须同时包含[题目]和[解答]两个部分
- 题目难度匹配考研真题，解答包含关键公式推导
- 禁止输出思考过程、开场白、无关讨论
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                        try:
                            # === 出题自纠重试：最多3次，格式不对让AI重新思考 ===
                            quiz_out = None
                            last_raw = ""
                            for attempt in range(3):
                                if attempt == 0:
                                    quiz_raw = call_llm_api(prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.2)
                                elif attempt == 1:
                                    correct_prompt = f"""<role>考研数学命题专家</role>
<task>你上一轮的输出格式不符合要求，请严格纠正后重新输出</task>
<knowledge>
{doc_text[:500]}
</knowledge>
<bad_example>
以下是你上一轮的错误输出（格式不符合要求）：
---
{last_raw[:400]}
---
</bad_example>
<rules>
- 必须严格使用 [题目] 和 [解答] 两个标记分栏
- [题目] 下写题目内容，[解答] 下写详细解答
- 禁止输出任何思考过程、开场白、解释文字
- 直接从 [题目] 开始输出
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                                    quiz_raw = call_llm_api(correct_prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.2)
                                else:
                                    strict_prompt = f"""<role>考研数学命题专家</role>
<task>这是最后一次机会。你必须严格遵循格式输出题目和解答。</task>
<knowledge>
{doc_text[:500]}
</knowledge>
<rules>
- 第一行必须是 [题目]
- 题目后必须有 [解答]
- 禁止任何前缀、后缀、思考文字
- 禁止输出除题目和解答之外的任何内容
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                                    quiz_raw = call_llm_api(strict_prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.1)
                                last_raw = quiz_raw or ""
                                # === 格式提取 ===
                                q_text = a_text = ""
                                qm = re.search(r'\[题目\]\s*\n?(.*?)(?=\[解答\]|$)', quiz_raw or "", re.DOTALL)
                                am = re.search(r'\[解答\]\s*\n?(.*?)$', quiz_raw or "", re.DOTALL)
                                if qm: q_text = qm.group(1).strip()
                                if am: a_text = am.group(1).strip()
                                if not q_text:
                                    qm2 = re.search(r'(?:^|\n)\s*题目[：:]\s*(.+?)(?=\n\s*解答[：:]|\Z)', quiz_raw or "", re.DOTALL)
                                    if qm2: q_text = qm2.group(1).strip()
                                if not a_text:
                                    am2 = re.search(r'(?:^|\n)\s*解答[：:]\s*(.+?)(?=\Z)', quiz_raw or "", re.DOTALL)
                                    if am2: a_text = am2.group(1).strip()
                                if len(q_text) > 5 and len(a_text) > 5:
                                    quiz_out = f"**[题目]**\n\n{q_text}\n\n**[解答]**\n\n{a_text}"
                                    if attempt > 0:
                                        quiz_out += f"\n\n*（第{attempt+1}次重试后符合格式）*"
                                    break
                            # === 展示结果 ===
                            if quiz_out:
                                with st.container(border=True):
                                    st.markdown(_escape_md(_collapse_math(_fix_latex(quiz_out))))
                                _katex_refresh()
                            else:
                                # 3次全部失败
                                st.error("AI 3次尝试均未输出符合格式的 [题目] + [解答]，请手动重试。")
                                if st.session_state.get("debug_mode"):
                                    st.text_area("最后一次原始输出", last_raw[:1000], height=200)
                        except Exception as e:
                            st.error(f"出题请求失败: {e}")
                            if st.session_state.get("debug_mode"):
                                st.exception(e)
                        except Exception as e:
                            st.error(f"出题请求失败: {e}")
                            if st.session_state.get("debug_mode"):
                                st.exception(e)
                with b3:
                    if st.button("概念自测", key=f"concept_{doc_id}", use_container_width=True):
                        prompt = f"""<role>考研数学老师</role>
<task>根据知识点生成3个概念自测问题，检验学生是否真正理解</task>
<knowledge>
{doc_text[:350]}
</knowledge>
<rules>
- 每个问题不超过25字
- 只输出3行，每行以序号开头
- 禁止输出分析、解释、开场白、思考过程
- 直接从问题列表开始输出，不要任何前缀文字
</rules>
<format>
[输出]
1. 问题一
2. 问题二
3. 问题三
</format>"""
                        try:
                            concept_raw = call_llm_api(prompt, model="mimo-v2.5", max_tokens=600, temperature=0.2)
                            # === 概念自测专属后处理：只提取有实质内容的编号行 ===
                            if concept_raw:
                                valid_lines = []
                                for line in concept_raw.split('\n'):
                                    s = line.strip()
                                    # 必须: 编号 + 非空白内容（至少5个字符的问题才有意义）
                                    if re.match(r'^\d+[\.\、\)\)]\s*\S.{3,}', s):
                                        # 排除思维链伪编号: "1. 首先分析" "2. 然后设计" 等
                                        if not any(kw in s for kw in (
                                            '首先', '然后', '接着', '最后', '分析', '设计', '考虑',
                                            '步骤', '思路', '方案', '策略', '方法', '需要', '应该',
                                            '了解', '掌握', '理解知识点', '回顾', '思考',
                                        )):
                                            valid_lines.append(s)
                                if valid_lines:
                                    concept = '\n'.join(valid_lines)
                                else:
                                    concept = concept_raw  # 没有有效编号行，保留原文
                            else:
                                concept = concept_raw
                            # === 验证 ===
                            has_valid = bool(valid_lines)
                            if concept and has_valid and len(concept) > 10:
                                with st.container(border=True):
                                    st.markdown(_escape_md(_collapse_math(_fix_latex(concept))))
                                _katex_refresh()
                            elif concept and len(concept) > 10 and not has_valid:
                                st.warning(f"AI 输出格式异常（未提取到有效问题），可能是思维链残留：\n\n{concept[:500]}")
                            else:
                                st.warning("AI 未能生成自测问题，请重试")
                        except Exception as e:
                            st.error(f"概念自测请求失败: {e}")
                            if st.session_state.get("debug_mode"):
                                st.exception(e)

                # 展开的内容区域（更长）
                if st.session_state.get(f"show_{doc_id}", False):
                    st.markdown(f"""
                    <div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;padding:16px;font-size:0.84rem;color:#475569;line-height:1.75;margin-bottom:16px;max-height:450px;overflow-y:auto;">
                        {doc_text}
                    </div>
                    """, unsafe_allow_html=True)

        with _tabs[1]:
            st.markdown("""
            <div style="font-size:0.88rem;font-weight:700;color:#1e293b;margin-bottom:12px;">遗忘曲线复习候选</div>
            """, unsafe_allow_html=True)

            # 从数据库读取需要复习的知识点
            uid = st.session_state.get("user_id")
            review_items = []
            if uid:
                try:
                    init_memory_db()
                    conn = sqlite3.connect(MEMORY_DB)
                    cur = conn.execute(
                        "SELECT knowledge_id, mastery_level, times_correct, times_wrong, error_type, last_review "
                        "FROM knowledge_mastery WHERE user_id=? AND mastery_level < 60 AND status != '已掌握' "
                        "ORDER BY mastery_level ASC, last_review ASC LIMIT 10",
                        (uid,))
                    for row in cur.fetchall():
                        kid, ml, tc, tw, et, lr = row
                        # 从语料库中查找知识点名称
                        corpus = load_corpus()
                        name = kid
                        desc = ""
                        for doc in corpus:
                            if kid in doc.get("id", ""):
                                name = doc["id"].replace(".md", "").replace("_", " ")
                                desc = doc.get("text", "")[:200]
                                break
                        review_items.append({
                            "kid": kid, "name": name, "mem": max(10, int(ml or 0)),
                            "desc": desc or f"正确{tw or 0}次 · 错误{tw or 0}次",
                            "correct": tw or 0, "wrong": tw or 0
                        })
                    conn.close()
                except Exception:
                    pass

            if not review_items:
                # demo fallback
                review_items = [
                    {"kid": "001", "name": "泰勒公式与麦克劳林展开", "mem": 45, "desc": "泰勒公式用于将函数在某点附近展开为多项式形式。麦克劳林展开是 x=0 处的特例。", "correct": 3, "wrong": 2},
                    {"kid": "006", "name": "矩阵的特征值与特征向量", "mem": 62, "desc": "特征值与特征向量用于矩阵对角化、二次型标准化。", "correct": 5, "wrong": 4},
                    {"kid": "012", "name": "拉格朗日中值定理", "mem": 38, "desc": "拉格朗日中值定理是微分学核心定理之一。", "correct": 2, "wrong": 3},
                ]

            for item in review_items:
                mem = item["mem"]
                _mem_color = "#10b981" if mem >= 60 else ("#f59e0b" if mem >= 40 else "#ef4444")
                kid = item["kid"]
                st.markdown(f"""
                <div style="background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:14px 16px;margin-bottom:8px;">
                    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px;">
                        <span style="font-weight:600;font-size:0.88rem;color:#1e293b;">{item['name']}</span>
                        <span style="font-size:0.78rem;padding:4px 10px;background:{_mem_color}15;color:{_mem_color};border-radius:20px;font-weight:600;">记忆率 {mem}%</span>
                    </div>
                    <div style="font-size:0.82rem;color:#64748b;margin-bottom:10px;max-height:120px;overflow-y:auto;">{item['desc']}</div>
                </div>
                """, unsafe_allow_html=True)
                b1, b2, b3 = st.columns(3)
                with b1:
                    if st.button("✅ 掌握", key=f"rev_ok_{kid}", use_container_width=True):
                        st.session_state["mastered_count"] = st.session_state.get("mastered_count", 0) + 1
                        st.success(f"已标记 {item['name']} 为掌握！")
                        st.rerun()
                with b2:
                    if st.button("🔄 再练", key=f"rev_retry_{kid}", use_container_width=True):
                        st.info(f"已将 {item['name']} 加入再练列表")
                with b3:
                    if st.button("🎲 出题", key=f"rev_quiz_{kid}", use_container_width=True):
                        corpus = load_corpus()
                        doc_text = ""
                        for d in corpus:
                            if kid in d.get("id", ""):
                                doc_text = d.get("text", "")[:500]
                                break
                        if doc_text:
                            try:
                                # === 出题自纠重试：最多3次 ===
                                quiz_out = None
                                last_raw = ""
                                base_prompt = f"""<role>考研数学命题专家</role>
<task>根据知识点出一道考研数学题，并给出详细解答</task>
<knowledge>
{doc_text}
</knowledge>
<rules>
- 必须同时包含[题目]和[解答]两个部分
- 题目难度匹配考研真题，解答包含关键公式推导
- 禁止输出思考过程、开场白、无关讨论
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                                for attempt in range(3):
                                    if attempt == 0:
                                        quiz_raw = call_llm_api(base_prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.2)
                                    elif attempt == 1:
                                        correct_prompt = f"""<role>考研数学命题专家</role>
<task>你上一轮的输出格式不符合要求，请严格纠正后重新输出</task>
<knowledge>
{doc_text}
</knowledge>
<bad_example>
以下是你上一轮的错误输出（格式不符合要求）：
---
{last_raw[:400]}
---
</bad_example>
<rules>
- 必须严格使用 [题目] 和 [解答] 两个标记分栏
- [题目] 下写题目内容，[解答] 下写详细解答
- 禁止输出任何思考过程、开场白、解释文字
- 直接从 [题目] 开始输出
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                                        quiz_raw = call_llm_api(correct_prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.2)
                                    else:
                                        strict_prompt = f"""<role>考研数学命题专家</role>
<task>这是最后一次机会。你必须严格遵循格式输出题目和解答。</task>
<knowledge>
{doc_text}
</knowledge>
<rules>
- 第一行必须是 [题目]
- 题目后必须有 [解答]
- 禁止任何前缀、后缀、思考文字
- 禁止输出除题目和解答之外的任何内容
</rules>
<format>
[题目]
（题目内容）
[解答]
（详细解答过程）
</format>"""
                                        quiz_raw = call_llm_api(strict_prompt, model="mimo-v2.5", max_tokens=3000, temperature=0.1)
                                    last_raw = quiz_raw or ""
                                    q_text = a_text = ""
                                    qm = re.search(r'\[题目\]\s*\n?(.*?)(?=\[解答\]|$)', quiz_raw or "", re.DOTALL)
                                    am = re.search(r'\[解答\]\s*\n?(.*?)$', quiz_raw or "", re.DOTALL)
                                    if qm: q_text = qm.group(1).strip()
                                    if am: a_text = am.group(1).strip()
                                    if not q_text:
                                        qm2 = re.search(r'(?:^|\n)\s*题目[：:]\s*(.+?)(?=\n\s*解答[：:]|\Z)', quiz_raw or "", re.DOTALL)
                                        if qm2: q_text = qm2.group(1).strip()
                                    if not a_text:
                                        am2 = re.search(r'(?:^|\n)\s*解答[：:]\s*(.+?)(?=\Z)', quiz_raw or "", re.DOTALL)
                                        if am2: a_text = am2.group(1).strip()
                                    if len(q_text) > 5 and len(a_text) > 5:
                                        quiz_out = f"**[题目]**\n\n{q_text}\n\n**[解答]**\n\n{a_text}"
                                        if attempt > 0:
                                            quiz_out += f"\n\n*（第{attempt+1}次重试后符合格式）*"
                                        break
                                if quiz_out:
                                    with st.container(border=True):
                                        st.markdown(_escape_md(_collapse_math(_fix_latex(quiz_out))))
                                    _katex_refresh()
                                else:
                                    st.error("AI 3次尝试均未输出符合格式的 [题目] + [解答]，请手动重试。")
                                    if st.session_state.get("debug_mode"):
                                        st.text_area("最后一次原始输出", last_raw[:1000], height=200)
                            except Exception as e:
                                st.error(f"出题请求失败: {e}")
                                if st.session_state.get("debug_mode"):
                                    st.exception(e)

            if not review_items:
                st.info("暂无待复习知识点。使用问答后系统会自动记录。")

        with _tabs[2]:
            st.markdown("""
            <div style="font-size:0.88rem;font-weight:700;color:#1e293b;margin-bottom:12px;">费曼学习法</div>
            <div style="padding:10px 14px;background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;font-size:0.82rem;color:#4f46e5;margin-bottom:14px;">
                选择模式 → 输入题目或上传图片 → 写下你的答案 → AI 评价
            </div>
            """, unsafe_allow_html=True)
            _fm = st.radio("模式", ["概念理解", "解题练习"], horizontal=True, key="feynman_mode")
            f1, f2 = st.columns(2)
            with f1:
                st.markdown('<div style="font-size:0.82rem;font-weight:600;color:#64748b;margin-bottom:4px;">输入题目</div>', unsafe_allow_html=True)
                _ftopic = st.text_area("题目", placeholder="例如：什么是洛必达法则？\n或：求函数 f(x)=x³-3x+2 的极值", key="feynman_topic", label_visibility="collapsed")
            with f2:
                st.markdown('<div style="font-size:0.82rem;font-weight:600;color:#64748b;margin-bottom:4px;">你的答案</div>', unsafe_allow_html=True)
                _fanswer = st.text_area("答案", placeholder="用自己的话写下答案...\n\n提示：尽量用自己的语言表达，展示你的理解过程。", key="feynman_answer", label_visibility="collapsed", height=200)
            if st.button("提交答案", use_container_width=True, type="primary"):
                if _ftopic and _fanswer:
                    prompt = f"""你是考研辅导专家。请评价学生对以下题目的回答。
题目: {_ftopic}
学生答案: {_fanswer}
请给出总评分(满分20)和分项评价(概念理解/表达能力)，指出不足并鼓励。"""
                    result = call_llm_api(prompt, model="mimo-v2.5", max_tokens=600)
                    st.success(result)

        with _tabs[3]:
            st.markdown("""
            <div style="font-size:0.88rem;font-weight:700;color:#1e293b;margin-bottom:14px;">知识点掌握情况</div>
            """, unsafe_allow_html=True)
            _all = len(load_corpus())
            uid = st.session_state.get("user_id")
            _mastered = 0
            _cards = []
            if uid:
                try:
                    init_memory_db()
                    conn = sqlite3.connect(MEMORY_DB)
                    # 统计已掌握数量
                    cur = conn.execute("SELECT COUNT(*) FROM knowledge_mastery WHERE user_id=? AND status='已掌握'", (uid,))
                    _mastered = cur.fetchone()[0]
                    # 获取最近的知识点状态
                    cur = conn.execute(
                        "SELECT knowledge_id, mastery_level, times_correct, times_wrong, status "
                        "FROM knowledge_mastery WHERE user_id=? ORDER BY last_review DESC LIMIT 20",
                        (uid,))
                    for row in cur.fetchall():
                        kid, ml, tc, tw, status = row
                        corpus = load_corpus()
                        name = kid
                        for doc in corpus:
                            if kid in doc.get("id", ""):
                                name = doc["id"].replace(".md", "").replace("_", " ")
                                break
                        _cards.append({
                            "status": "mastered" if (status == "已掌握" or (ml or 0) >= 80) else "learning",
                            "name": name,
                            "correct": tc or 0, "wrong": tw or 0
                        })
                    conn.close()
                except Exception:
                    pass
            else:
                _mastered = st.session_state.get("mastered_count", 0)

            _pct = int(_mastered / max(_all, 1) * 100)
            st.markdown(f"""
            <div style="display:flex;justify-content:space-between;font-size:0.82rem;margin-bottom:6px;">
                <span style="color:#64748b;">掌握进度</span>
                <strong style="color:#1e293b;">{_mastered}/{_all} ({_pct}%)</strong>
            </div>
            """, unsafe_allow_html=True)
            st.progress(_pct / 100)
            st.markdown('<div style="height:8px;"></div>', unsafe_allow_html=True)

            if _cards:
                for item in _cards:
                    status = item["status"]
                    _bg = "#f0fdf4" if status == "mastered" else "#fff7ed"
                    _border = "#22c55e" if status == "mastered" else "#f97316"
                    _color = "#166534" if status == "mastered" else "#c2410c"
                    st.markdown(f"""
                    <div style="padding:10px 14px;margin-bottom:4px;border-radius:8px;font-size:0.82rem;font-weight:500;background:{_bg};color:{_color};border-left:3px solid {_border};">
                        {item['name']} · ✓{item['correct']} ✗{item['wrong']}
                    </div>
                    """, unsafe_allow_html=True)
            else:
                # 无数据时显示提示
                _sample_cards = [
                    ("mastered", "01. 泰勒公式与麦克劳林展开", "✓12 ✗2"),
                    ("mastered", "02. 定积分的几何与物理应用", "✓8 ✗3"),
                    ("learning", "03. 矩阵的特征值与特征向量", "✓5 ✗6"),
                ]
                for status, name, stats in _sample_cards:
                    _bg = "#f0fdf4" if status == "mastered" else "#fff7ed"
                    _border = "#22c55e" if status == "mastered" else "#f97316"
                    _color = "#166534" if status == "mastered" else "#c2410c"
                    st.markdown(f"""
                    <div style="padding:10px 14px;margin-bottom:4px;border-radius:8px;font-size:0.82rem;font-weight:500;background:{_bg};color:{_color};border-left:3px solid {_border};">
                        {name} · {stats}
                    </div>
                    """, unsafe_allow_html=True)
                st.caption("👆 以上为示例数据。开始使用问答后，系统会自动追踪你的掌握情况。")
            st.markdown('<div style="text-align:center;padding:12px;font-size:0.75rem;color:#94a3b8;">使用问答后系统自动追踪 · 共 110 个知识点</div>', unsafe_allow_html=True)

    st.stop()

# ==================== Hub 主界面 ====================
if st.session_state.page == "hub":

    # ══════════════════════════════════════
    # MAIN CONTENT
    # ══════════════════════════════════════

    # ── Row 1: Hero (2/3) + Metric (1/3) ──
    _exam_date = datetime(2026, 12, 21)
    _days_left = (_exam_date - datetime.now()).days
    h_left, h_right = st.columns([2, 1])
    with h_left:
        st.markdown(f"""
        <div class="main-title" style="text-align:left;display:flex;justify-content:space-between;align-items:center;">
            <div>
                <h1 style="font-size:1.5rem!important;">2026 年全国统考倒计时</h1>
                <p>星光不问赶路人，时光不负有心人</p>
            </div>
            <div style="background:rgba(255,255,255,0.1);padding:1rem 1.5rem;border-radius:12px;border:1px solid rgba(255,255,255,0.12);text-align:center;">
                <span style="font-size:2.5rem;font-weight:900;display:block;line-height:1;">{_days_left}</span>
                <span style="font-size:0.7rem;letter-spacing:2px;opacity:0.55;">DAYS LEFT</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with h_right:
        hub_user_id = st.session_state.get("user_id")
        if hub_user_id:
            flow_data = get_flow_focus_data(hub_user_id)
            flow_msg = pick_flow_message(hub_user_id)
            focus_hours = flow_data["total_hours"]
            progress_pct = flow_data["progress_pct"]
            # 进度颜色：绿(≥80%) / 靛(≥50%) / 琥珀(>0%) / 灰(0%)
            if progress_pct >= 80:
                bar_color = "#22c55e"; bg_color = "#f0fdf4"
            elif progress_pct >= 50:
                bar_color = "#4f46e5"; bg_color = "#eef2ff"
            elif progress_pct > 0:
                bar_color = "#f59e0b"; bg_color = "#fffbeb"
            else:
                bar_color = "#94a3b8"; bg_color = "#f8fafc"
            st.markdown(f"""
            <div class="qa-card" style="text-align:center;">
                <div style="font-size:0.82rem;font-weight:700;color:#64748b;margin-bottom:6px;">⚡ 今日聚焦心流</div>
                <div style="font-size:2.2rem;font-weight:800;color:{bar_color};">{focus_hours}
                    <span style="font-size:0.85rem;color:#64748b;font-weight:400;">小时</span>
                </div>
                <div style="display:inline-flex;align-items:center;gap:6px;font-size:0.75rem;padding:4px 12px;background:{bg_color};color:{bar_color};border-radius:20px;font-weight:500;">
                    <span style="width:6px;height:6px;border-radius:50%;background:{bar_color};"></span> {'已达每日目标 ' + str(progress_pct) + '%' if progress_pct > 0 else '今日尚未开始'}
                </div>
                <div style="margin-top:10px;font-size:0.78rem;color:#475569;line-height:1.5;padding:0 4px;">{flow_msg}</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="qa-card" style="text-align:center;">
                <div style="font-size:0.82rem;font-weight:700;color:#64748b;margin-bottom:6px;">⚡ 今日聚焦心流</div>
                <div style="font-size:2.2rem;font-weight:800;color:#94a3b8;">—
                    <span style="font-size:0.85rem;color:#64748b;font-weight:400;">小时</span>
                </div>
                <div style="display:inline-flex;align-items:center;gap:6px;font-size:0.75rem;padding:4px 12px;background:#f8fafc;color:#94a3b8;border-radius:20px;font-weight:500;">
                    登录后开始记录
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Row 2: 4 Feature Cards ──
    fc1, fc2, fc3, fc4 = st.columns(4)

    _card_data = [
        ("math", "icon-math", """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="7" height="7" rx="1"/><rect x="14" y="3" width="7" height="7" rx="1"/><rect x="3" y="14" width="7" height="7" rx="1"/><rect x="14" y="14" width="7" height="7" rx="1"/></svg>""",
         "数学问答", "110 个知识点 · 智能问答 · 遗忘曲线复习",
         ["微积分", "线代", "概率"], "main"),
        ("popularity", "icon-fire", """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>""",
         "高校热度", "院校专业热度 · 报录比 · 复试线",
         ["报录比", "趋势分析", "专业查询"], "popularity"),
        ("english", "icon-eng", """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/><line x1="8" y1="7" x2="16" y2="7"/><line x1="8" y1="11" x2="14" y2="11"/></svg>""",
         "英语专家", "作文批改 · 长难句解析 · 翻译 · 单词记忆",
         ["作文", "翻译", "单词"], "english"),
        ("suggest", "icon-fb", """<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/><path d="M8 9h8"/><path d="M8 13h5"/></svg>""",
         "反馈建议", "欢迎提出建议，会及时跟进修改",
         ["需求", "反馈"], "suggest"),
    ]

    for col, (key, icon_cls, icon_svg, title, desc, tags, target) in zip(
        [fc1, fc2, fc3, fc4], _card_data
    ):
        with col:
            tags_html = "".join(f'<span class="card-tag">{t}</span>' for t in tags) if tags else ""
            st.markdown(f"""
            <div class="feature-card">
                <div class="card-icon {icon_cls}">{icon_svg}</div>
                <div class="card-title">{title}</div>
                <div class="card-desc">{desc}</div>
                {f'<div class="card-tags">{tags_html}</div>' if tags_html else ''}
            </div>
            """, unsafe_allow_html=True)
            if st.button("进入", key=f"hub_{key}", use_container_width=True):
                st.session_state.page = target
                st.rerun()

    # ── Row 3: 2 Wide Cards ──
    wc1, wc2 = st.columns(2)

    with wc1:
        st.markdown("""
        <div class="feature-card" style="display:flex;align-items:center;gap:16px;">
            <div class="card-icon icon-mat" style="width:48px;height:48px;margin-bottom:0;flex-shrink:0;"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg></div>
            <div style="flex:1;">
                <div class="card-title">学习资料生成</div>
                <div class="card-desc">AI 生成习题册 · 知识点整理 · DOCX 导出</div>
                <div class="card-tags" style="margin-top:6px;">
                    <span class="card-tag">习题生成</span><span class="card-tag">模考卷</span><span class="card-tag">DOCX</span>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("进入", key="hub_material", use_container_width=True):
            st.session_state.page = "material"
            st.rerun()

    with wc2:
        st.markdown("""
        <div class="feature-card" style="display:flex;align-items:center;gap:16px;">
            <div class="card-icon icon-ck" style="width:48px;height:48px;margin-bottom:0;flex-shrink:0;"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg></div>
            <div style="flex:1;">
                <div class="card-title">打卡督学</div>
                <div class="card-desc">每日打卡 · 学习计划 · 学习日记 · 番茄计时 · 学习画像</div>
                <div class="card-tags" style="margin-top:6px;">
                    <span class="card-tag">打卡</span><span class="card-tag">番茄钟</span><span class="card-tag">日记</span><span class="card-tag">画像</span>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("进入", key="hub_checkin", use_container_width=True):
            st.session_state.page = "checkin"
            st.rerun()

    st.stop()

# ==================== 高校热度查询 ====================
if st.session_state.page == "popularity":
    st.markdown("""
    <div class="main-title" style="text-align:left;padding:1.2rem 1.8rem;">
        <div style="display:flex;align-items:center;gap:14px;">
            <div style="width:42px;height:42px;border-radius:12px;background:linear-gradient(135deg,#db2777,#ec4899);display:flex;align-items:center;justify-content:center;color:#fff;flex-shrink:0;box-shadow:0 4px 12px rgba(219,39,119,0.3);"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" width="24" height="24"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg></div>
            <div>
                <h1 style="margin:0;font-size:1.4rem!important;">高校热度查询</h1>
                <p style="margin:2px 0 0;font-size:0.82rem;opacity:0.75;">院校信息与报考热度分析</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    col_back, col_status = st.columns([1, 3])
    with col_back:
        if st.button("← 返回首页"):
            st.session_state.page = "hub"
            st.rerun()
    with col_status:
        node_ok = kaoyan_predict.check_node_available()
        if not node_ok:
            st.warning("⚠️ 未检测到 Node.js，预测功能不可用")

    st.markdown("---")

    with st.form("popularity_form"):
        col_sch, col_maj = st.columns(2)
        with col_sch:
            school = st.text_input("🏫 学校名称", placeholder="例如：华东师范大学")
        with col_maj:
            major = st.text_input("专业名称（可留空）", placeholder="例如：生物学")
        submitted = st.form_submit_button("查询热度", use_container_width=True, type="primary")

    if "_kaoyan_cache" not in st.session_state:
        st.session_state._kaoyan_cache = {}

    if submitted and school:
        cache_key = f"{school.strip()}|{major.strip()}"
        with st.spinner("正在分析院校热度（数据 + 媒体双引擎）..."):
            try:
                raw = kaoyan_predict.predict(school, major)
                data = kaoyan_predict.normalize_for_ui(raw)
                st.session_state._kaoyan_cache[cache_key] = data
                data["_school"] = school.strip()
                data["_major"] = major.strip() or school.strip()
                st.session_state._kaoyan_last = cache_key
            except kaoyan_predict.KaoyanPredictError as e:
                st.error(f"预测失败：{e}")
                st.session_state._kaoyan_last = None

    last_key = st.session_state.get("_kaoyan_last")
    if last_key and last_key in st.session_state._kaoyan_cache:
        data = st.session_state._kaoyan_cache[last_key]
        heat = data["compositeHeat"]
        level = data["heatLevel"]

        # 显示查询的学校和专业
        _school_name = data.get("_school", "?")
        _major_name = data.get("_major", "?")
        st.markdown(f"#### 🎓 {_school_name} · {_major_name}")

        st.markdown('<div class="qa-card">', unsafe_allow_html=True)
        st.markdown(f"### {level['color']} 综合热度 {heat}/100  ·  {level['label']}")
        st.progress(heat / 100)

        col_d, col_m = st.columns(2)
        with col_d:
            st.metric("数据热度", f"{data['dataHeat']}/100")
        with col_m:
            st.metric("📱 媒体热度", f"{data['mediaHeat']}/100")

        st.caption(f"📡 {data['dataSource']}  |  🎯 置信度 {data['confidence']}%  |  趋势 {data['trend']}")
        st.markdown('</div>', unsafe_allow_html=True)

        if data.get("admissionHistory"):
            st.markdown("### 录取历史")
            rows = []
            has_valid = False
            for h in data["admissionHistory"]:
                row = {
                    "年份": h.get("year", "?"),
                    "报考": f"{h.get('applicants', 0)}人",
                    "录取": f"{h.get('admitted', 0)}人",
                    "报录比": f"{h.get('ratio', '?')}:1" if h.get("ratio") else "?",
                    "复试线": f"{h.get('cutScore', '?')}分" if h.get("cutScore") else "?",
                }
                if h.get("admitted"):
                    has_valid = True
                rows.append(row)
            st.dataframe(rows, use_container_width=True, hide_index=True)
            if not has_valid:
                st.caption("📌 录取历史数据来自公开API，部分年份数据不完整。建议结合目标院校官网核实。")

        pred = data.get("prediction", {})
        st.markdown("### 🔮 27届预测")
        col_p1, col_p2, col_p3 = st.columns(3)
        with col_p1:
            val = pred.get('estimatedApplicants', 0) or 0
            st.metric("预计报考人数", f"{val}人" if val > 0 else "暂无数据")
        with col_p2:
            val = pred.get('estimatedRatio', 0) or 0
            st.metric("预计报录比", f"{val}:1" if val > 0 else "暂无数据")
        with col_p3:
            val = pred.get('estimatedCutScore', 0) or 0
            st.metric("预计复试线", f"{val}分" if val > 0 else "暂无数据")

        st.caption("预测基于历史录取数据和媒体热度，仅供参考。实际数据以目标院校官网为准。")

        # ── 可点击展开的详情卡片（参考数学Hub知识点展示方案）──
        st.markdown("---")
        subjects = data.get("examSubjects", [])
        platforms = data.get("platforms", [])
        si = data.get("schoolInfo", {})
        notes = data.get("notes", [])

        # 卡片1：考试科目
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border-radius:12px;padding:14px 18px;margin-bottom:4px;border:1px solid #bae6fd;">
            <span class="info-badge">科目</span>
            <span style="font-size:0.9rem;font-weight:600;color:#0369a1;">考试科目（{len(subjects)}门）</span>
        </div>
        """, unsafe_allow_html=True)
        if st.button("展开 / 收起", key="pop_toggle_subjects", use_container_width=True):
            st.session_state["pop_show_subjects"] = not st.session_state.get("pop_show_subjects", False)
        if st.session_state.get("pop_show_subjects", False):
            if subjects:
                for s in subjects:
                    st.markdown(f"- **{s.get('code','?')}** {s.get('name','?')}（{s.get('type','?')}）")
            else:
                st.caption("暂无考试科目数据，请参考目标院校官网招生简章。")

        st.markdown("<br>", unsafe_allow_html=True)

        # 卡片2：平台热度
        plat_ok = sum(1 for p in platforms if p.get("score") is not None)
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border-radius:12px;padding:14px 18px;margin-bottom:4px;border:1px solid #bae6fd;">
            <span class="info-badge">热度</span>
            <span style="font-size:0.9rem;font-weight:600;color:#0369a1;">平台热度详情（{plat_ok}/{len(platforms)}平台有数据）</span>
        </div>
        """, unsafe_allow_html=True)
        if st.button("展开 / 收起", key="pop_toggle_platforms", use_container_width=True):
            st.session_state["pop_show_platforms"] = not st.session_state.get("pop_show_platforms", False)
        if st.session_state.get("pop_show_platforms", False):
            if platforms:
                cols = st.columns(2)
                for i, p in enumerate(platforms):
                    name = p.get("name", "?")
                    score = p.get("score")
                    weight = p.get("weight", 0)
                    with cols[i % 2]:
                        if score is None:
                            st.markdown(f"❌ **{name}** — 抓取失败")
                        else:
                            bar_char = "▓" * max(1, int(score / 100 * 20))
                            w_pct = f"({int(weight * 100)}%)" if weight else ""
                            st.markdown(f"**{name}** {w_pct}  \n`{bar_char}` {score}/100")
                if data.get("failedPlatforms"):
                    st.caption(f"⚠️ 部分平台抓取失败：{', '.join(data['failedPlatforms'])}")
                st.caption("数据来源：B站、百度、微信、QQ 等公开平台搜索热度")
            else:
                st.caption("暂无平台热度数据。")

        st.markdown("<br>", unsafe_allow_html=True)

        # 卡片3：院校信息
        level = si.get("schoolLevel", "") or "暂未收录"
        dept = si.get("department", "") or "暂未收录"
        push = si.get("pushRatioDesc", "")
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border-radius:12px;padding:14px 18px;margin-bottom:4px;border:1px solid #bae6fd;">
            <span class="info-badge">院校</span>
            <span style="font-size:0.9rem;font-weight:600;color:#0369a1;">院校信息</span>
        </div>
        """, unsafe_allow_html=True)
        if st.button("展开 / 收起", key="pop_toggle_school", use_container_width=True):
            st.session_state["pop_show_school"] = not st.session_state.get("pop_show_school", False)
        if st.session_state.get("pop_show_school", False):
            st.markdown(f"- **层次**：{level}")
            st.markdown(f"- **院系**：{dept}")
            if push:
                st.markdown(f"- **推免**：{push}")
            else:
                st.caption("更详细的院校信息请查阅学校官网。")

        st.markdown("<br>", unsafe_allow_html=True)

        # 卡片4：备注
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#e0f2fe,#bae6fd);border-radius:12px;padding:14px 18px;margin-bottom:4px;border:1px solid #bae6fd;">
            <span class="info-badge">备注</span>
            <span style="font-size:0.9rem;font-weight:600;color:#0369a1;">备注（{len(notes)}条）</span>
        </div>
        """, unsafe_allow_html=True)
        if st.button("展开 / 收起", key="pop_toggle_notes", use_container_width=True):
            st.session_state["pop_show_notes"] = not st.session_state.get("pop_show_notes", False)
        if st.session_state.get("pop_show_notes", False):
            for n in notes:
                st.markdown(f"- {n}")
        else:
            st.caption("暂无备注信息。")

        # ── 个人建议 ──
        st.markdown("---")
        st.markdown("### 💡 个人建议")

        uid = st.session_state.get("user_id", 1)
        profile = get_user_profile(uid)

        if not profile:
            st.info("请先在「打卡督学 → 学习画像」中填写个人信息，以获取个性化报考建议。")
        else:
            summary_parts = []
            if profile.get("undergraduate_level"):
                summary_parts.append(f"本科{profile['undergraduate_level']}")
            if profile.get("grade"):
                summary_parts.append(profile["grade"])
            if profile.get("target_major"):
                summary_parts.append(f"目标{profile['target_major']}")
            if summary_parts:
                st.caption("当前画像：" + " · ".join(summary_parts))

        # 生成建议
        if st.button("生成/刷新个人建议", use_container_width=True, key="gen_rec"):
            if not profile:
                st.info("请先填写个人画像")
            else:
                with st.spinner("🤔 正在结合你的个人画像和院校数据生成建议..."):
                    try:
                        rec_text = generate_recommendation(
                            uid, data,
                            get_profile_fn=get_user_profile,
                            call_llm_fn=call_llm_api,
                        )
                        if rec_text:
                            st.session_state._rec_text = rec_text
                        else:
                            st.info("请先完善个人画像以获取建议。")
                    except Exception as e:
                        st.warning(f"⚠️ 建议生成失败：{e}")

        if st.session_state.get("_rec_text"):
            st.markdown(st.session_state._rec_text)

    elif not submitted:
        st.info("👆 输入学校和专业名称，点击「查询热度」开始分析")

    st.stop()

# ==================== 学习资料 ====================
if st.session_state.page == "material":
    if st.button("← 返回首页"):
        st.session_state.page = "hub"
        st.rerun()
    st.markdown("""
    <div class="main-title" style="text-align:left;padding:1.2rem 1.8rem;">
        <div style="display:flex;align-items:center;gap:14px;">
            <div style="width:42px;height:42px;border-radius:12px;background:linear-gradient(135deg,#ca8a04,#eab308);display:flex;align-items:center;justify-content:center;color:#fff;flex-shrink:0;box-shadow:0 4px 12px rgba(202,138,4,0.3);"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" width="24" height="24"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg></div>
            <div>
                <h1 style="margin:0;font-size:1.4rem!important;">学习资料生成</h1>
                <p style="margin:2px 0 0;font-size:0.82rem;opacity:0.75;">AI 生成习题册 · 知识点整理 · 备考资料</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    tab_math, tab_english = st.tabs(["数学资料", "英语资料"])

    # ── 数学资料 ──
    with tab_math:
        # 知识点选择
        corpus_files = sorted([
            f.stem for f in DATA_DIR.glob("*.md")
        ])
        selected_topics = st.multiselect(
            "选择知识点范围（可多选，留空则使用全部）",
            corpus_files,
            placeholder="例如：定积分、微分方程、中值定理...",
            key="mat_topics"
        )

        # 用户需求输入
        user_requirement = st.text_area(
            "描述你想要的资料",
            height=100,
            placeholder="帮我生成一份积分典型题习题集，要基础题和难题都有",
            key="mat_requirement"
        )

        # 生成按钮
        gen_col1, gen_col2 = st.columns([1, 3])
        with gen_col1:
            generate_btn = st.button("🚀 生成资料", type="primary", use_container_width=True, key="mat_gen")
        with gen_col2:
            if not (user_requirement or "").strip():
                st.caption("💡 在上方输入框中描述你想要的资料类型，然后点击生成")

        if generate_btn:
            if not (user_requirement or "").strip():
                st.warning("请先输入你对资料的需求描述")
            elif not API_KEY:
                st.error("未配置 AI API Key，无法生成")
            else:
                with st.spinner("AI 正在生成资料，请稍候..."):
                    prompt = _build_material_prompt(selected_topics, user_requirement)
                    try:
                        reasoning, result_text = _generate_material(prompt)
                        docx_bytes = _ai_output_to_docx_via_pandoc(result_text)

                        # 简短展示思考过程
                        if reasoning:
                            reasoning_lines = [l for l in reasoning.split("\n") if l.strip()]
                            brief = reasoning_lines[:3] if len(reasoning_lines) > 3 else reasoning_lines
                            with st.expander(f"💭 AI 思考过程（共 {len(reasoning)} 字）"):
                                st.caption("\n".join(brief))
                                if len(reasoning_lines) > 3:
                                    st.caption(f"...（共 {len(reasoning_lines)} 行思考内容）")

                        # 动态文件名：优先用选中的知识点，否则从用户需求中提取
                        if selected_topics:
                            base = selected_topics[0][:20] if len(selected_topics) == 1 else f"{selected_topics[0][:12]}等{len(selected_topics)}个知识点"
                            file_name = f"{base}资料.docx"
                        else:
                            kw = user_requirement.replace("帮我", "").replace("生成", "").replace("一份", "").strip()[:20]
                            file_name = f"{kw}资料.docx" if kw else "考研数学资料.docx"

                        st.success(f"生成完成！共 {len(result_text)} 字")
                        st.download_button(
                            label="📥 下载 docx",
                            data=docx_bytes,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="mat_dl",
                            type="primary",
                        )
                    except Exception as e:
                        st.error(f"生成失败: {e}")

    # ── 英语资料 ──
    with tab_english:
        eng_categories = [
            "语法专题", "阅读技巧", "写作模板", "翻译技巧",
            "完形填空", "新题型", "单词归类总结"
        ]
        selected_eng_category = st.selectbox(
            "选择知识点分类",
            eng_categories,
            key="eng_category"
        )

        vocab_category = ""
        if selected_eng_category == "单词归类总结":
            vocab_types = ["按主题分类", "按词性分类", "按难度分类", "按真题频率分类"]
            vocab_category = st.selectbox(
                "选择分类方式",
                vocab_types,
                key="vocab_category"
            )

        eng_requirement = st.text_area(
            "描述你想要的资料",
            height=100,
            placeholder="例如：帮我整理定语从句的笔记，要有例句和真题",
            key="eng_requirement"
        )

        eng_col1, eng_col2 = st.columns([1, 3])
        with eng_col1:
            eng_generate_btn = st.button("🚀 生成资料", type="primary", use_container_width=True, key="eng_gen")
        with eng_col2:
            if not (eng_requirement or "").strip():
                st.caption("💡 在上方输入框中描述你想要的资料类型，然后点击生成")

        if eng_generate_btn:
            if not (eng_requirement or "").strip():
                st.warning("请先输入你对资料的需求描述")
            elif not API_KEY:
                st.error("未配置 AI API Key，无法生成")
            else:
                with st.spinner("AI 正在生成资料，请稍候..."):
                    prompt = _build_english_material_prompt(
                        selected_eng_category, eng_requirement, vocab_category
                    )
                    try:
                        reasoning, result_text = _generate_material(prompt)
                        docx_bytes = _ai_output_to_docx_via_pandoc(result_text)

                        if reasoning:
                            reasoning_lines = [l for l in reasoning.split("\n") if l.strip()]
                            brief = reasoning_lines[:3] if len(reasoning_lines) > 3 else reasoning_lines
                            with st.expander(f"💭 AI 思考过程（共 {len(reasoning)} 字）"):
                                st.caption("\n".join(brief))
                                if len(reasoning_lines) > 3:
                                    st.caption(f"...（共 {len(reasoning_lines)} 行思考内容）")

                        file_name = f"考研英语_{selected_eng_category}.docx"

                        st.success(f"生成完成！共 {len(result_text)} 字")
                        st.download_button(
                            label="📥 下载 docx",
                            data=docx_bytes,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="eng_dl",
                            type="primary",
                        )
                    except Exception as e:
                        st.error(f"生成失败: {e}")

    st.stop()


# ==================== 提建议 ====================
if st.session_state.page == "suggest":
    if st.button("← 返回首页"):
        st.session_state.page = "hub"
        st.rerun()
    st.markdown("""
    <div class="main-title">
        <h1>提建议</h1>
        <p>有什么想法？尽管说——</p>
    </div>
    """, unsafe_allow_html=True)

    with st.form("suggest_form"):
        content = st.text_area("你的建议", height=200, placeholder="反馈问题、提出需求、随便聊聊...")
        submitted = st.form_submit_button("提交", use_container_width=True, type="primary")
        if submitted and (content or "").strip():
            init_memory_db()
            conn = sqlite3.connect(MEMORY_DB)
            conn.execute("INSERT INTO suggestions (username, content) VALUES (?, ?)",
                        (st.session_state.get("username", ""), content.strip()))
            conn.commit()
            conn.close()
            st.success("收到！")
            log_visit("提建议", content[:50])

    st.stop()

# ==================== 英语专家 ====================
if st.session_state.page == "english":
    if st.button("← 返回首页"):
        st.session_state.page = "hub"
        st.rerun()
    st.markdown("""
    <div class="main-title" style="text-align:left;padding:1.2rem 1.8rem;">
        <div style="display:flex;align-items:center;gap:14px;">
            <div style="width:42px;height:42px;border-radius:12px;background:linear-gradient(135deg,#059669,#10b981);display:flex;align-items:center;justify-content:center;color:#fff;flex-shrink:0;box-shadow:0 4px 12px rgba(5,150,105,0.3);"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" width="24" height="24"><path d="M4 19.5A2.5 2.5 0 0 1 6.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 0 1 4 19.5v-15A2.5 2.5 0 0 1 6.5 2z"/><line x1="8" y1="7" x2="16" y2="7"/><line x1="8" y1="11" x2="14" y2="11"/></svg></div>
            <div>
                <h1 style="margin:0;font-size:1.4rem!important;">考研英语专家</h1>
                <p style="margin:2px 0 0;font-size:0.82rem;opacity:0.75;">作文批改 · 长难句解析 · 翻译练习 · 单词记忆</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    tab_essay, tab_sentence, tab_translate, tab_vocab = st.tabs([
        "作文批改", "长难句解析", "翻译与新题型", "单词记忆"
    ])

    # ── 作文批改 ──
    with tab_essay:
        st.subheader("语法纠错 · 句子升级 · 三维评分")

        col_exam, col_part = st.columns(2)
        with col_exam:
            exam_type = st.radio("考试类型", ["英语一", "英语二"], horizontal=True, key="essay_exam")
        with col_part:
            part_type = st.radio("Part 类型", ["Part A 应用文", "Part B 大作文"], horizontal=True, key="essay_part")

        max_score = 10 if part_type == "Part A 应用文" else (20 if exam_type == "英语一" else 15)

        # 历年真题库
        essay_topics_data = {}
        topics_file = Path("data/essay_topics.json")
        if topics_file.exists():
            try:
                essay_topics_data = json.loads(topics_file.read_text(encoding="utf-8"))
            except:
                pass

        if essay_topics_data:
            with st.expander("历年真题 (2016-2025)", expanded=False):
                years = sorted([y for y in essay_topics_data.get(exam_type, {}).keys() if int(y) >= 2016], reverse=True)
                if years:
                    selected_year = st.selectbox("选择年份", years, key="essay_year")
                    topic_data = essay_topics_data.get(exam_type, {}).get(selected_year, {})
                    
                    # 根据 Part 类型选择对应的图片
                    if "Part A" in part_type:
                        img_path = topic_data.get("partA_image", "")
                        part_label = "应用文"
                    else:
                        img_path = topic_data.get("partB_image", "")
                        part_label = "大作文"
                    
                    if img_path:
                        img_file = Path(img_path)
                        if img_file.exists():
                            st.image(str(img_file), caption=f"{selected_year}年 {exam_type} {part_label}", use_container_width=True)
                        else:
                            st.warning(f"图片未找到，请先截图放入 data/英语真题图片/ 目录")
                    else:
                        st.warning(f"该年份暂未收录截图")
                    
                    if st.button("使用此题目", key="use_topic"):
                        st.session_state._essay_topic = f"请参考上方 {selected_year} 年 {exam_type} {part_label}图片"
                        st.rerun()

        essay_topic = st.text_input("作文题目（选填）", value=st.session_state.get("_essay_topic", ""),
            placeholder="例如：Write a letter to... / The chart shows...")

        # 图片上传 → OCR 识别
        uploaded_img = st.file_uploader("上传手写作文照片（选填）", type=["png", "jpg", "jpeg"], key="essay_img")
        ocr_text = st.session_state.get("_essay_ocr_text", "")

        if uploaded_img is not None:
            img_data = base64.b64encode(uploaded_img.getvalue()).decode()
            if st.button("识别照片文字", use_container_width=True):
                with st.spinner("OCR 识别中..."):
                    try:
                        ocr_prompt = "请识别这张照片中的英语作文内容，只输出英文文本，保持原文格式和段落分隔。"
                        data = {"model": "mimo-v2.5", "messages": [
                            {"role": "user", "content": [
                                {"type": "text", "text": ocr_prompt},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_data}"}}
                            ]}
                        ], "max_tokens": 2000, "temperature": 0}
                        req = urllib.request.Request(API_BASE + "/chat/completions",
                            data=json.dumps(data).encode("utf-8"),
                            headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                            method="POST")
                        with urllib.request.urlopen(req, timeout=60) as resp:
                            ocr_result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                        st.session_state._essay_ocr_text = ocr_result
                        st.rerun()
                    except Exception as e:
                        st.error(f"识别失败: {e}")

        # OCR 结果可编辑区域
        if ocr_text:
            st.info("以下是 OCR 识别结果，请修正识别错误后点击「确认并批改」")
            edited_text = st.text_area("识别结果（可编辑）", value=ocr_text, height=250, key="essay_ocr_edit")
            col_confirm, col_redo = st.columns([3, 1])
            with col_confirm:
                confirm_ocr = st.button("确认并批改", use_container_width=True, type="primary", key="essay_confirm")
            with col_redo:
                if st.button("重新识别", key="essay_redo"):
                    st.session_state._essay_ocr_text = ""
                    st.rerun()
        else:
            edited_text = st.text_area("作文内容", height=250, placeholder="粘贴你的英语作文...", key="essay_manual")
            confirm_ocr = False

        # 批改逻辑
        should_grade = confirm_ocr and edited_text.strip()
        if not ocr_text:
            with st.form("essay_form"):
                essay_submitted = st.form_submit_button("批改作文", use_container_width=True, type="primary")
            should_grade = essay_submitted and edited_text.strip()

        if should_grade:
            # 构建动态提示词
            part_desc = "应用文（约100词，书信/通知/邮件）" if "Part A" in part_type else "大作文"
            if "Part A" in part_type:
                task_desc = "检查格式是否正确（书信格式、通知格式等），内容是否覆盖所有要点"
            elif exam_type == "英语一":
                task_desc = "图画作文，3步：描述图画 → 解读寓意 → 给出评论（160-200词，满分20分）"
            else:
                task_desc = "图表作文，2步：描述图表数据 → 给出评论（≥150词，满分15分）"

            if "Part A" in part_type:
                # 小作文批改提示词（用户提供的详细版）
                prompt = f"""你是考研英语小作文阅卷AI，经验丰富、洞察深刻。你的点评风格专业严谨、一针见血，同时又极具建设性和鼓励性。

## 评分标准（满分10分）
考生需根据给出的提示信息，写一篇100词左右的应用文（信函、通知等）。

### 六档评分
- 第一档 (9-10分) 很好地完成了任务：包含所有内容要点。运用丰富的语法结构和词汇，语言自然流畅，语法错误极少。有效地采用了多种衔接方法，文字连贯，层次清晰。格式与语域恰当贴切。
- 第二档 (7-8分) 较好地完成了任务：包含所有内容要点，允许漏掉1-2个次重点。使用较丰富的语法结构和词汇，只有在试图使用较复杂结构或较高级词汇时才有个别语法错误。采用了适当的衔接手法，层次清晰，组织较严密。格式与语域较恰当。
- 第三档 (5-6分) 基本完成了任务：虽漏掉一些内容，但包含多数内容要点。应用的语法结构和词汇能满足任务的需求，有一些语法及词汇错误，但不影响理解。采用了简单的衔接手法，内容较连贯，层次较清晰。格式和语域基本合理。
- 第四档 (3-4分) 未能按要求完成任务：漏掉或未能有效阐述一些内容要点，写了一些无关内容。语法结构单调、词汇有限，有较多语法结构及词汇方面的错误，影响了对写作内容的理解。未采用恰当的衔接手法，内容缺少连贯性。格式和语域不恰当。
- 第五档 (1-2分) 未完成任务：明显遗漏主要内容，且有许多不相关的内容。语法和词汇单调、重复，语言错误多，严重影响理解。无衔接，缺少组织、分段。无格式和语域概念。
- 零分档 (0分)：所传达的信息或所使用语言太少，内容与要求无关或无法辨认。

### 分项评分
- 内容完整性（3分）：是否覆盖所有内容要点
- 语言准确性（3分）：语法、词汇、拼写
- 结构与格式（4分）：格式是否规范，衔接是否流畅

## 输出格式

**得分**
[总分/10分]
属于第N档的作文

**分项评估**
内容完整性：[X分/3分]……
语言准确性：[X分/3分]……
结构与格式：[X分/4分]……

**点评 (Comments)**
优点：总结文章最突出的1-2个优点，尤其肯定其任务完成度和格式规范性
待改进处：以积极、引导的口吻，指出文章在语气得体性、语言精炼度或细节覆盖上最值得提升的1-2个问题
具体建议：提供具体的、可操作的修改建议，基于原句优化，不超过4000词汇量

**语法错误检测**
逐句检查作文，找出所有语法错误，每个错误标注：[错误] 原句 → [修正] 正确写法

**句子升级建议**
对用户已写的句子，给出更高级、优雅的平替版本：[原句] → [升级]（基于原句优化，不超过4000词汇量）

**修改后版本**
给出完整的改进后作文。

---
作文题目：{essay_topic if essay_topic else '未指定'}
作文内容：
{edited_text}"""
            else:
                # 大作文批改提示词（用户提供的详细版）
                prompt = f"""你是考研英语大作文阅卷AI，经验丰富、洞察深刻。你的点评风格专业严谨、一针见血，同时又极具建设性和鼓励性。你善于发现学生作文中的闪光点并予以肯定，对于不足之处，则以启发式、引导式的口吻提出具体的改进方案。

## 评分标准（满分20分）
考生需根据给出的提示信息（文字、图画、图表等），写一篇160-200词的短文。

### 五档评分
- 第五档 (17-20分) 优秀：很好地完成了任务，包含并有效阐述所有内容要点。使用了丰富的语法结构和词汇，错误极少。有效使用了多种衔接手段，内容连贯、流畅，层次清晰。文体格式和语体恰当贴切。
- 第四档 (13-16分) 良好：较好地完成了任务，包含所有内容要点，少数要点阐述不够充分。使用了较丰富的语法结构和词汇，仅在尝试复杂结构/词汇时有个别错误。比较有效地使用了一些衔接手段，内容较连贯，层次较清晰。
- 第三档 (9-12分) 合格：基本完成了任务，虽漏掉一些内容，但包含多数要点。语法结构和词汇基本满足需求，存在一些错误，但基本不影响理解。使用了简单的衔接手段，内容基本连贯，层次基本清晰。
- 第二档 (5-8分) 较差：未能按要求完成任务，漏掉或未有效阐述要点，有无关内容。语法结构单调，词汇有限，存在较多错误，影响理解。缺乏必要的衔接，内容不连贯。
- 第一档 (1-4分) 很差：明显遗漏主要内容，有大量不相关内容。语法结构很单调，词汇很有限，语言错误很多，内容很难理解。

### 分项评分（各5分，共20分）
- 内容：是否覆盖并有效阐述所有任务点（描述、寓意、评论）
- 语言：语法结构和词汇是否丰富多样，语言表达是否准确、规范
- 结构：结构是否合理，层次是否清晰，衔接是否流畅
- 语体：文体格式和语体是否恰当

### 特别说明
- 引用扣分：使用提示语中的部分或整个语句，将被酌情扣分
- 词数要求：不符合160-200词的要求将酌情扣分
- 拼写与标点：视为语言准确性的一个方面，视其对交际的影响程度予以考虑

## 输出格式

**得分**
[总分/20分]
属于第N档的作文

**分项评估**
内容：[X分/5分]……
语言：[X分/5分]……
结构：[X分/5分]……
语体：[X分/5分]……

**点评 (Comments)**
优点：总结文章最突出的1-2个优点，尤其肯定其思路与结构
待改进处：以积极、引导的口吻，指出文章在内容、结构或语言上最值得提升的1-2个问题
具体建议：提供具体的、可操作的修改建议，基于原句优化，不超过4000词汇量

**语法错误检测**
逐句检查作文，找出所有语法错误：
- 时态错误、主谓一致、冠词使用、介词搭配
- 从句结构、虚拟语气、非谓语动词
- 拼写错误、标点错误
每个错误标注：[错误] 原句 → [修正] 正确写法

**句子升级建议**
对用户已写的句子，给出更高级、优雅的平替版本：
[原句] 用户写的句子 → [升级] 更优雅的版本（基于原句优化，不超过4000词汇量）

**修改后版本**
给出完整的改进后作文。

---
考试类型：{exam_type}
题型：{part_desc}
任务要求：{task_desc}
作文题目：{essay_topic if essay_topic else '未指定'}
作文内容：
{edited_text}"""
            with st.spinner("批改中..."):
                try:
                    data = {"model": "mimo-v2.5", "messages": [
                        {"role": "user", "content": prompt}
                    ], "max_tokens": 3000, "temperature": 0.3}
                    req = urllib.request.Request(API_BASE + "/chat/completions",
                        data=json.dumps(data).encode("utf-8"),
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                        method="POST")
                    with urllib.request.urlopen(req, timeout=120) as resp:
                        result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                    st.markdown("---")
                    st.markdown(_escape_md(_collapse_math(_fix_latex(result))))
                    st.html("<script>if(typeof renderMathInElement!=='undefined'){renderMathInElement(document.body,{delimiters:[{left:'$$',right:'$$',display:true},{left:'$',right:'$',display:false}],throwOnError:!1})}</script>")
                    log_visit("英语作文批改", f"{exam_type} {part_type}: {essay_topic or edited_text[:30]}")
                except Exception as e:
                    st.error(f"批改失败: {e}")

    # ── 长难句解析 ──
    with tab_sentence:
        st.subheader("主干提取 · 修饰分析 · 翻译")
        with st.form("sentence_form"):
            sentence_text = st.text_area("输入英语长难句", height=120,
                placeholder="例如: The fact that the defendant had previously been convicted of a similar offense, which was not disclosed to the jury, raises serious questions about the fairness of the trial.")
            sentence_submitted = st.form_submit_button("解析", use_container_width=True, type="primary")

        if sentence_submitted and sentence_text.strip():
            prompt = f"""你是考研英语长难句解析专家。按以下步骤解析：

1. 找主干：主语 + 谓语 + 宾语
2. 标修饰：定语从句、状语从句、插入语、同位语等
3. 理逻辑：因果、转折、并列等逻辑关系
4. 译全文：准确中文翻译
5. 语法点：涉及的语法知识点

输出格式：
[主干] 主语 + 谓语 + 宾语
[修饰] 各修饰成分分析
[逻辑] 句子逻辑关系
[翻译] 中文翻译
[语法点] 涉及的语法知识点

待解析句子：
{sentence_text}"""
            with st.spinner("解析中..."):
                try:
                    data = {"model": "mimo-v2.5", "messages": [
                        {"role": "user", "content": prompt}
                    ], "max_tokens": 1500, "temperature": 0.3}
                    req = urllib.request.Request(API_BASE + "/chat/completions",
                        data=json.dumps(data).encode("utf-8"),
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                        method="POST")
                    with urllib.request.urlopen(req, timeout=90) as resp:
                        result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                    st.markdown(_escape_md(_collapse_math(_fix_latex(result))))
                    log_visit("长难句解析", sentence_text[:40])
                except Exception as e:
                    st.error(f"解析失败: {e}")

    # ── 翻译与新题型 ──
    with tab_translate:
        st.subheader("翻译 · 7选5 · 排序 · 小标题匹配")
        translate_mode = st.radio("练习类型",
            ["英译中", "中译英", "7选5", "排序题", "小标题匹配"],
            horizontal=True, key="translate_mode")
        with st.form("translate_form"):
            translate_text = st.text_area("输入文本", height=200,
                placeholder="粘贴需要翻译或练习的文本...")
            translate_submitted = st.form_submit_button("开始练习", use_container_width=True, type="primary")

        if translate_submitted and translate_text.strip():
            mode_prompts = {
                "英译中": "将以下英文翻译为中文，给出准确译文 + 关键词汇解析 + 语法结构 + 翻译技巧",
                "中译英": "将以下中文翻译为英文，给出准确译文 + 关键词汇 + 语法结构 + 写作技巧",
                "7选5": "你是考研英语新题型专家（7选5）。分析上下文逻辑，从选项中选出最佳答案填入空白处",
                "排序题": "你是考研英语新题型专家（排序题）。找段落间的衔接词和逻辑关系，给出正确排序",
                "小标题匹配": "你是考研英语新题型专家（小标题匹配）。提炼每段主旨，匹配最佳标题",
            }
            prompt = f"""{mode_prompts.get(translate_mode, '')}

输出格式：
[答案/翻译] 准确译文或答案
[关键词] 重要词汇解析
[语法] 涉及的语法结构
[技巧] 使用的解题技巧

待处理文本：
{translate_text}"""
            with st.spinner("处理中..."):
                try:
                    data = {"model": "mimo-v2.5", "messages": [
                        {"role": "user", "content": prompt}
                    ], "max_tokens": 1500, "temperature": 0.3}
                    req = urllib.request.Request(API_BASE + "/chat/completions",
                        data=json.dumps(data).encode("utf-8"),
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                        method="POST")
                    with urllib.request.urlopen(req, timeout=90) as resp:
                        result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                    st.markdown(_escape_md(_collapse_math(_fix_latex(result))))
                    log_visit("英语翻译", f"{translate_mode}: {translate_text[:30]}")
                except Exception as e:
                    st.error(f"处理失败: {e}")

    # ── 单词记忆 ──
    with tab_vocab:
        st.subheader("词根词缀 · 联想记忆 · 搭配例句")
        with st.form("vocab_form"):
            vocab_input = st.text_input("输入单词或主题",
                placeholder="例如：innovation / 考研高频词 / 经济类词汇")
            vocab_submitted = st.form_submit_button("查询", use_container_width=True, type="primary")

        if vocab_submitted and vocab_input.strip():
            prompt = f"""你是考研英语单词记忆专家。针对以下单词或主题，提供：

1. 词根词缀分析
2. 联想记忆法
3. 同义词/反义词
4. 常考搭配
5. 经典例句（考研真题风格）

输出格式：
[词根] 词根词缀拆解
[联想] 记忆联想
[同义] 同义词 / [反义] 反义词
[搭配] 常考搭配
[例句] 经典例句

单词/主题：{vocab_input}"""
            with st.spinner("查询中..."):
                try:
                    data = {"model": "mimo-v2.5", "messages": [
                        {"role": "user", "content": prompt}
                    ], "max_tokens": 1000, "temperature": 0.3}
                    req = urllib.request.Request(API_BASE + "/chat/completions",
                        data=json.dumps(data).encode("utf-8"),
                        headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                        method="POST")
                    with urllib.request.urlopen(req, timeout=60) as resp:
                        result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                    st.markdown(_escape_md(_collapse_math(_fix_latex(result))))
                    log_visit("英语单词记忆", vocab_input[:30])
                except Exception as e:
                    st.error(f"查询失败: {e}")

    st.stop()

# ==================== 打卡与督学 ====================
if st.session_state.page == "checkin":
    if st.button("← 返回首页"):
        st.session_state.page = "hub"
        st.rerun()
    st.markdown("""
    <div class="main-title" style="text-align:left;padding:1.2rem 1.8rem;">
        <div style="display:flex;align-items:center;gap:14px;">
            <div style="width:42px;height:42px;border-radius:12px;background:linear-gradient(135deg,#16a34a,#22c55e);display:flex;align-items:center;justify-content:center;color:#fff;flex-shrink:0;box-shadow:0 4px 12px rgba(22,163,74,0.3);"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" width="24" height="24"><path d="M9 11l3 3L22 4"/><path d="M21 12v7a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h11"/></svg></div>
            <div>
                <h1 style="margin:0;font-size:1.4rem!important;">打卡与督学</h1>
                <p style="margin:2px 0 0;font-size:0.82rem;opacity:0.75;">每日打卡 · 学习计划 · 学习日记 · 番茄计时</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    checkin_user_id = st.session_state.get("user_id")
    checkin_username = st.session_state.get("username", "用户")

    st.caption(f"当前用户：{checkin_username}（user_id={checkin_user_id}）")

    # 概览指标
    ck_c1, ck_c2, ck_c3, ck_c4 = st.columns(4)
    with ck_c1:
        st.metric("连续打卡", f"{get_consecutive_days(checkin_user_id)} 天")
    with ck_c2:
        st.metric("今日学习", f"{get_today_duration(checkin_user_id)} 分钟")
    with ck_c3:
        st.metric("计划完成率", f"{get_checkin_plan_progress(checkin_user_id)}%")
    with ck_c4:
        st.metric("今日心情", get_today_mood(checkin_user_id))

    # 智能提醒
    for level, msg in check_checkin_reminders(checkin_user_id):
        if level == "warning":
            st.warning(msg)
        elif level == "error":
            st.error(msg)
        elif level == "success":
            st.success(msg)
        else:
            st.info(msg)

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["每日打卡", "🌙 学习日记", "学习计划", "番茄计时", "学习画像"])

    # ── 每日打卡 ──
    with tab1:
        st.subheader("每日打卡")
        current_ck = get_today_checkin(checkin_user_id)
        if current_ck:
            st.success("今天已打卡，可重新提交覆盖今日记录。")
            ck_col1, ck_col2, ck_col3 = st.columns(3)
            with ck_col1:
                st.write(f"科目：{current_ck['subject'] or '-'}")
            with ck_col2:
                st.write(f"时长：{current_ck['duration_minutes'] or 0} 分钟")
            with ck_col3:
                st.write(f"完成度：{current_ck['completion_rate'] or 0:.0f}%")
            st.write(f"心情：{current_ck['mood'] or '-'}")
            if current_ck["notes"]:
                st.caption(f"备注：{current_ck['notes']}")

        with st.form("checkin_form"):
            SUBJECTS = ["数学", "英语", "政治", "综合"]
            MOODS = ["😊 开心", "😐 一般", "😢 疲惫", "😤 焦虑", "🤯 崩溃"]
            subject_index = SUBJECTS.index(current_ck["subject"]) if current_ck and current_ck["subject"] in SUBJECTS else 0
            mood_index = MOODS.index(current_ck["mood"]) if current_ck and current_ck["mood"] in MOODS else 1
            subject = st.selectbox("学习科目", SUBJECTS, index=subject_index)
            duration = st.slider("学习时长（分钟）", 0, 480,
                int(current_ck["duration_minutes"]) if current_ck else 120, step=15)
            completion = st.slider("完成度", 0, 100,
                int(current_ck["completion_rate"]) if current_ck else 80, step=5)
            mood = st.radio("心情", MOODS, index=mood_index, horizontal=True)
            notes = st.text_input("备注（选填）",
                value=current_ck["notes"] if current_ck and current_ck["notes"] else "")
            submitted = st.form_submit_button("提交打卡", use_container_width=True, type="primary")

        if submitted:
            save_checkin(checkin_user_id, checkin_today_str(), subject, duration, completion, mood, notes)
            st.success("打卡已保存。")
            st.rerun()

        # 打卡日历
        st.markdown("#### 最近 30 天")
        rows_30 = get_recent_checkins(checkin_user_id, 30)
        by_date_30 = {row["checkin_date"]: row for row in rows_30}
        start_date = date.today() - timedelta(days=29)
        cal_cells = []
        for idx in range(30):
            day = start_date + timedelta(days=idx)
            key = day.strftime("%Y-%m-%d")
            row = by_date_30.get(key)
            label = day.strftime("%m-%d")
            if row:
                comp = float(row["completion_rate"] or 0)
                marker = "🟢" if comp >= 80 else "🟡" if comp >= 60 else "🔴"
                cal_cells.append(f'<div class="cal-cell">{marker}<br><small>{label}</small></div>')
            else:
                cal_cells.append(f'<div class="cal-cell">⚪<br><small>{label}</small></div>')
        st.markdown(f'<div class="cal-grid">{"".join(cal_cells)}</div>', unsafe_allow_html=True)

        # 学习趋势
        st.markdown("#### 学习趋势")
        chart_data = []
        for idx in range(30):
            day = start_date + timedelta(days=idx)
            key = day.strftime("%Y-%m-%d")
            row = by_date_30.get(key)
            chart_data.append({
                "日期": key[5:],
                "学习时长": int(row["duration_minutes"] or 0) if row else 0,
                "完成度": float(row["completion_rate"] or 0) if row else 0,
            })
        st.line_chart(chart_data, x="日期", y=["学习时长", "完成度"])

    # ── 晚间复盘（日记形式） ──
    with tab2:
        st.subheader("🌙 今日日记")
        current_rv = get_today_review(checkin_user_id)
        
        with st.form("diary_form"):
            diary_default = ""
            if current_rv and current_rv["diary_content"]:
                diary_default = current_rv["diary_content"]
            elif current_rv and current_rv["what_learned"]:
                # 兼容旧格式
                diary_default = current_rv["what_learned"]
                if current_rv["what_difficult"]:
                    diary_default += "\n\n遇到的困难：" + current_rv["what_difficult"]
                if current_rv["what_improve"]:
                    diary_default += "\n\n改进计划：" + current_rv["what_improve"]
            
            diary_content = st.text_area(
                "写下今天的学习心得...",
                value=diary_default,
                height=200,
                placeholder="今天学了什么？有什么收获？遇到了什么困难？明天打算怎么做？",
                key="diary_input")
            submitted_diary = st.form_submit_button("保存日记", use_container_width=True, type="primary")

        if submitted_diary and diary_content.strip():
            save_review(checkin_user_id, checkin_today_str(), diary_content.strip())
            st.success("日记已保存。")
            st.rerun()

        # ── 学习足迹（日记+计划穿插） ──
        st.markdown("---")
        st.subheader("学习足迹")
        
        timeline = get_timeline(checkin_user_id, days=14)
        if timeline:
            current_date = ""
            for item in timeline:
                item_date = str(item["date"] or "")
                if item_date != current_date:
                    current_date = item_date
                    st.markdown(f"**{item_date}**")
                
                if item["type"] == "diary":
                    content = str(item["content"] or "")
                    st.markdown(f"🌙 {content}")
                elif item["type"] == "plan_create":
                    content = str(item["content"] or "")
                    st.markdown(f"{content}")
        else:
            st.info("还没有记录，开始你的第一次打卡吧！")

    # ── 学习计划 ──
    with tab3:
        st.subheader("学习计划")

        # 获取用户画像
        profile = get_user_profile(checkin_user_id)
        target_schools = _display_target_schools(profile)
        target_major = profile.get("target_major") or "未设置"
        undergraduate_major = profile.get("undergraduate_major") or "未设置"
        undergraduate_level = profile.get("undergraduate_level") or "未设置"
        weak_subjects = _safe_json_loads(profile.get("weak_subjects"))
        strong_subjects = _safe_json_loads(profile.get("strong_subjects"))
        is_cross_major = profile.get("is_cross_major") or "否"
        anxiety = profile.get("anxiety_level") or 3

        # 自动判断当前阶段
        current_month = datetime.now().month
        if 3 <= current_month <= 6:
            current_phase = "基础阶段"
        elif 7 <= current_month <= 9:
            current_phase = "强化阶段"
        elif 10 <= current_month <= 11:
            current_phase = "提升阶段"
        else:
            current_phase = "冲刺阶段"

        # 阶段说明卡片
        st.markdown(f"""
        <div style="background:var(--bg-elevated, #f8fafc); border-radius:12px; padding:clamp(12px,2vw,20px); margin-bottom:20px; border:1px solid var(--border, #e2e8f0);">
            <h3 style="margin:0 0 12px 0; color:var(--text-main, #1e293b);">当前阶段：{current_phase}</h3>
            <table style="width:100%; border-collapse:collapse; font-size:clamp(12px,1.5vw,14px);">
                <tr style="background:#4f46e5; color:#fff;">
                    <th style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px); text-align:left;">阶段</th>
                    <th style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px); text-align:left;">时间</th>
                    <th style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px); text-align:left;">核心任务</th>
                </tr>
                <tr style="background:{'#eef2ff' if current_phase=='基础阶段' else '#fff'}">
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">基础阶段</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">3-6月</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">全面打基础、吃透教材和基础题</td>
                </tr>
                <tr style="background:{'#eef2ff' if current_phase=='强化阶段' else '#f8fafc'}">
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">强化阶段</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">7-9月</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">专项突破、大量刷题、建立做题体系</td>
                </tr>
                <tr style="background:{'#eef2ff' if current_phase=='提升阶段' else '#fff'}">
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">提升阶段</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">10-11月</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">真题实战、查漏补缺、模考检验</td>
                </tr>
                <tr style="background:{'#eef2ff' if current_phase=='冲刺阶段' else '#f8fafc'}">
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">冲刺阶段</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">12月</td>
                    <td style="padding:clamp(4px,1vw,8px) clamp(6px,1.5vw,12px);">高频考点押题、心理调整、保持手感</td>
                </tr>
            </table>
        </div>
        """, unsafe_allow_html=True)

        # 生成学习计划表单
        with st.form("generate_plan_form"):
            col1, col2 = st.columns(2)
            with col1:
                plan_phase = st.selectbox("学习阶段", ["基础阶段", "强化阶段", "提升阶段", "冲刺阶段"],
                                          index=["基础阶段", "强化阶段", "提升阶段", "冲刺阶段"].index(current_phase))
            with col2:
                plan_subjects = st.multiselect("选择科目", ["数学", "英语", "政治", "专业课"],
                                               default=["数学", "英语"])
            daily_hours = st.slider("每日学习时长（小时）", min_value=2, max_value=12, value=6)
            submit_plan = st.form_submit_button("🚀 生成学习计划", use_container_width=True, type="primary")

        # 生成计划
        if submit_plan and plan_subjects:
            # 构建 prompt
            subjects_str = "、".join(plan_subjects)
            gap_desc = ""
            if target_schools and target_schools != "未设置":
                gap_desc = f"目标院校：{target_schools}，目标专业：{target_major}"
                if undergraduate_level and undergraduate_level != "未设置":
                    gap_desc += f"，本科院校：{undergraduate_level}"
                if is_cross_major == "是":
                    gap_desc += "，跨考生"

            prompt = f"""结合用户基础和院校差距，生成一份实用、可执行的学习计划。

## 用户画像
- {gap_desc if gap_desc else '目标院校：未设置'}
- 本专业：{undergraduate_major}
- 强科：{', '.join(strong_subjects) if strong_subjects else '未设置'}
- 弱科：{', '.join(weak_subjects) if weak_subjects else '未设置'}
- 焦虑程度：{anxiety}/5
- 当前阶段：{plan_phase}
- 学习科目：{subjects_str}
- 每日学习时长：{daily_hours}小时

## 输出要求
请按科目分段，每段给出具体的复习安排和建议。
要求：
1. 直接给出复习安排，不写鼓励、共情或称呼（如"亲爱的同学""孩子""老师"）
2. 语气朴素、理性，像说明文档而非个人书信
3. 结合用户的院校差距，说明该阶段的学习重点
4. 结合强弱科，给出针对性建议
5. 可用短段落 + 要点形式，总字数约 400-600 字
"""
            with st.spinner("正在生成学习计划..."):
                try:
                    result = call_llm_api(prompt, model="mimo-v2.5")
                    st.session_state._plan_result = result
                    st.session_state._plan_phase = plan_phase
                    st.session_state._plan_subjects = plan_subjects
                    st.rerun()
                except Exception as e:
                    st.error(f"生成失败：{e}")

        # 显示生成的计划
        if st.session_state.get("_plan_result"):
            st.markdown("---")
            st.markdown(f"### 学习计划 - {st.session_state._plan_phase}")
            st.markdown(st.session_state._plan_result)
            if st.button("💾 保存此计划"):
                save_checkin_plan(
                    checkin_user_id,
                    st.session_state._plan_phase + "计划",
                    date.today().strftime("%Y-%m-%d"),
                    st.session_state._plan_result
                )
                st.success("计划已保存！")
                st.rerun()

        # 已保存的计划
        plans = get_checkin_plans(checkin_user_id)
        if plans:
            st.markdown("---")
            st.subheader("已保存的计划")
            for plan in plans:
                with st.expander(f"{plan['plan_name']} - {plan['target_date']}"):
                    st.markdown(plan["tasks"] or "")
                    if st.button("🗑️ 删除", key=f"del_plan_{plan['id']}"):
                        delete_plan(checkin_user_id, plan["id"])
                        st.rerun()

    # ── 番茄计时 ──
    with tab4:
        st.subheader("番茄计时器")
        if "pomo_running" not in st.session_state:
            st.session_state.pomo_running = False
        if "pomo_start_time" not in st.session_state:
            st.session_state.pomo_start_time = None
        if "pomo_duration" not in st.session_state:
            st.session_state.pomo_duration = 25
        if "pomo_subject" not in st.session_state:
            st.session_state.pomo_subject = "数学"

        pomo_count, pomo_minutes = get_today_pomodoros(checkin_user_id)
        st.caption(f"今日已完成 {pomo_count} 个番茄，共 {pomo_minutes} 分钟")

        pc1, pc2 = st.columns(2)
        with pc1:
            pomo_subject = st.selectbox("学习科目", SUBJECTS, key="pomo_subject_select")
        with pc2:
            pomo_duration = st.selectbox("时长（分钟）", [15, 25, 45, 60], index=1, key="pomo_duration_select")

        ps1, ps2, ps3 = st.columns(3)
        with ps1:
            if st.button("开始", use_container_width=True, type="primary"):
                st.session_state.pomo_running = True
                st.session_state.pomo_start_time = time.time()
                st.session_state.pomo_duration = pomo_duration
                st.session_state.pomo_subject = pomo_subject
                st.rerun()
        with ps2:
            if st.button("暂停/重置", use_container_width=True):
                st.session_state.pomo_running = False
                st.session_state.pomo_start_time = None
                st.rerun()
        with ps3:
            if st.button("手动完成", use_container_width=True):
                save_pomodoro(checkin_user_id, pomo_subject, pomo_duration, pomo_duration, 1)
                st.session_state.pomo_running = False
                st.session_state.pomo_start_time = None
                st.success("番茄记录已保存。")
                st.rerun()

        if st.session_state.pomo_running and st.session_state.pomo_start_time:
            elapsed = int(time.time() - st.session_state.pomo_start_time)
            total = int(st.session_state.pomo_duration) * 60
            remaining = max(0, total - elapsed)
            minutes_left = remaining // 60
            seconds_left = remaining % 60

            st.markdown(f"""
            <div style="text-align:center; font-size:clamp(2rem,8vw,4rem); font-weight:700; margin:1rem 0;">
                {minutes_left:02d}:{seconds_left:02d}
            </div>
            """, unsafe_allow_html=True)
            st.progress(min(elapsed / total, 1.0))

            if remaining <= 0:
                save_pomodoro(checkin_user_id, st.session_state.pomo_subject,
                    st.session_state.pomo_duration, st.session_state.pomo_duration, 1)
                st.session_state.pomo_running = False
                st.session_state.pomo_start_time = None
                st.success("番茄完成，建议休息 5 分钟。")
                st.balloons()
            else:
                time.sleep(1)
                st.rerun()
        else:
            st.info("选择科目和时长后点击开始。")

    # ── 学习画像（问卷建档） ──
    with tab5:
        st.subheader("学习画像问卷")
        st.info("填写以下信息，系统会为你生成更精准的学习计划。")
        
        existing_profile = get_user_profile(checkin_user_id)
        
        with st.form("profile_form"):
            target_school = st.text_input("目标院校", value=existing_profile.get("target_schools") or "", placeholder="例如：清华大学、华东师范大学")
            target_major = st.text_input("目标专业", value=existing_profile.get("target_major") or "", placeholder="例如：计算机科学与技术")
            
            undergraduate_major = st.text_input("本专业", value=existing_profile.get("undergraduate_major") or "", placeholder="例如：计算机科学与技术、电气工程")
            undergraduate_level = st.selectbox("本科院校级别", ["双非", "双一流", "211", "985"],
                                               index=["双非", "双一流", "211", "985"].index(existing_profile.get("undergraduate_level") or "双非"))
            is_cross_major = st.selectbox("是否跨考", ["否", "是"],
                                          index=["否", "是"].index(existing_profile.get("is_cross_major") or "否"))
            
            st.markdown("**强科/弱科（多选）：**")
            subjects = ["数学", "英语", "政治", "专业课"]
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                strong = st.multiselect("强科", subjects, default=_safe_json_loads(existing_profile.get("strong_subjects")))
            with col_s2:
                weak = st.multiselect("弱科", subjects, default=_safe_json_loads(existing_profile.get("weak_subjects")))
            
            anxiety = st.slider("焦虑程度", 1, 5, value=int(existing_profile.get("anxiety_level") or 3))
            
            if st.form_submit_button("保存画像", use_container_width=True, type="primary"):
                save_profile_field(checkin_user_id, "target_schools", json.dumps({"冲刺": target_school}, ensure_ascii=False) if target_school else "")
                save_profile_field(checkin_user_id, "target_major", target_major)
                save_profile_field(checkin_user_id, "undergraduate_major", undergraduate_major)
                save_profile_field(checkin_user_id, "undergraduate_level", undergraduate_level)
                save_profile_field(checkin_user_id, "is_cross_major", is_cross_major)
                save_profile_field(checkin_user_id, "strong_subjects", json.dumps(strong, ensure_ascii=False))
                save_profile_field(checkin_user_id, "weak_subjects", json.dumps(weak, ensure_ascii=False))
                save_profile_field(checkin_user_id, "anxiety_level", anxiety)
                st.success("画像已保存！生成学习计划时会自动参考。")
                st.rerun()

        # 显示当前画像
        st.markdown("---")
        st.subheader("当前画像")
        if existing_profile:
            st.markdown(f"**目标院校**：{_display_target_schools(existing_profile)}")
            st.markdown(f"**目标专业**：{existing_profile.get('target_major') or '未设置'}")
            st.markdown(f"**本专业**：{existing_profile.get('undergraduate_major') or '未设置'}")
            st.markdown(f"**本科院校级别**：{existing_profile.get('undergraduate_level') or '未设置'}")
            st.markdown(f"**是否跨考**：{existing_profile.get('is_cross_major') or '否'}")
            st.markdown(f"**强科**：{', '.join(_safe_json_loads(existing_profile.get('strong_subjects'))) or '未设置'}")
            st.markdown(f"**弱科**：{', '.join(_safe_json_loads(existing_profile.get('weak_subjects'))) or '未设置'}")
            st.markdown(f"**焦虑程度**：{existing_profile.get('anxiety_level') or '未设置'}/5")
        else:
            st.info("尚未建档，请填写上方问卷。")

    st.stop()

# ==================== 考研数学问答工具 ====================

# 初始化（已登录）
corpus = load_corpus()
experience = load_agent_experience()
stats = get_memory_stats()
add_thinking(f"用户 {st.session_state.get('username','?')} 登录")

# ← 返回首页按钮
if st.button("← 返回首页", key="back_hub"):
    st.session_state.page = "hub"
    st.rerun()

# 顶部标题
st.markdown("""
<div class="main-title">
    <h1>考研学习助手</h1>
    <p>基于本地知识库的智能问答系统 | 支持自学习、遗忘曲线、经验积累</p>
</div>
""", unsafe_allow_html=True)

# 使用columns实现三栏布局
left_col, mid_col = st.columns([1, 2])

# ==================== 左侧面板 ====================
with left_col:
    st.markdown("### 👤 当前用户")
    st.markdown(f"**{st.session_state.get('username','?')}**")
    if st.button("🚪 退出登录", use_container_width=True):
        clear_login_token(st.session_state.get("user_id", 0))
        cookie_manager.delete("auth_token")
        st.session_state.logged_in = False
        st.session_state.user_id = None
        st.session_state.page = "hub"
        st.rerun()

    st.markdown("---")

    # 模型
    st.session_state.selected_model = "mimo-v2.5"
    st.caption("模型: mimo-v2.5")

    st.markdown("---")

    # Skill 技能切换
    st.markdown("### 🎯 回答方式")
    all_skills = load_all_skills()
    if all_skills:
        options = ["无 (默认)"] + list(all_skills.keys())
        labels = ["无 (默认)"] + [f"{m.get('label', n)}" for n, m in all_skills.items()]
        choice = st.selectbox("选择回答风格", range(len(options)), format_func=lambda x: labels[x])
        st.session_state.active_skills = [options[choice]] if choice > 0 else []
    else:
        st.caption("`skills/` 目录下暂无 Skill")
    st.markdown("---")

    # 系统状态
    st.markdown("### ⚙️ 系统状态")
    st.markdown(f"📁 {len(corpus)} 个文档 · {stats['total']} 知识点")
    col_r1, col_r2 = st.columns([1, 1])
    with col_r1:
        if st.button("刷新知识库", use_container_width=True):
            st.cache_data.clear()
            st.rerun()
    with col_r2:
        if st.button("🔌 重连API", use_container_width=True):
            st.cache_data.clear()
            st.rerun()

# ==================== 中间面板 ====================
with mid_col:
    with st.expander("💡 新手指南", expanded=False):
        st.markdown("本系统内置 **110 个考研数学核心知识点**，覆盖高等数学、线性代数、概率论三大模块，内容对齐 2025 年考试大纲。")
        st.markdown("**🧭 智能路由** — 数学问题检索知识库，英语/政治问题由 AI 直接回答。")
        st.markdown("**🎯 回答方式** — 侧边栏可选择 AI 的输出风格：分步解题、概念讲解、错题分析，以及纯要点、问答、纯公式等格式。")
        st.markdown("**记忆系统** — 自动追踪每个知识点的掌握程度，根据遗忘曲线推送复习内容。")

    st.markdown("### 智能问答")
    with st.form("qa_form", clear_on_submit=False):
        query = st.text_input("输入你的考研问题", placeholder="例如：什么是导数？", key="query_input")
        uploaded_img = st.file_uploader("题目截图", type=["png","jpg","jpeg"], label_visibility="collapsed")
        submitted = st.form_submit_button("提问", use_container_width=True)

    img_data = None
    if uploaded_img is not None:
        try:
            img_data = base64.b64encode(uploaded_img.getvalue()).decode()
        except:
            pass

    if submitted and (query or img_data):
        add_thinking(f"查询: {query[:30]}..." if query else "图片识别...")
        results = search_corpus(query, corpus, top_k=3) if query else []

        # 流式接收（只收集，不立即显示）
        st.markdown('<div class="qa-card">', unsafe_allow_html=True)
        st.markdown("### 💡 回答")
        thinking_placeholder = st.empty()
        thinking_placeholder.markdown("<span style='color:#4f46e5;font-weight:600;'>AI 正在思考...</span>", unsafe_allow_html=True)

        raw_full = ""
        output = None
        import time as _time

        for event in run_pipeline(query or "请识别并解答图中的数学题目", results, st.session_state.selected_model, img_data):
            if event["type"] == "token":
                raw_full += event["content"]
            elif event["type"] == "done":
                output = event["result"]

        # 流结束后，打字效果显示 [ANSWER] 部分
        thinking_placeholder.empty()
        answer_text = ""
        if output and output.get("answer"):
            answer_text = output["answer"]
        elif raw_full:
            # fallback: 从原始文本中提取
            if "[ANSWER]" in raw_full:
                answer_text = raw_full.split("[ANSWER]", 1)[1]
                if "[KNOWLEDGE]" in answer_text:
                    answer_text = answer_text.split("[KNOWLEDGE]")[0]
            else:
                answer_text = raw_full

        if answer_text.strip():
            answer_placeholder = st.empty()
            _typing_display(answer_placeholder, _escape_md(_collapse_math(_fix_latex(answer_text.strip()))), delay=0.02)
            _katex_refresh()
            # 保存 answer_text，供 rerun 时重新渲染（如切标签页后 WebSocket 断连）
            st.session_state._last_answer_text = answer_text.strip()
        st.markdown('</div>', unsafe_allow_html=True)
        # 诊断：GLM 原始输出
        if output.get("_raw_debug"):
            with st.expander("🔧 GLM原始输出（诊断）"):
                st.code(output["_raw_debug"])
        add_thinking(f"回答完成")
        log_visit("提问", f"{query[:50]}")

        # 知识点归纳（用 ALIAS 表归一化为实际文件名）
        if output.get("knowledge"):
            validated = []
            for kid in output["knowledge"]:
                match = smart_match_knowledge(kid.strip())
                validated.append(match[0] if len(match) > 0 else kid.strip())
            validated = list(dict.fromkeys(validated))
            for kid in validated:
                update_memory(kid, False, error_type="自动归纳")
            add_thinking(f"自动归纳知识点: {validated}")
            st.session_state._matched_knowledge = validated

        # 参考来源
        if results:
            st.markdown("### 使用的参考资料")
            ref_html = ""
            for r in results:
                ref_html += f"<span class='ref-tag'>📄 {_clean_knowledge_name(r['id'])} ×{r['score']}</span>"
            st.markdown(ref_html, unsafe_allow_html=True)
        else:
            st.caption("📡 回答来自LLM自身知识")

        # 保存上下文到 session_state（供后续按钮使用）
        st.session_state._last_output = output
        st.session_state._last_query = query
        st.session_state._last_results = results

    elif st.session_state.get("_last_answer_text"):
        # 重新渲染上次回答（如切标签页 WebSocket 断连后 rerun，submitted=False 但答案仍在）
        st.markdown('<div class="qa-card">', unsafe_allow_html=True)
        st.markdown("### 💡 回答")
        answer_placeholder = st.empty()
        answer_placeholder.markdown(_escape_md(_collapse_math(_fix_latex(st.session_state._last_answer_text))))
        _katex_refresh()
        st.markdown('</div>', unsafe_allow_html=True)
        last_results = st.session_state.get("_last_results", [])
        if last_results:
            st.markdown("### 使用的参考资料")
            ref_html = ""
            for r in last_results:
                ref_html += f"<span class='ref-tag'>📄 {_clean_knowledge_name(r['id'])} ×{r['score']}</span>"
            st.markdown(ref_html, unsafe_allow_html=True)
        else:
            st.caption("📡 回答来自LLM自身知识")

    # 出2道练习题按钮 + 评价按钮（在 mid_col 内，不在 if submitted 内）
    # 显示上一次操作的反馈消息
    act_msg = st.session_state.pop("_action_msg", "")
    act_diag = st.session_state.pop("_action_diag", "")
    if act_msg:
        st.success(f"{act_msg}  ·  {act_diag}" if act_diag else act_msg)

    # 出题结果显示
    btn_quiz = st.session_state.pop("_btn_quiz", None)
    if btn_quiz and btn_quiz.get("success"):
        st.markdown("#### 练习题")
        render_qa_cards(btn_quiz['questions'], columns=2, typing=True)

    last_output = st.session_state.get("_last_output")
    if last_output:
        # 出2道练习题按钮
        if st.button("🎲 生成复习题", use_container_width=True):
            last_query = st.session_state.get("_last_query", "")
            matched = st.session_state.get("_matched_knowledge") or smart_match_knowledge(last_query)
            if matched:
                progress_bar = st.progress(0, text="🎲 开始生成题目...")
                progress_bar.progress(30, text="正在分析知识点...")
                st.session_state._btn_quiz = generate_review_questions([{"knowledge_id": m} for m in matched[:2]])
                progress_bar.progress(100, text="题目生成完成！")
                st.rerun()
            else:
                st.info("未匹配到相关知识点")

        st.markdown("### 这个回答对你有帮助吗？")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("掌握了", use_container_width=True):
                last_results = st.session_state.get("_last_results", [])
                last_query = st.session_state.get("_last_query", "")
                if last_results:
                    for r in last_results:
                        update_memory(r['id'], True)
                else:
                    matched = smart_match_knowledge(last_query)
                    if matched:
                        for kid in matched:
                            update_memory(kid, True)
                        add_thinking(f"智能匹配知识点: {matched}")
                add_thinking("用户点击: 掌握了")
                st.session_state._action_msg = "已记录为掌握！"
                st.rerun()
        with col2:
            if st.button("加入复习库", use_container_width=True):
                last_query = st.session_state.get("_last_query", "")
                matched = st.session_state.get("_matched_knowledge") or smart_match_knowledge(last_query)
                if matched:
                    for kid in matched:
                        update_memory(kid, False, error_type="用户标记")
                    # 立即验证 DB 写入
                    vconn = sqlite3.connect(MEMORY_DB)
                    vc = vconn.cursor()
                    uid = st.session_state.get("user_id", 1)
                    vc.execute("SELECT knowledge_id, status FROM knowledge_mastery WHERE knowledge_id=? AND user_id=?", (matched[0], uid))
                    verify = vc.fetchone()
                    vconn.close()
                    st.session_state._action_msg = f"已加入复习库 ({len(matched)}个知识点)"
                    st.session_state._action_diag = f"匹配: {matched} | DB验证: {verify}"
                else:
                    st.session_state._action_msg = "未匹配到具体知识点"
                log_visit("加入复习库", last_query[:50] if last_query else "")
                st.rerun()

# ==================== 底部Tab ====================
st.markdown("---")

# 考纲分类选择
if "math_type" not in st.session_state:
    st.session_state.math_type = "全部"
math_type = st.radio("考纲分类", ["全部", "数学一专属", "数学三专属"], horizontal=True, key="math_type_radio", label_visibility="collapsed")
st.session_state.math_type = math_type

filtered_corpus = _filter_corpus(corpus, math_type)
filtered_corpus_ids = {d["id"] for d in filtered_corpus}
total_all = len(corpus)
total_filtered = len(filtered_corpus)

tab1, tab2, tab4, tab3 = st.tabs(["知识库", "复习挑战", "🎓 费曼学习法", "记忆系统"])

with tab1:
    label = f"（{math_type}）" if math_type != "全部" else ""
    st.subheader(f"知识库{label} {total_filtered}/{total_all} 个文档")
    search_kw = st.text_input("搜索知识库", label_visibility="collapsed", placeholder="搜索...")
    if search_kw:
        results = search_corpus(search_kw, filtered_corpus, top_k=20)
        for r in results:
            kid = r['id']
            with st.expander(f"📄 {_clean_knowledge_name(kid)} ({r['score']})"):
                st.markdown(r['text'][:1500])
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("🎲 出题", key=f"kb_s_{kid}", use_container_width=True):
                        st.session_state._kb_quiz = generate_review_questions([{"knowledge_id": kid}])
                        st.session_state._kb_qid = kid
                        st.rerun()
                with c2:
                    if st.button("🎓 概念自测", key=f"kb_s_cp_{kid}", use_container_width=True):
                        st.session_state._kb_concept_qid = kid
                        st.rerun()
                if st.session_state.get("_kb_qid") == kid:
                    quiz = st.session_state.get("_kb_quiz")
                    if st.session_state.get("_kb_result"):
                        st.markdown("### 评分结果")
                        st.markdown(_escape_md(_collapse_math(_fix_latex(st.session_state._kb_result))))
                        if st.button("关闭", key=f"kb_s_close_res_{kid}", use_container_width=True):
                            st.session_state.pop("_kb_result", None)
                            st.session_state.pop("_kb_qid", None)
                            st.rerun()
                    elif quiz and quiz.get("success"):
                        render_qa_cards(quiz['questions'], columns=1, typing=True)
                        ans = st.text_area("你的解法", key=f"kb_s_ans_{kid}", height=150,
                            placeholder="写下你的解题思路和答案...")
                        if st.button("提交自测", key=f"kb_s_sub_{kid}", use_container_width=True):
                            if ans.strip():
                                with st.spinner("AI 正在评分..."):
                                    try:
                                        prompt = PROBLEM_EVAL_PROMPT.format(question=quiz['questions'], answer=ans)
                                        result = call_llm_api(prompt, model="mimo-v2.5")
                                        total = 0; sc = 0; se = 0; sa = 0
                                        m = re.search(r'\[总分\]\s*(\d+)/(\d+)分', result)
                                        if m: total = int(m.group(1))
                                        m = re.search(r'\[解题正确性\]\s*(\d+)/(\d+)分', result)
                                        if m: sc = int(m.group(1))
                                        m = re.search(r'\[解题过程\]\s*(\d+)/(\d+)分', result)
                                        if m: se = int(m.group(1))
                                        m = re.search(r'\[书写真实性\]\s*(\d+)/(\d+)分', result)
                                        if m: sa = int(m.group(1))
                                        q_raw = quiz['questions']
                                        qm = re.search(r'Q:\s*(.+)', q_raw)
                                        display_q = qm.group(1).strip()[:300] if qm else q_raw[:300]
                                        save_feynman_record(st.session_state.get("user_id"), "problem", display_q, ans, result, sc, se, sa, total)
                                        st.session_state._kb_result = result
                                        st.session_state._kb_quiz = None
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"评价失败: {e}")
                            else:
                                st.warning("请输入你的解法")
                if st.session_state.get("_kb_concept_qid") == kid:
                    concept_quiz = st.session_state.get("_kb_concept_quiz")
                    if st.session_state.get("_kb_concept_result"):
                        st.markdown("### 评分结果")
                        st.markdown(_escape_md(_collapse_math(_fix_latex(st.session_state._kb_concept_result))))
                        if st.button("关闭", key=f"kb_s_close_cres_{kid}", use_container_width=True):
                            st.session_state.pop("_kb_concept_result", None)
                            st.session_state.pop("_kb_concept_qid", None)
                            st.rerun()
                    elif not st.session_state.get("_kb_concept_result"):
                        concept_quiz_text = f"概念自测：「{_clean_knowledge_name(kid)}」"
                        st.info(f"{concept_quiz_text}")
                        ans = st.text_area("你的回答", key=f"kb_s_ans_{kid}", height=120)
                        if st.button("提交自测", key=f"kb_s_cp_sub_{kid}", use_container_width=True):
                            if ans.strip():
                                with st.spinner("AI 正在评分..."):
                                    try:
                                        prompt = CONCEPT_EVAL_PROMPT.format(question=concept_quiz_text, answer=ans)
                                        result = call_llm_api(prompt, model="mimo-v2.5")
                                        total = 0; sc = 0; se = 0; sa = 0
                                        m = re.search(r'\[总分\]\s*(\d+)/(\d+)分', result)
                                        if m: total = int(m.group(1))
                                        m = re.search(r'\[概念理解\]\s*(\d+)/(\d+)分', result)
                                        if m: sc = int(m.group(1))
                                        m = re.search(r'\[表达能力\]\s*(\d+)/(\d+)分', result)
                                        if m: se = int(m.group(1))
                                        m = re.search(r'\[书写真实性\]\s*(\d+)/(\d+)分', result)
                                        if m: sa = int(m.group(1))
                                        save_feynman_record(st.session_state.get("user_id"), "concept", concept_quiz, ans, result, sc, se, sa, total)
                                        st.session_state._kb_concept_result = result
                                        st.session_state._kb_concept_quiz = None
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"评价失败: {e}")
                            else:
                                st.warning("请输入你的回答")

    else:
        for doc in filtered_corpus:
            kid = doc['id']
            with st.expander(f"📄 {_clean_knowledge_name(kid)}"):
                st.markdown(doc['text'][:1500])
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("🎲 出题", key=f"kb_d_{kid}", use_container_width=True):
                        st.session_state._kb_quiz = generate_review_questions([{"knowledge_id": kid}])
                        st.session_state._kb_qid = kid
                        st.rerun()
                with c2:
                    if st.button("🎓 概念自测", key=f"kb_d_cp_{kid}", use_container_width=True):
                        st.session_state._kb_concept_qid = kid
                        st.rerun()
                if st.session_state.get("_kb_qid") == kid:
                    quiz = st.session_state.get("_kb_quiz")
                    if st.session_state.get("_kb_result"):
                        st.markdown("### 评分结果")
                        st.markdown(_escape_md(_collapse_math(_fix_latex(st.session_state._kb_result))))
                        if st.button("关闭", key=f"kb_d_close_res_{kid}", use_container_width=True):
                            st.session_state.pop("_kb_result", None)
                            st.session_state.pop("_kb_qid", None)
                            st.rerun()
                    elif quiz and quiz.get("success"):
                        render_qa_cards(quiz['questions'], columns=1, typing=True)
                        ans = st.text_area("你的解法", key=f"kb_d_ans_{kid}", height=150,
                            placeholder="写下你的解题思路和答案...")
                        if st.button("提交自测", key=f"kb_d_sub_{kid}", use_container_width=True):
                            if ans.strip():
                                with st.spinner("AI 正在评分..."):
                                    try:
                                        prompt = PROBLEM_EVAL_PROMPT.format(question=quiz['questions'], answer=ans)
                                        result = call_llm_api(prompt, model="mimo-v2.5")
                                        total = 0; sc = 0; se = 0; sa = 0
                                        m = re.search(r'\[总分\]\s*(\d+)/(\d+)分', result)
                                        if m: total = int(m.group(1))
                                        m = re.search(r'\[解题正确性\]\s*(\d+)/(\d+)分', result)
                                        if m: sc = int(m.group(1))
                                        m = re.search(r'\[解题过程\]\s*(\d+)/(\d+)分', result)
                                        if m: se = int(m.group(1))
                                        m = re.search(r'\[书写真实性\]\s*(\d+)/(\d+)分', result)
                                        if m: sa = int(m.group(1))
                                        q_raw = quiz['questions']
                                        qm = re.search(r'Q:\s*(.+)', q_raw)
                                        display_q = qm.group(1).strip()[:300] if qm else q_raw[:300]
                                        save_feynman_record(st.session_state.get("user_id"), "problem", display_q, ans, result, sc, se, sa, total)
                                        st.session_state._kb_result = result
                                        st.session_state._kb_quiz = None
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"评价失败: {e}")
                            else:
                                st.warning("请输入你的解法")
                if st.session_state.get("_kb_concept_qid") == kid:
                    concept_quiz = st.session_state.get("_kb_concept_quiz")
                    if st.session_state.get("_kb_concept_result"):
                        st.markdown("### 评分结果")
                        st.markdown(_escape_md(_collapse_math(_fix_latex(st.session_state._kb_concept_result))))
                        if st.button("关闭", key=f"kb_d_close_cres_{kid}", use_container_width=True):
                            st.session_state.pop("_kb_concept_result", None)
                            st.session_state.pop("_kb_concept_qid", None)
                            st.rerun()
                    elif not st.session_state.get("_kb_concept_result"):
                        concept_quiz_text = f"概念自测：「{_clean_knowledge_name(kid)}」"
                        st.info(f"{concept_quiz_text}")
                        ans = st.text_area("你的回答", key=f"kb_d_cp_ans_{kid}", height=120)
                        if st.button("提交自测", key=f"kb_d_cp_sub_{kid}", use_container_width=True):
                            if ans.strip():
                                with st.spinner("AI 正在评分..."):
                                    try:
                                        prompt = CONCEPT_EVAL_PROMPT.format(question=concept_quiz_text, answer=ans)
                                        result = call_llm_api(prompt, model="mimo-v2.5")
                                        total = 0; sc = 0; se = 0; sa = 0
                                        m = re.search(r'\[总分\]\s*(\d+)/(\d+)分', result)
                                        if m: total = int(m.group(1))
                                        m = re.search(r'\[概念理解\]\s*(\d+)/(\d+)分', result)
                                        if m: sc = int(m.group(1))
                                        m = re.search(r'\[表达能力\]\s*(\d+)/(\d+)分', result)
                                        if m: se = int(m.group(1))
                                        m = re.search(r'\[书写真实性\]\s*(\d+)/(\d+)分', result)
                                        if m: sa = int(m.group(1))
                                        save_feynman_record(st.session_state.get("user_id"), "concept", concept_quiz, ans, result, sc, se, sa, total)
                                        st.session_state._kb_concept_result = result
                                        st.session_state._kb_concept_quiz = None
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"评价失败: {e}")
                            else:
                                st.warning("请输入你的回答")

with tab2:
    st.subheader("🎯 复习挑战")
    candidates = [c for c in get_review_candidates() if c['knowledge_id'] in filtered_corpus_ids]
    if candidates:
        for i, c in enumerate(candidates, 1):
            recall_pct = int(c['recall'] * 100)
            with st.expander(f"第{i}题: {_clean_knowledge_name(c['knowledge_id'])[:35]} (记忆: {recall_pct}%)"):
                knowledge_text = get_knowledge_text(c['knowledge_id'], corpus)
                st.markdown(knowledge_text[:1500])
                c1, c2, c3 = st.columns(3)
                with c1:
                    if st.button(f"掌握", key=f"rev_m_{i}"):
                        update_memory(c['knowledge_id'], True)
                        st.rerun()
                with c2:
                    if st.button(f"❌ 再练", key=f"rev_w_{i}"):
                        update_memory(c['knowledge_id'], False, error_type="遗忘")
                        st.rerun()
                with c3:
                    gen_key = f"rev_gen_{i}"
                    if st.button(f"🎲 出题", key=gen_key):
                        progress = st.progress(0, text="🎲 生成中...")
                        gen_r = generate_review_questions([{"knowledge_id": c['knowledge_id']}])
                        progress.progress(100, text="完成")
                        st.session_state._rev_quiz = gen_r
                        st.session_state._rev_quiz_id = i
                        st.rerun()

            if st.session_state.get("_rev_quiz_id") == i:
                quiz = st.session_state.pop("_rev_quiz", None)
                st.session_state.pop("_rev_quiz_id", None)
                if quiz and quiz.get("success"):
                    render_qa_cards(quiz['questions'], columns=1)

        if not candidates:
            st.success("🎉 暂无待复习知识点。使用问答后自动添加。")

with tab4:
    st.subheader("🎓 费曼学习法")
    st.info("选择模式 → 输入题目或上传图片 → 写下你的答案 → AI 评价")

    # 模式选择
    feynman_mode = st.radio("学习模式", ["概念理解", "解题练习"], horizontal=True, key="feynman_mode")
    mode_key = "concept" if feynman_mode == "概念理解" else "problem"

    # 出题方式
    col_input, col_img = st.columns(2)
    with col_input:
        feynman_question = st.text_area("输入题目", height=100,
            placeholder="例如：什么是洛必达法则？\n或：求函数 f(x)=x³-3x+2 的极值",
            key="feynman_question")
    with col_img:
        st.markdown("**或上传题目图片**")
        feynman_img = st.file_uploader("上传图片", type=["png", "jpg", "jpeg"], key="feynman_img",
            label_visibility="collapsed")
        if feynman_img is not None:
            img_b64 = base64.b64encode(feynman_img.getvalue()).decode()
            if st.button("识别图片文字", key="feynman_ocr", use_container_width=True):
                with st.spinner("识别中..."):
                    try:
                        ocr_data = {"model": "mimo-v2.5", "messages": [
                            {"role": "user", "content": [
                                {"type": "text", "text": "请识别这张图片中的数学题目内容，只输出题目文字。"},
                                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
                            ]}
                        ], "max_tokens": 1000, "temperature": 0}
                        req = urllib.request.Request(API_BASE + "/chat/completions",
                            data=json.dumps(ocr_data).encode("utf-8"),
                            headers={"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"},
                            method="POST")
                        with urllib.request.urlopen(req, timeout=30) as resp:
                            ocr_result = _extract_content(json.loads(resp.read().decode("utf-8"))["choices"][0]["message"])
                        st.session_state._feynman_ocr = ocr_result
                        st.rerun()
                    except Exception as e:
                        st.error(f"识别失败: {e}")
        if st.session_state.get("_feynman_ocr"):
            feynman_question = st.session_state.pop("_feynman_ocr")
            st.success("已识别图片文字，请在左侧确认题目内容")

    # 用户答案
    feynman_answer = st.text_area("你的答案", height=200,
        placeholder="用自己的话写下答案...\n\n提示：尽量用自己的语言表达，展示你的理解过程。",
        key="feynman_answer")

    # 提交按钮
    if st.button("提交答案", key="feynman_submit", use_container_width=True, type="primary"):
        if not feynman_question.strip():
            st.warning("请输入题目内容")
        elif not feynman_answer.strip():
            st.warning("请输入你的答案")
        else:
            with st.spinner("AI 正在评价..."):
                try:
                    if mode_key == "concept":
                        eval_prompt = CONCEPT_EVAL_PROMPT
                    else:
                        eval_prompt = PROBLEM_EVAL_PROMPT

                    prompt = eval_prompt.format(question=feynman_question, answer=feynman_answer)
                    result = call_llm_api(prompt, model="mimo-v2.5")

                    # 解析分数
                    score_correct = 0
                    score_expression = 0
                    score_authentic = 0
                    total_score = 0

                    score_match = re.search(r'\[总分\]\s*(\d+)/(\d+)分', result)
                    if score_match:
                        total_score = int(score_match.group(1))

                    correct_match = re.search(r'\[(?:概念理解|解题正确性)\]\s*(\d+)/(\d+)分', result)
                    if correct_match:
                        score_correct = int(correct_match.group(1))

                    expr_match = re.search(r'\[(?:表达能力|解题过程)\]\s*(\d+)/(\d+)分', result)
                    if expr_match:
                        score_expression = int(expr_match.group(1))

                    auth_match = re.search(r'\[书写真实性\]\s*(\d+)/(\d+)分', result)
                    if auth_match:
                        score_authentic = int(auth_match.group(1))

                    # 保存记录
                    save_feynman_record(
                        st.session_state.get("user_id"),
                        mode_key,
                        feynman_question,
                        feynman_answer,
                        result,
                        score_correct,
                        score_expression,
                        score_authentic,
                        total_score
                    )

                    # 显示结果
                    st.markdown("---")
                    st.markdown("### 评价结果")
                    st.markdown(_escape_md(_collapse_math(_fix_latex(result))))

                except Exception as e:
                    st.error(f"评价失败: {e}")

    # 历史记录
    st.markdown("---")
    st.markdown("### 📜 历史记录")
    feynman_history = get_feynman_history(st.session_state.get("user_id", 1))
    if feynman_history:
        for record in feynman_history:
            mode_label = "概念" if record["mode"] == "concept" else "解题"
            score = record["total_score"]
            time_str = str(record["created_at"])[:16]
            q_text = record['question_text'] or '(概念自测)'
            with st.expander(f"[{mode_label}] {q_text[:40]}... | {score}分 | {time_str}"):
                st.markdown(f"**题目**: {q_text}")
                st.markdown(f"**你的答案**: {record['user_answer']}")
                st.markdown("---")
                st.markdown(record["ai_evaluation"])
    else:
        st.info("暂无记录，开始你的第一次练习吧！")

with tab3:
    st.subheader("知识点掌握情况")
    conn = sqlite3.connect(MEMORY_DB)
    c = conn.cursor()
    c.execute("SELECT knowledge_id, status, times_correct, times_wrong, stability FROM knowledge_mastery WHERE user_id=? ORDER BY last_review DESC", (st.session_state.get("user_id", 1),))
    rows = c.fetchall()
    conn.close()
    filtered_rows = [r for r in rows if r[0] in filtered_corpus_ids]
    mastered = sum(1 for r in filtered_rows if r[1] == "掌握")
    learning = sum(1 for r in filtered_rows if r[1] == "学习中")
    total = len(filtered_rows)
    progress = mastered / max(total, 1)
    st.progress(progress)
    st.markdown(f"**掌握进度**: {mastered}/{total} ({progress*100:.1f}%)")

    for r in filtered_rows:
        name = r[0]
        if len(name) > 30:
            name = name[:27] + "..."
        if r[1] == "掌握":
            st.markdown(f"<div class='mastered-card'>{name} | ✓{r[2]} ✗{r[3]}</div>", unsafe_allow_html=True)
        elif r[1] == "学习中":
            st.markdown(f"<div class='learning-card'>{name} | ✓{r[2]} ✗{r[3]}</div>", unsafe_allow_html=True)


